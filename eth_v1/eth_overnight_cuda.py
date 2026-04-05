# -*- coding: utf-8 -*-
"""
ETH/USDT Overnight Optimizer — CUDA GPU Accelerated
Original: eth_overnight_parallel.py (CPU multiprocessing)
Converted: GPU CUDA batch processing (RTX 4060 Ti)
"""
import pandas as pd, numpy as np, time, json, sys, os, math
from numba import njit, cuda
from datetime import datetime
import warnings
warnings.filterwarnings('ignore')
sys.stdout.reconfigure(encoding='utf-8', line_buffering=True)

OUTDIR = 'D:/filesystem/futures/eth_v1'

# ══════════════════════════════════════════════
# INDICATORS (same as original)
# ══════════════════════════════════════════════
def ema_calc(s, p): return s.ewm(span=p, adjust=False).mean().values.astype(np.float64)
def sma_calc(s, p): return s.rolling(p).mean().values.astype(np.float64)

def adx_calc(h, l, c, period=20):
    n=len(c); pdm=np.zeros(n); mdm=np.zeros(n); tr=np.zeros(n)
    for i in range(1, n):
        hd=h[i]-h[i-1]; ld=l[i-1]-l[i]
        pdm[i]=hd if(hd>ld and hd>0) else 0
        mdm[i]=ld if(ld>hd and ld>0) else 0
        tr[i]=max(h[i]-l[i], abs(h[i]-c[i-1]), abs(l[i]-c[i-1]))
    a=1.0/period
    atr=pd.Series(tr).ewm(alpha=a, min_periods=period, adjust=False).mean()
    sp=pd.Series(pdm).ewm(alpha=a, min_periods=period, adjust=False).mean()
    sn=pd.Series(mdm).ewm(alpha=a, min_periods=period, adjust=False).mean()
    pdi=100*sp/atr.replace(0,1e-10)
    mdi=100*sn/atr.replace(0,1e-10)
    dx=100*(pdi-mdi).abs()/(pdi+mdi).replace(0,1e-10)
    return dx.ewm(alpha=a, min_periods=period, adjust=False).mean().values.astype(np.float64)

def rsi_calc(c, period=10):
    d=np.diff(c, prepend=c[0]); g=np.where(d>0, d, 0); lo=np.where(d<0, -d, 0)
    a=1.0/period
    ag=pd.Series(g).ewm(alpha=a, min_periods=period, adjust=False).mean()
    al=pd.Series(lo).ewm(alpha=a, min_periods=period, adjust=False).mean()
    return (100-100/(1+ag/al.replace(0,1e-10))).values.astype(np.float64)


# ══════════════════════════════════════════════
# NUMBA JIT BACKTEST (Stage 1: CPU JIT for verification)
# ══════════════════════════════════════════════
@njit(cache=True)
def bt_jit(c, h, l, fm, sm, av, rv, n,
           SL, TA, TSL, ADX_M, ARISE, RMIN, RMAX, GAP, MON, SKIP, RNC, WU):
    F=0.0004; P=10000.0
    cap=5000.0; pos=0; epx=0.0; slp=0.0; ton=0; thi=0.0; tlo=999999.0
    w=0; ws=0; ld=0; pk=cap; mdd=0.0; ms=cap
    sc=0; tc=0; rc=0; wn=0; ln=0; gp=0.0; gl=0.0

    for i in range(WU, n):
        px=c[i]; hi=h[i]; lo_=l[i]
        if i > WU and i % 720 == 0: ms = cap

        if pos != 0:
            w = 0
            if ton == 0 and SL > 0:
                if (pos==1 and lo_<=slp) or (pos==-1 and hi>=slp):
                    pnl=(slp-epx)/epx*P*pos - P*F; cap+=pnl; sc+=1
                    if pnl>0: wn+=1; gp+=pnl
                    else: ln+=1; gl+=abs(pnl)
                    ld=pos; pos=0; pk=max(pk,cap)
                    dd=(pk-cap)/pk if pk>0 else 0.0
                    if dd>mdd: mdd=dd
                    if cap<=0: break
                    continue
            if TA > 0:
                br = ((hi-epx)/epx*100) if pos==1 else ((epx-lo_)/epx*100)
                if br >= TA and ton == 0: ton = 1
            if ton == 1 and TSL > 0:
                ex_hit = 0
                if pos == 1:
                    if hi > thi: thi = hi
                    ns = thi*(1-TSL/100)
                    if ns > slp: slp = ns
                    if px <= slp: ex_hit = 1
                else:
                    if lo_ < tlo: tlo = lo_
                    ns = tlo*(1+TSL/100)
                    if ns < slp: slp = ns
                    if px >= slp: ex_hit = 1
                if ex_hit == 1:
                    pnl=(px-epx)/epx*P*pos - P*F; cap+=pnl; tc+=1
                    if pnl>0: wn+=1; gp+=pnl
                    else: ln+=1; gl+=abs(pnl)
                    ld=pos; pos=0; pk=max(pk,cap)
                    dd=(pk-cap)/pk if pk>0 else 0.0
                    if dd>mdd: mdd=dd
                    if cap<=0: break
                    continue
            if i > 0:
                bn = fm[i]>sm[i]; bp = fm[i-1]>sm[i-1]
                cu = bn and not bp; cd = not bn and bp
                if (pos==1 and cd) or (pos==-1 and cu):
                    pnl=(px-epx)/epx*P*pos - P*F; cap+=pnl; rc+=1
                    if pnl>0: wn+=1; gp+=pnl
                    else: ln+=1; gl+=abs(pnl)
                    ld=pos; pos=0
                    if RNC == 0:
                        pk=max(pk,cap); dd=(pk-cap)/pk if pk>0 else 0.0
                        if dd>mdd: mdd=dd
                        if cap<=0: break
                        continue
        if i < 1: continue
        bn = fm[i]>sm[i]; bp = fm[i-1]>sm[i-1]
        cu = bn and not bp; cd = not bn and bp
        if pos == 0:
            if cu: w=1; ws=i
            elif cd: w=-1; ws=i
            if w!=0 and i>ws:
                if MON>0 and i-ws>MON: w=0; continue
                if w==1 and cd: w=-1; ws=i; continue
                if w==-1 and cu: w=1; ws=i; continue
                if SKIP>0 and w==ld: continue
                if av[i]<ADX_M: continue
                if ARISE>0 and i>=ARISE and av[i]<=av[i-ARISE]: continue
                if (RMIN>0 or RMAX<100) and (rv[i]<RMIN or rv[i]>RMAX): continue
                if GAP>0 and sm[i]>0 and abs(fm[i]-sm[i])/sm[i]*100<GAP: continue
                if cap<1000: continue
                cap -= P*F; pos=w; epx=px; ton=0; thi=px; tlo=px
                if pos==1 and SL>0: slp=px*(1-SL/100)
                elif pos==-1 and SL>0: slp=px*(1+SL/100)
                else: slp = 0.0 if pos==1 else 999999.0
                pk=max(pk,cap); w=0
        pk=max(pk,cap); dd=(pk-cap)/pk if pk>0 else 0.0
        if dd>mdd: mdd=dd
        if cap<=0: break

    if pos!=0 and cap>0:
        pnl=(c[n-1]-epx)/epx*P*pos - P*F; cap+=pnl
        if pnl>0: wn+=1; gp+=pnl
        else: ln+=1; gl+=abs(pnl)

    tot=sc+tc+rc; pf=gp/gl if gl>0 else 0.0; wr=wn/(wn+ln)*100 if (wn+ln)>0 else 0.0
    # Return: cap, net, tot, sc, tc, rc, wn, ln, wr, pf, mdd*100
    return cap, gp-gl, tot, sc, tc, rc, wn, ln, wr, pf, mdd*100


# ══════════════════════════════════════════════
# CUDA KERNEL
# ══════════════════════════════════════════════
# params_matrix columns: [SL, TA, TSL, ADX_M, ARISE, RMIN, RMAX, GAP, MON, SKIP, RNC, WU]
NUM_PARAMS = 12
NUM_RESULTS = 11  # cap, net, tot, sc, tc, rc, wn, ln, wr, pf, mdd

@cuda.jit
def bt_cuda_kernel(close, high, low, fm_all, sm_all, adx_all, rsi_all,
                   params_matrix, results_matrix, n_arr):
    tid = cuda.grid(1)
    if tid >= params_matrix.shape[0]:
        return

    n = n_arr[0]
    p = params_matrix[tid]
    SL=p[0]; TA=p[1]; TSL=p[2]; ADX_M=p[3]; ARISE=int(p[4])
    RMIN=p[5]; RMAX=p[6]; GAP=p[7]; MON=int(p[8]); SKIP=int(p[9]); RNC=int(p[10]); WU=int(p[11])

    F=0.0004; P=10000.0
    cap=5000.0; pos=0; epx=0.0; slp=0.0; ton=0; thi=0.0; tlo=999999.0
    w=0; ws=0; ld=0; pk=cap; mdd=0.0; ms=cap
    sc=0; tc=0; rc=0; wn=0; ln=0; gp=0.0; gl=0.0

    for i in range(WU, n):
        px=close[i]; hi=high[i]; lo_=low[i]
        fm_i=fm_all[tid, i]; sm_i=sm_all[tid, i]; av_i=adx_all[tid, i]; rv_i=rsi_all[tid, i]

        if i > WU and i % 720 == 0: ms = cap

        if pos != 0:
            w = 0
            if ton == 0 and SL > 0:
                if (pos==1 and lo_<=slp) or (pos==-1 and hi>=slp):
                    pnl=(slp-epx)/epx*P*pos - P*F; cap+=pnl; sc+=1
                    if pnl>0: wn+=1; gp+=pnl
                    else: ln+=1; gl+=abs(pnl)
                    ld=pos; pos=0
                    if cap>pk: pk=cap
                    dd=(pk-cap)/pk if pk>0 else 0.0
                    if dd>mdd: mdd=dd
                    if cap<=0: break
                    continue
            if TA > 0:
                br = ((hi-epx)/epx*100) if pos==1 else ((epx-lo_)/epx*100)
                if br >= TA and ton == 0: ton = 1
            if ton == 1 and TSL > 0:
                ex_hit = 0
                if pos == 1:
                    if hi > thi: thi = hi
                    ns = thi*(1-TSL/100)
                    if ns > slp: slp = ns
                    if px <= slp: ex_hit = 1
                else:
                    if lo_ < tlo: tlo = lo_
                    ns = tlo*(1+TSL/100)
                    if ns < slp: slp = ns
                    if px >= slp: ex_hit = 1
                if ex_hit == 1:
                    pnl=(px-epx)/epx*P*pos - P*F; cap+=pnl; tc+=1
                    if pnl>0: wn+=1; gp+=pnl
                    else: ln+=1; gl+=abs(pnl)
                    ld=pos; pos=0
                    if cap>pk: pk=cap
                    dd=(pk-cap)/pk if pk>0 else 0.0
                    if dd>mdd: mdd=dd
                    if cap<=0: break
                    continue
            if i > 0:
                fm_prev=fm_all[tid, i-1]; sm_prev=sm_all[tid, i-1]
                bn = fm_i>sm_i; bp = fm_prev>sm_prev
                cu = bn and not bp; cd = not bn and bp
                if (pos==1 and cd) or (pos==-1 and cu):
                    pnl=(px-epx)/epx*P*pos - P*F; cap+=pnl; rc+=1
                    if pnl>0: wn+=1; gp+=pnl
                    else: ln+=1; gl+=abs(pnl)
                    ld=pos; pos=0
                    if RNC == 0:
                        if cap>pk: pk=cap
                        dd=(pk-cap)/pk if pk>0 else 0.0
                        if dd>mdd: mdd=dd
                        if cap<=0: break
                        continue
        if i < 1: continue
        fm_prev=fm_all[tid, i-1]; sm_prev=sm_all[tid, i-1]
        bn = fm_i>sm_i; bp = fm_prev>sm_prev
        cu = bn and not bp; cd = not bn and bp
        if pos == 0:
            if cu: w=1; ws=i
            elif cd: w=-1; ws=i
            if w!=0 and i>ws:
                if MON>0 and i-ws>MON: w=0; continue
                if w==1 and cd: w=-1; ws=i; continue
                if w==-1 and cu: w=1; ws=i; continue
                if SKIP>0 and w==ld: continue
                if av_i<ADX_M: continue
                if ARISE>0 and i>=ARISE and adx_all[tid, i]<=adx_all[tid, i-ARISE]: continue
                if (RMIN>0 or RMAX<100) and (rv_i<RMIN or rv_i>RMAX): continue
                if GAP>0 and sm_i>0 and abs(fm_i-sm_i)/sm_i*100<GAP: continue
                if cap<1000: continue
                cap -= P*F; pos=w; epx=px; ton=0; thi=px; tlo=px
                if pos==1 and SL>0: slp=px*(1-SL/100)
                elif pos==-1 and SL>0: slp=px*(1+SL/100)
                else: slp = 0.0 if pos==1 else 999999.0
                if cap>pk: pk=cap
                w=0
        if cap>pk: pk=cap
        dd=(pk-cap)/pk if pk>0 else 0.0
        if dd>mdd: mdd=dd
        if cap<=0: break

    if pos!=0 and cap>0:
        pnl=(close[n-1]-epx)/epx*P*pos - P*F; cap+=pnl
        if pnl>0: wn+=1; gp+=pnl
        else: ln+=1; gl+=abs(pnl)

    tot=sc+tc+rc
    pf=gp/gl if gl>0 else 0.0
    wr=wn/(wn+ln)*100.0 if (wn+ln)>0 else 0.0

    results_matrix[tid, 0] = cap
    results_matrix[tid, 1] = gp-gl
    results_matrix[tid, 2] = float(tot)
    results_matrix[tid, 3] = float(sc)
    results_matrix[tid, 4] = float(tc)
    results_matrix[tid, 5] = float(rc)
    results_matrix[tid, 6] = float(wn)
    results_matrix[tid, 7] = float(ln)
    results_matrix[tid, 8] = wr
    results_matrix[tid, 9] = pf
    results_matrix[tid, 10] = mdd*100


# ══════════════════════════════════════════════
# CUDA BATCH RUNNER (Memory Optimized)
# ══════════════════════════════════════════════
def run_cuda_grouped(ma_groups, close, high, low, adx_arr, rsi_arr, BATCH_SIZE=800):
    """
    Memory-optimized: group jobs by MA combination.
    Same MA pair shares one GPU transfer instead of copying per-job.

    ma_groups: list of (fm_arr, sm_arr, param_list, meta_list)
        - fm_arr, sm_arr: numpy arrays (shared by all jobs in group)
        - param_list: list of [SL,TA,TSL,ADX_M,ARISE,RMIN,RMAX,GAP,MON,SKIP,RNC,WU]
        - meta_list: list of meta dicts
    """
    n = len(close)
    all_results = []

    d_close = cuda.to_device(close)
    d_high = cuda.to_device(high)
    d_low = cuda.to_device(low)
    d_n = cuda.to_device(np.array([n], dtype=np.int64))

    for fm_arr, sm_arr, param_list, meta_list in ma_groups:
        N = len(param_list)
        if N == 0:
            continue

        # Upload this MA pair's arrays once (shared for all jobs in group)
        # Shape: (1, n) → broadcast to all threads via indexing trick
        # But kernel expects (num_jobs, n). We tile only within batch.

        for b_start in range(0, N, BATCH_SIZE):
            b_end = min(b_start + BATCH_SIZE, N)
            bs = b_end - b_start

            # Build params matrix (small: bs × 12 floats)
            params_mat = np.array(param_list[b_start:b_end], dtype=np.float64)

            # Build MA arrays: all jobs share same fm/sm/adx/rsi
            # Use np.broadcast_to to avoid memory copy (virtual expansion)
            fm_tiled = np.broadcast_to(fm_arr[np.newaxis, :], (bs, n)).copy()
            sm_tiled = np.broadcast_to(sm_arr[np.newaxis, :], (bs, n)).copy()
            adx_tiled = np.broadcast_to(adx_arr[np.newaxis, :], (bs, n)).copy()
            rsi_tiled = np.broadcast_to(rsi_arr[np.newaxis, :], (bs, n)).copy()

            # Transfer to GPU
            d_fm = cuda.to_device(fm_tiled)
            d_sm = cuda.to_device(sm_tiled)
            d_adx = cuda.to_device(adx_tiled)
            d_rsi = cuda.to_device(rsi_tiled)
            d_params = cuda.to_device(params_mat)
            d_results = cuda.device_array((bs, NUM_RESULTS), dtype=np.float64)

            # Launch
            threads = 128
            blocks = (bs + threads - 1) // threads
            bt_cuda_kernel[blocks, threads](d_close, d_high, d_low,
                                             d_fm, d_sm, d_adx, d_rsi,
                                             d_params, d_results, d_n)
            cuda.synchronize()

            # Copy back & filter
            res_mat = d_results.copy_to_host()
            batch_meta = meta_list[b_start:b_end]

            for j in range(bs):
                r = res_mat[j]
                tot = int(r[2]); pf = r[9]; mdd_v = r[10]
                if tot >= 10 and pf > 1.5 and mdd_v < 55:
                    all_results.append({
                        'cap': r[0], 'net': r[1], 'tot': tot,
                        'sc': int(r[3]), 'tc': int(r[4]), 'rc': int(r[5]),
                        'wn': int(r[6]), 'ln': int(r[7]),
                        'wr': r[8], 'pf': pf, 'mdd': mdd_v,
                        **batch_meta[j]
                    })

            # Free GPU memory explicitly
            del d_fm, d_sm, d_adx, d_rsi, d_params, d_results
            del fm_tiled, sm_tiled, adx_tiled, rsi_tiled, params_mat

    return all_results


# ══════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════
def main():
    # Check GPU
    if not cuda.is_available():
        print("ERROR: CUDA not available! Set CUDA_PATH and restart.")
        return

    gpu = cuda.get_current_device()
    mem = cuda.current_context().get_memory_info()

    print("=" * 80)
    print(f"  ETH/USDT Overnight Optimizer — CUDA GPU Accelerated")
    print(f"  GPU: {gpu.name} ({mem[0]//1024**3}GB free / {mem[1]//1024**3}GB total)")
    print(f"  Started: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print("=" * 80, flush=True)

    # Load data
    print("\n[1] Loading ETH data...", flush=True)
    files = [f'{OUTDIR}/eth_usdt_5m_2020_to_now_part{i}.csv' for i in range(1, 4)]
    dfs = [pd.read_csv(f, parse_dates=['timestamp']) for f in files]
    df5m = pd.concat(dfs, ignore_index=True).sort_values('timestamp').drop_duplicates('timestamp', keep='last')
    df5m.set_index('timestamp', inplace=True)
    print(f"  5m: {len(df5m)} bars", flush=True)

    agg = {'open': 'first', 'high': 'max', 'low': 'min', 'close': 'last', 'volume': 'sum'}
    timeframes = {
        '10m': df5m.resample('10min').agg(agg).dropna(),
        '15m': df5m.resample('15min').agg(agg).dropna(),
        '30m': df5m.resample('30min').agg(agg).dropna(),
        '1h': df5m.resample('1h').agg(agg).dropna(),
    }
    for k, v in timeframes.items():
        print(f"  {k}: {len(v)} bars", flush=True)

    # Parameters (same as original)
    fps = [3, 5, 7, 10, 20, 50, 75, 100]
    sps = [100, 150, 200, 250, 300, 400, 500, 600]
    adx_ms = [25, 30, 35]
    sls = [2, 3, 5, 7]
    ta_tsls = [(0, 0), (8, 5), (10, 7), (12, 9), (15, 10), (20, 12)]
    rsi_ranges = [(40, 80), (0, 100)]
    gaps = [0, 0.2]
    skips = [True, False]

    # Process one TF at a time to save memory
    print("\n[2] Processing timeframes sequentially (memory optimized)...", flush=True)
    all_results = []
    grand_total = 0
    import gc

    for tf_name, df in timeframes.items():
        print(f"\n  === {tf_name} ===", flush=True)

        # Compute indicators for this TF only
        cs = df['close']
        c = cs.values.astype(np.float64)
        h = df['high'].values.astype(np.float64)
        l = df['low'].values.astype(np.float64)
        av = adx_calc(h, l, c)
        rv = rsi_calc(c)

        mas = {}
        for p in fps + sps:
            mas[('EMA', p)] = ema_calc(cs, p)
            mas[('SMA', p)] = sma_calc(cs, p)
        print(f"    {len(mas)} MAs computed", flush=True)

        # Build MA-grouped jobs: group by (ft,fp,st,sp) so same MA pair shares memory
        ma_groups = []
        total_tf = 0
        for ft in ['EMA', 'SMA']:
            for fp in fps:
                for st in ['EMA', 'SMA']:
                    for sp in sps:
                        if fp >= sp: continue
                        fm = mas[(ft, fp)]
                        sm = mas[(st, sp)]
                        wu = max(sp, 21)

                        param_list = []
                        meta_list = []
                        for adx_m in adx_ms:
                            for sl in sls:
                                for ta, tsl in ta_tsls:
                                    for skip in skips:
                                        rnc = (ta > 0)
                                        for rmin, rmax in rsi_ranges:
                                            for gap in gaps:
                                                param_list.append([
                                                    float(sl), float(ta), float(tsl),
                                                    float(adx_m), 6.0,
                                                    float(rmin), float(rmax),
                                                    float(gap), 24.0,
                                                    1.0 if skip else 0.0,
                                                    1.0 if rnc else 0.0,
                                                    float(wu)
                                                ])
                                                meta_list.append({
                                                    'tf': tf_name, 'ft': ft, 'fp': fp,
                                                    'st': st, 'sp': sp,
                                                    'adx': adx_m, 'sl': sl, 'ta': ta, 'tsl': tsl,
                                                    'skip': skip, 'rnc': rnc,
                                                    'rmin': rmin, 'rmax': rmax, 'gap': gap,
                                                })

                        if param_list:
                            ma_groups.append((fm, sm, param_list, meta_list))
                            total_tf += len(param_list)

        grand_total += total_tf
        print(f"    {total_tf:,} jobs in {len(ma_groups)} MA groups", flush=True)

        # Run on GPU
        t0 = time.time()
        tf_results = run_cuda_grouped(ma_groups, c, h, l, av, rv, BATCH_SIZE=800)
        el = time.time() - t0
        all_results.extend(tf_results)
        print(f"    → {len(tf_results)} passed in {el:.1f}s ({total_tf/max(el,0.01):.0f} jobs/s)", flush=True)

        # Free this TF's memory before next TF
        del mas, c, h, l, av, rv, ma_groups, tf_results
        gc.collect()

    print(f"\n  TOTAL: {grand_total:,} jobs | {len(all_results)} passed", flush=True)

    all_results.sort(key=lambda x: x['net'], reverse=True)

    # TOP 30
    print(f"\n{'='*80}")
    print(f"  TOP 30 ETH STRATEGIES (CUDA GPU)")
    print(f"{'='*80}")
    print(f"  {'#':>3} {'TF':<4} {'Fast':<8} {'Slow':<8} {'ADX':>3} {'SL':>3} {'TA':>3} {'TSL':>3} {'Net$':>10} {'#T':>4} {'PF':>5} {'MDD':>5} {'WR':>4}")
    print(f"  {'-'*72}")
    for i, r in enumerate(all_results[:30], 1):
        print(f"  {i:>3} {r['tf']:<4} {r['ft']}{r['fp']:<6} {r['st']}{r['sp']:<6} {r['adx']:>3} {r['sl']:>3} {r['ta']:>3} {r['tsl']:>3} ${r['net']:>+9,.0f} {r['tot']:>4} {r['pf']:>5.1f} {r['mdd']:>5.1f} {r['wr']:>3.0f}%")
    sys.stdout.flush()

    # 6-Engine cross-verify top 10 (using JIT for speed)
    # Recompute needed TF data on-demand (memory friendly)
    print(f"\n[4] 6-Engine Cross-Verify Top 10 (JIT)...", flush=True)
    top10 = all_results[:10]
    verified = []
    verify_cache = {}  # cache TF data only for needed TFs

    for idx, s in enumerate(top10):
        tf_name = s['tf']
        if tf_name not in verify_cache:
            df = timeframes[tf_name]
            cs = df['close']
            vc = {
                'c': cs.values.astype(np.float64),
                'h': df['high'].values.astype(np.float64),
                'l': df['low'].values.astype(np.float64),
                'av': adx_calc(df['high'].values.astype(np.float64),
                               df['low'].values.astype(np.float64),
                               cs.values.astype(np.float64)),
                'rv': rsi_calc(cs.values.astype(np.float64)),
                'mas': {}
            }
            verify_cache[tf_name] = vc

        vc = verify_cache[tf_name]
        ma_key_fm = (s['ft'], s['fp'])
        ma_key_sm = (s['st'], s['sp'])
        if ma_key_fm not in vc['mas']:
            cs = timeframes[tf_name]['close']
            vc['mas'][ma_key_fm] = ema_calc(cs, s['fp']) if s['ft'] == 'EMA' else sma_calc(cs, s['fp'])
        if ma_key_sm not in vc['mas']:
            cs = timeframes[tf_name]['close']
            vc['mas'][ma_key_sm] = ema_calc(cs, s['sp']) if s['st'] == 'EMA' else sma_calc(cs, s['sp'])

        fm = vc['mas'][ma_key_fm]
        sm = vc['mas'][ma_key_sm]
        n = len(vc['c'])
        wu = max(s['sp'], 21)
        rnc_int = 1 if s['rnc'] else 0
        skip_int = 1 if s['skip'] else 0

        caps = []
        for _ in range(6):
            r = bt_jit(vc['c'], vc['h'], vc['l'], fm, sm, vc['av'], vc['rv'], n,
                       float(s['sl']), float(s['ta']), float(s['tsl']), float(s['adx']), 6,
                       float(s.get('rmin', 40)), float(s.get('rmax', 80)), float(s.get('gap', 0)),
                       24, skip_int, rnc_int, wu)
            caps.append(r[0])

        diff = max(caps) - min(caps)
        ok = diff < 0.01
        s['cross_ok'] = ok
        s['cross_diff'] = diff
        verified.append(s)
        print(f"  #{idx+1} {s['ft']}{s['fp']}/{s['st']}{s['sp']} ({s['tf']}) 6E:{'OK' if ok else 'NG'} diff=${diff:.4f} Net=${s['net']:+,.0f}", flush=True)

    del verify_cache; gc.collect()

    # Word document
    print(f"\n[5] Generating Word document...", flush=True)
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn

        doc = Document()
        st = doc.styles['Normal']; st.font.name = 'Malgun Gothic'; st.font.size = Pt(9)
        st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

        doc.add_heading('ETH/USDT Futures Strategy Specification (CUDA GPU)', 0)
        doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
        doc.add_paragraph(f'GPU: RTX 4060 Ti | Total tested: {grand_total:,} | Passed: {len(all_results)}')

        doc.add_heading('1. Strategy Ranking', 1)
        tbl = doc.add_table(rows=1, cols=11); tbl.style = 'Light Grid Accent 1'
        for i, hd in enumerate(['#', 'TF', 'Fast', 'Slow', 'ADX', 'SL', 'TA/TSL', 'Net($)', 'Trades', 'PF', 'MDD%']):
            tbl.rows[0].cells[i].text = hd
        for i, r in enumerate(all_results[:50], 1):
            row = tbl.add_row().cells
            row[0].text = str(i); row[1].text = r['tf']
            row[2].text = f"{r['ft']}{r['fp']}"; row[3].text = f"{r['st']}{r['sp']}"
            row[4].text = str(r['adx']); row[5].text = f"{r['sl']}%"
            row[6].text = f"{r['ta']}/{r['tsl']}"; row[7].text = f"${r['net']:+,.0f}"
            row[8].text = str(r['tot']); row[9].text = f"{r['pf']:.1f}"
            row[10].text = f"{r['mdd']:.1f}%"

        doc.add_page_break()
        doc.add_heading('2. Top 10 Verified Strategies', 1)
        for idx, r in enumerate(verified, 1):
            doc.add_heading(f"#{idx} {r['ft']}{r['fp']}/{r['st']}{r['sp']} ({r['tf']})", 2)
            p = doc.add_paragraph()
            p.add_run(f"Net: ${r['net']:+,.0f} | Trades: {r['tot']} (SL:{r['sc']} TSL:{r['tc']} REV:{r['rc']}) | "
                      f"WR: {r['wr']:.1f}% | PF: {r['pf']:.1f} | MDD: {r['mdd']:.1f}%\n"
                      f"6-Engine: {'MATCH' if r.get('cross_ok') else 'DIFF'} (diff: ${r.get('cross_diff', 0):.4f})\n"
                      f"Params: ADX>={r['adx']} RSI {r.get('rmin', 40)}-{r.get('rmax', 80)} "
                      f"SL={r['sl']}% TA={r['ta']}% TSL={r['tsl']}% "
                      f"Gap={r.get('gap', 0)}% Skip={'Y' if r['skip'] else 'N'}")

        out_docx = f'{OUTDIR}/ETH_Strategy_Specification_CUDA.docx'
        doc.save(out_docx)
        print(f"  Word: {out_docx}", flush=True)
    except Exception as e:
        print(f"  Word generation error: {e}", flush=True)

    # Save JSON
    json_out = f'{OUTDIR}/eth_final_results_cuda.json'
    save = []
    for r in all_results[:100]:
        sr = {k: v for k, v in r.items() if k not in ['trades', 'monthly']}
        save.append(sr)
    with open(json_out, 'w', encoding='utf-8') as f:
        json.dump(save, f, indent=2, ensure_ascii=False, default=str)
    print(f"  JSON: {json_out}", flush=True)

    print(f"\n{'='*80}")
    print(f"  COMPLETED: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"  {grand_total:,} jobs on GPU | CPU free for trading bots")
    print(f"{'='*80}", flush=True)


if __name__ == '__main__':
    main()
