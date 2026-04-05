"""
v28 전체 기획서 6엔진 교차검증 결과 리뷰
==========================================
25개 전략 x 6엔진 교차검증 + 연도별 상세 분석 + Word 문서 생성

Phase 1: 데이터 로드 및 리샘플링 (5m, 10m, 15m, 30m, 1h)
Phase 2: 25개 전략 x 6엔진 백테스트
Phase 3: Engine 1 상세 백테스트 (연도별/월별 분석)
Phase 4: Word 문서 생성 (수익률 BEST 10, 안정형 BEST 10, 폐기형 BEST 10)
"""

import os
import sys
import time
import numpy as np
import pandas as pd
import warnings
warnings.filterwarnings('ignore')

os.environ['PYTHONUNBUFFERED'] = '1'

# =============================================================================
# Imports from existing modules
# =============================================================================
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from v28_backtest_engine import (
    load_5m_data, resample_ohlcv,
    calc_ma, calc_adx_wilder, calc_rsi_wilder,
    calc_ema, calc_wma, calc_sma, calc_hma, calc_vwma,
    backtest_core,
)

from v28_detailed_verify import detailed_backtest

from numba import njit

BASE = os.path.dirname(os.path.abspath(__file__))
FEE_RATE = 0.0004  # 0.04% x 2 sides

# =============================================================================
# 25 strategies
# =============================================================================
strategies = [
    # ft: fast MA type (0=EMA, 1=WMA, 2=SMA, 3=HMA, 4=VWMA)
    # st: slow MA type
    # fp/sp: fast/slow period
    # ap: ADX period, at: ADX threshold
    # rl/rh: RSI low/high
    # sl: stop loss %, ta: trail activation %, tp: trail pct %
    # m: margin %, lv: leverage
    # d: entry delay bars, o: entry offset pct, sk: skip same dir
    # cap: initial capital
    {"name": "v12.3",   "tf": "5m",  "ft": 0, "fp": 7,   "st": 0, "sp": 100, "ap": 14, "at": 30, "rl": 30, "rh": 58, "sl": -9,  "ta": 8,  "tp": 6, "m": 20, "lv": 10, "d": 0, "o": 0.0,  "sk": 0, "cap": 3000},
    {"name": "v13.5",   "tf": "5m",  "ft": 0, "fp": 7,   "st": 0, "sp": 100, "ap": 14, "at": 30, "rl": 30, "rh": 58, "sl": -7,  "ta": 8,  "tp": 6, "m": 20, "lv": 10, "d": 0, "o": 0.0,  "sk": 0, "cap": 3000},
    {"name": "v14.2",   "tf": "30m", "ft": 3, "fp": 7,   "st": 0, "sp": 200, "ap": 20, "at": 25, "rl": 25, "rh": 65, "sl": -7,  "ta": 10, "tp": 1, "m": 30, "lv": 10, "d": 3, "o": 0.0,  "sk": 0, "cap": 3000},
    {"name": "v14.4",   "tf": "30m", "ft": 0, "fp": 3,   "st": 0, "sp": 200, "ap": 14, "at": 35, "rl": 30, "rh": 65, "sl": -7,  "ta": 6,  "tp": 3, "m": 25, "lv": 10, "d": 0, "o": 0.0,  "sk": 0, "cap": 3000},
    {"name": "v15.2",   "tf": "30m", "ft": 0, "fp": 3,   "st": 0, "sp": 200, "ap": 14, "at": 35, "rl": 30, "rh": 65, "sl": -5,  "ta": 6,  "tp": 5, "m": 30, "lv": 10, "d": 6, "o": 0.0,  "sk": 0, "cap": 3000},
    {"name": "v15.4",   "tf": "30m", "ft": 0, "fp": 3,   "st": 0, "sp": 200, "ap": 14, "at": 35, "rl": 30, "rh": 65, "sl": -7,  "ta": 6,  "tp": 3, "m": 40, "lv": 10, "d": 0, "o": 0.0,  "sk": 0, "cap": 3000},
    {"name": "v15.5",   "tf": "30m", "ft": 0, "fp": 3,   "st": 0, "sp": 200, "ap": 14, "at": 35, "rl": 35, "rh": 65, "sl": -7,  "ta": 6,  "tp": 5, "m": 35, "lv": 10, "d": 0, "o": 0.0,  "sk": 0, "cap": 3000},
    {"name": "v16.0",   "tf": "30m", "ft": 1, "fp": 3,   "st": 0, "sp": 200, "ap": 20, "at": 35, "rl": 35, "rh": 65, "sl": -8,  "ta": 4,  "tp": 3, "m": 50, "lv": 10, "d": 0, "o": 0.0,  "sk": 0, "cap": 3000},
    {"name": "v16.4",   "tf": "30m", "ft": 1, "fp": 3,   "st": 0, "sp": 200, "ap": 20, "at": 35, "rl": 35, "rh": 65, "sl": -8,  "ta": 4,  "tp": 3, "m": 30, "lv": 10, "d": 0, "o": 0.0,  "sk": 0, "cap": 3000},
    {"name": "v16.6",   "tf": "30m", "ft": 1, "fp": 3,   "st": 0, "sp": 200, "ap": 20, "at": 35, "rl": 30, "rh": 70, "sl": -8,  "ta": 3,  "tp": 2, "m": 50, "lv": 10, "d": 5, "o": 0.0,  "sk": 1, "cap": 3000},
    {"name": "v22.0F",  "tf": "30m", "ft": 1, "fp": 3,   "st": 0, "sp": 200, "ap": 20, "at": 35, "rl": 30, "rh": 70, "sl": -8,  "ta": 3,  "tp": 2, "m": 50, "lv": 10, "d": 5, "o": 0.0,  "sk": 1, "cap": 5000},
    {"name": "v22.2",   "tf": "30m", "ft": 0, "fp": 3,   "st": 0, "sp": 200, "ap": 14, "at": 35, "rl": 30, "rh": 65, "sl": -7,  "ta": 6,  "tp": 3, "m": 60, "lv": 10, "d": 0, "o": 0.0,  "sk": 1, "cap": 3000},
    {"name": "v22.3",   "tf": "30m", "ft": 0, "fp": 3,   "st": 1, "sp": 250, "ap": 20, "at": 25, "rl": 35, "rh": 65, "sl": -8,  "ta": 5,  "tp": 4, "m": 60, "lv": 10, "d": 0, "o": 0.0,  "sk": 0, "cap": 3000},
    {"name": "v22.8",   "tf": "30m", "ft": 0, "fp": 100, "st": 0, "sp": 600, "ap": 20, "at": 30, "rl": 35, "rh": 75, "sl": -8,  "ta": 6,  "tp": 5, "m": 35, "lv": 10, "d": 0, "o": 0.0,  "sk": 1, "cap": 5000},
    {"name": "v23.4",   "tf": "30m", "ft": 0, "fp": 3,   "st": 0, "sp": 200, "ap": 14, "at": 35, "rl": 30, "rh": 65, "sl": -7,  "ta": 6,  "tp": 3, "m": 30, "lv": 10, "d": 0, "o": 0.0,  "sk": 1, "cap": 5000},
    {"name": "v23.5",   "tf": "10m", "ft": 0, "fp": 3,   "st": 2, "sp": 200, "ap": 14, "at": 35, "rl": 40, "rh": 75, "sl": -10, "ta": 8,  "tp": 4, "m": 25, "lv": 3,  "d": 5, "o": 0.0,  "sk": 0, "cap": 5000},
    {"name": "v23.5b",  "tf": "30m", "ft": 3, "fp": 5,   "st": 0, "sp": 150, "ap": 20, "at": 25, "rl": 30, "rh": 65, "sl": -10, "ta": 10, "tp": 1, "m": 25, "lv": 10, "d": 3, "o": 0.0,  "sk": 0, "cap": 5000},
    {"name": "v24.2",   "tf": "1h",  "ft": 0, "fp": 3,   "st": 0, "sp": 100, "ap": 20, "at": 30, "rl": 30, "rh": 70, "sl": -8,  "ta": 6,  "tp": 5, "m": 70, "lv": 10, "d": 0, "o": 0.0,  "sk": 0, "cap": 3000},
    {"name": "v25.0",   "tf": "5m",  "ft": 0, "fp": 5,   "st": 0, "sp": 100, "ap": 14, "at": 30, "rl": 40, "rh": 60, "sl": -4,  "ta": 5,  "tp": 3, "m": 30, "lv": 10, "d": 0, "o": 0.0,  "sk": 0, "cap": 3000},
    {"name": "v25.1A",  "tf": "10m", "ft": 3, "fp": 21,  "st": 0, "sp": 250, "ap": 20, "at": 35, "rl": 40, "rh": 75, "sl": -6,  "ta": 7,  "tp": 3, "m": 50, "lv": 10, "d": 0, "o": 0.0,  "sk": 0, "cap": 3000},
    {"name": "v28_T1",  "tf": "15m", "ft": 1, "fp": 5,   "st": 4, "sp": 300, "ap": 20, "at": 35, "rl": 35, "rh": 75, "sl": -8,  "ta": 10, "tp": 5, "m": 15, "lv": 15, "d": 2, "o": 0.0,  "sk": 0, "cap": 3000},
    {"name": "v28_T2",  "tf": "15m", "ft": 3, "fp": 14,  "st": 4, "sp": 300, "ap": 20, "at": 35, "rl": 35, "rh": 70, "sl": -7,  "ta": 10, "tp": 3, "m": 50, "lv": 10, "d": 3, "o": -1.5, "sk": 0, "cap": 3000},
    {"name": "v28_T3",  "tf": "15m", "ft": 4, "fp": 2,   "st": 4, "sp": 300, "ap": 14, "at": 45, "rl": 30, "rh": 75, "sl": -9,  "ta": 5,  "tp": 5, "m": 40, "lv": 10, "d": 0, "o": 0.0,  "sk": 0, "cap": 3000},
    {"name": "v32.2",   "tf": "30m", "ft": 0, "fp": 100, "st": 0, "sp": 600, "ap": 20, "at": 30, "rl": 40, "rh": 80, "sl": -3,  "ta": 12, "tp": 9, "m": 35, "lv": 10, "d": 0, "o": 0.0,  "sk": 1, "cap": 5000},
    {"name": "v32.3",   "tf": "30m", "ft": 0, "fp": 75,  "st": 2, "sp": 750, "ap": 20, "at": 30, "rl": 40, "rh": 80, "sl": -3,  "ta": 12, "tp": 9, "m": 35, "lv": 10, "d": 0, "o": 0.0,  "sk": 1, "cap": 5000},
]


# =============================================================================
# ADX with Standard EMA (Engine 6)
# =============================================================================
@njit
def calc_adx_standard_ema(high, low, close, period):
    """ADX using standard EMA (k=2/(p+1)) instead of Wilder's smoothing"""
    n = len(close)
    adx = np.empty(n)
    adx[:] = np.nan
    plus_dm = np.zeros(n)
    minus_dm = np.zeros(n)
    tr = np.zeros(n)

    for i in range(1, n):
        h_diff = high[i] - high[i - 1]
        l_diff = low[i - 1] - low[i]
        if h_diff > l_diff and h_diff > 0:
            plus_dm[i] = h_diff
        if l_diff > h_diff and l_diff > 0:
            minus_dm[i] = l_diff
        tr1 = high[i] - low[i]
        tr2 = abs(high[i] - close[i - 1])
        tr3 = abs(low[i] - close[i - 1])
        tr[i] = max(tr1, max(tr2, tr3))

    if n < 2 * period + 1:
        return adx

    k = 2.0 / (period + 1.0)

    atr = np.mean(tr[1:period + 1])
    a_plus = np.mean(plus_dm[1:period + 1])
    a_minus = np.mean(minus_dm[1:period + 1])

    plus_di_arr = np.zeros(n)
    minus_di_arr = np.zeros(n)
    dx_arr = np.zeros(n)

    if atr > 0:
        plus_di_arr[period] = 100.0 * a_plus / atr
        minus_di_arr[period] = 100.0 * a_minus / atr
    di_sum = plus_di_arr[period] + minus_di_arr[period]
    if di_sum > 0:
        dx_arr[period] = 100.0 * abs(plus_di_arr[period] - minus_di_arr[period]) / di_sum

    for i in range(period + 1, n):
        atr = tr[i] * k + atr * (1.0 - k)
        a_plus = plus_dm[i] * k + a_plus * (1.0 - k)
        a_minus = minus_dm[i] * k + a_minus * (1.0 - k)
        if atr > 0:
            plus_di_arr[i] = 100.0 * a_plus / atr
            minus_di_arr[i] = 100.0 * a_minus / atr
        di_sum = plus_di_arr[i] + minus_di_arr[i]
        if di_sum > 0:
            dx_arr[i] = 100.0 * abs(plus_di_arr[i] - minus_di_arr[i]) / di_sum

    adx_start = 2 * period
    if adx_start >= n:
        return adx
    adx[adx_start] = np.mean(dx_arr[period:adx_start + 1])
    for i in range(adx_start + 1, n):
        adx[i] = dx_arr[i] * k + adx[i - 1] * (1.0 - k)

    return adx


# =============================================================================
# Engine 2: Numba-HL (SL on high/low intrabar)
# =============================================================================
@njit
def backtest_engine2_hl(
    close, high, low, volume, timestamps_epoch,
    fast_ma, slow_ma, adx, rsi,
    adx_thresh, rsi_lo, rsi_hi,
    sl_pct, trail_act_pct, trail_pct,
    margin_pct, leverage,
    entry_delay_bars, entry_offset_pct,
    fee_rate, skip_same_dir, weight_start_idx
):
    """Engine 2: SL checked on high/low (intrabar), TSL on close"""
    n = len(close)
    balance = 3000.0
    peak_balance = 3000.0
    max_dd = 0.0

    position = 0
    entry_price = 0.0
    position_size = 0.0
    position_margin = 0.0
    highest_roi = 0.0
    trail_active = False

    pending_signal = 0
    pending_bar = 0
    pending_price = 0.0
    last_closed_dir = 0

    total_trades = 0
    wins = 0
    losses = 0
    total_profit = 0.0
    total_loss = 0.0
    sl_count = 0
    tsl_count = 0
    rev_count = 0
    consec_loss = 0
    max_consec_loss = 0

    month_start_balance = 3000.0
    current_month = -1

    for i in range(1, n):
        if np.isnan(fast_ma[i]) or np.isnan(slow_ma[i]) or np.isnan(adx[i]) or np.isnan(rsi[i]):
            continue

        cur_price = close[i]

        approx_month = i // 8640
        if approx_month != current_month:
            current_month = approx_month
            month_start_balance = balance

        if balance < month_start_balance * 0.80:
            if position != 0:
                if position == 1:
                    roi = (cur_price - entry_price) / entry_price
                else:
                    roi = (entry_price - cur_price) / entry_price
                pnl = position_margin * leverage * roi - position_margin * leverage * fee_rate * 2
                balance += pnl
                if pnl > 0:
                    total_profit += pnl
                    wins += 1
                else:
                    total_loss += abs(pnl)
                    losses += 1
                total_trades += 1
                position = 0
            continue

        # Position management - SL on high/low
        if position != 0:
            if position == 1:
                worst_roi = (low[i] - entry_price) / entry_price
                roi = (cur_price - entry_price) / entry_price
            else:
                worst_roi = (entry_price - high[i]) / entry_price
                roi = (entry_price - cur_price) / entry_price

            # SL check on intrabar extreme
            if worst_roi <= sl_pct / 100.0:
                sl_roi = sl_pct / 100.0
                pnl = position_margin * leverage * sl_roi - position_margin * leverage * fee_rate * 2
                balance += pnl
                total_loss += abs(pnl)
                losses += 1
                sl_count += 1
                total_trades += 1
                consec_loss += 1
                if consec_loss > max_consec_loss:
                    max_consec_loss = consec_loss
                last_closed_dir = position
                position = 0
                trail_active = False
                continue

            # TSL on close
            if roi >= trail_act_pct / 100.0:
                trail_active = True
                if roi > highest_roi:
                    highest_roi = roi

            if trail_active and highest_roi > 0:
                drop = highest_roi - roi
                if drop >= trail_pct / 100.0:
                    pnl = position_margin * leverage * roi - position_margin * leverage * fee_rate * 2
                    balance += pnl
                    if pnl > 0:
                        total_profit += pnl
                        wins += 1
                        consec_loss = 0
                    else:
                        total_loss += abs(pnl)
                        losses += 1
                        consec_loss += 1
                        if consec_loss > max_consec_loss:
                            max_consec_loss = consec_loss
                    tsl_count += 1
                    total_trades += 1
                    last_closed_dir = position
                    position = 0
                    trail_active = False
                    continue

            if roi > highest_roi:
                highest_roi = roi

        # Cross detection
        cross_signal = 0
        if fast_ma[i - 1] <= slow_ma[i - 1] and fast_ma[i] > slow_ma[i]:
            cross_signal = 1
        elif fast_ma[i - 1] >= slow_ma[i - 1] and fast_ma[i] < slow_ma[i]:
            cross_signal = -1

        if cross_signal != 0:
            if position != 0 and position != cross_signal:
                if position == 1:
                    roi = (cur_price - entry_price) / entry_price
                else:
                    roi = (entry_price - cur_price) / entry_price
                pnl = position_margin * leverage * roi - position_margin * leverage * fee_rate * 2
                balance += pnl
                if pnl > 0:
                    total_profit += pnl
                    wins += 1
                    consec_loss = 0
                else:
                    total_loss += abs(pnl)
                    losses += 1
                    consec_loss += 1
                    if consec_loss > max_consec_loss:
                        max_consec_loss = consec_loss
                rev_count += 1
                total_trades += 1
                last_closed_dir = position
                position = 0
                trail_active = False

            if adx[i] >= adx_thresh and rsi_lo <= rsi[i] <= rsi_hi:
                if skip_same_dir and cross_signal == last_closed_dir:
                    pass
                else:
                    pending_signal = cross_signal
                    pending_bar = i
                    pending_price = cur_price

        # Delayed entry
        if pending_signal != 0 and position == 0:
            bars_elapsed = i - pending_bar
            if bars_elapsed >= entry_delay_bars:
                if pending_signal == 1:
                    price_change = (cur_price - pending_price) / pending_price * 100.0
                    ok = price_change >= entry_offset_pct
                else:
                    price_change = (pending_price - cur_price) / pending_price * 100.0
                    ok = price_change >= entry_offset_pct

                if ok and adx[i] >= adx_thresh and rsi_lo <= rsi[i] <= rsi_hi:
                    entry_price = cur_price
                    position_margin = balance * margin_pct / 100.0
                    position_size = position_margin * leverage / cur_price
                    position = pending_signal
                    highest_roi = 0.0
                    trail_active = False
                    pending_signal = 0
                elif bars_elapsed > entry_delay_bars + 12:
                    pending_signal = 0

        if balance > peak_balance:
            peak_balance = balance
        dd = (peak_balance - balance) / peak_balance * 100.0
        if dd > max_dd:
            max_dd = dd

    # Close remaining
    if position != 0:
        if position == 1:
            roi = (close[n - 1] - entry_price) / entry_price
        else:
            roi = (entry_price - close[n - 1]) / entry_price
        pnl = position_margin * leverage * roi - position_margin * leverage * fee_rate * 2
        balance += pnl
        if pnl > 0:
            total_profit += pnl
            wins += 1
        else:
            total_loss += abs(pnl)
            losses += 1
        total_trades += 1

    pf = total_profit / total_loss if total_loss > 0 else 999.0
    return (balance, total_trades, wins, losses, pf, max_dd,
            sl_count, tsl_count, rev_count)


# =============================================================================
# Engine 3: Pandas-Close (pure numpy, no Numba)
# =============================================================================
def calc_ma_numpy(close, volume, ma_type, period):
    """MA calculation without Numba"""
    n = len(close)
    out = np.full(n, np.nan)
    if ma_type == 0:  # EMA
        if n < period:
            return out
        k = 2.0 / (period + 1.0)
        out[period - 1] = np.mean(close[:period])
        for i in range(period, n):
            out[i] = close[i] * k + out[i - 1] * (1.0 - k)
    elif ma_type == 1:  # WMA
        w_sum = period * (period + 1) / 2.0
        weights = np.arange(1, period + 1, dtype=np.float64)
        for i in range(period - 1, n):
            out[i] = np.sum(close[i - period + 1:i + 1] * weights) / w_sum
    elif ma_type == 2:  # SMA
        for i in range(period - 1, n):
            out[i] = np.mean(close[i - period + 1:i + 1])
    elif ma_type == 3:  # HMA
        half_p = max(int(period / 2), 1)
        sqrt_p = max(int(np.sqrt(period)), 1)
        wma_half = calc_ma_numpy(close, volume, 1, half_p)
        wma_full = calc_ma_numpy(close, volume, 1, period)
        diff = np.full(n, np.nan)
        valid = ~np.isnan(wma_half) & ~np.isnan(wma_full)
        diff[valid] = 2.0 * wma_half[valid] - wma_full[valid]
        out = calc_ma_numpy(diff, volume, 1, sqrt_p)
    elif ma_type == 4:  # VWMA
        for i in range(period - 1, n):
            sl = slice(i - period + 1, i + 1)
            v_sum = np.sum(volume[sl])
            if v_sum > 0:
                out[i] = np.sum(close[sl] * volume[sl]) / v_sum
    return out


def calc_rsi_numpy(data, period):
    """RSI with Wilder's smoothing (numpy)"""
    n = len(data)
    out = np.full(n, np.nan)
    if n < period + 1:
        return out
    deltas = np.diff(data)
    gains = np.where(deltas > 0, deltas, 0.0)
    losses_arr = np.where(deltas < 0, -deltas, 0.0)
    avg_gain = np.mean(gains[:period])
    avg_loss = np.mean(losses_arr[:period])
    if avg_loss == 0:
        out[period] = 100.0
    else:
        out[period] = 100.0 - 100.0 / (1.0 + avg_gain / avg_loss)
    for i in range(period, len(gains)):
        avg_gain = (avg_gain * (period - 1) + gains[i]) / period
        avg_loss = (avg_loss * (period - 1) + losses_arr[i]) / period
        if avg_loss == 0:
            out[i + 1] = 100.0
        else:
            out[i + 1] = 100.0 - 100.0 / (1.0 + avg_gain / avg_loss)
    return out


def calc_adx_numpy(high, low, close, period):
    """ADX with Wilder's smoothing (numpy, no Numba)"""
    n = len(close)
    adx = np.full(n, np.nan)
    plus_dm = np.zeros(n)
    minus_dm = np.zeros(n)
    tr = np.zeros(n)
    for i in range(1, n):
        h_diff = high[i] - high[i - 1]
        l_diff = low[i - 1] - low[i]
        if h_diff > l_diff and h_diff > 0:
            plus_dm[i] = h_diff
        if l_diff > h_diff and l_diff > 0:
            minus_dm[i] = l_diff
        tr1 = high[i] - low[i]
        tr2 = abs(high[i] - close[i - 1])
        tr3 = abs(low[i] - close[i - 1])
        tr[i] = max(tr1, max(tr2, tr3))
    if n < 2 * period + 1:
        return adx
    atr = np.mean(tr[1:period + 1])
    a_plus = np.mean(plus_dm[1:period + 1])
    a_minus = np.mean(minus_dm[1:period + 1])
    plus_di_arr = np.zeros(n)
    minus_di_arr = np.zeros(n)
    dx_arr = np.zeros(n)
    if atr > 0:
        plus_di_arr[period] = 100.0 * a_plus / atr
        minus_di_arr[period] = 100.0 * a_minus / atr
    di_sum = plus_di_arr[period] + minus_di_arr[period]
    if di_sum > 0:
        dx_arr[period] = 100.0 * abs(plus_di_arr[period] - minus_di_arr[period]) / di_sum
    for i in range(period + 1, n):
        atr = (atr * (period - 1) + tr[i]) / period
        a_plus = (a_plus * (period - 1) + plus_dm[i]) / period
        a_minus = (a_minus * (period - 1) + minus_dm[i]) / period
        if atr > 0:
            plus_di_arr[i] = 100.0 * a_plus / atr
            minus_di_arr[i] = 100.0 * a_minus / atr
        di_sum = plus_di_arr[i] + minus_di_arr[i]
        if di_sum > 0:
            dx_arr[i] = 100.0 * abs(plus_di_arr[i] - minus_di_arr[i]) / di_sum
    adx_start = 2 * period
    if adx_start >= n:
        return adx
    adx[adx_start] = np.mean(dx_arr[period:adx_start + 1])
    for i in range(adx_start + 1, n):
        adx[i] = (adx[i - 1] * (period - 1) + dx_arr[i]) / period
    return adx


def backtest_engine3_pandas(
    close, high, low, volume,
    fast_ma, slow_ma, adx, rsi,
    adx_thresh, rsi_lo, rsi_hi,
    sl_pct, trail_act_pct, trail_pct,
    margin_pct, leverage,
    entry_delay_bars, entry_offset_pct,
    fee_rate, skip_same_dir
):
    """Engine 3: Pure pandas/numpy backtest, SL/TSL on close"""
    n = len(close)
    balance = 3000.0
    peak_balance = 3000.0
    max_dd = 0.0

    position = 0
    entry_price = 0.0
    position_margin = 0.0
    highest_roi = 0.0
    trail_active = False

    pending_signal = 0
    pending_bar = 0
    pending_price = 0.0
    last_closed_dir = 0

    total_trades = 0
    wins = 0
    losses = 0
    total_profit = 0.0
    total_loss = 0.0
    sl_count = 0
    tsl_count = 0
    rev_count = 0
    consec_loss = 0
    max_consec_loss = 0

    month_start_balance = 3000.0
    current_month = -1

    for i in range(1, n):
        if np.isnan(fast_ma[i]) or np.isnan(slow_ma[i]) or np.isnan(adx[i]) or np.isnan(rsi[i]):
            continue
        cur_price = close[i]

        approx_month = i // 8640
        if approx_month != current_month:
            current_month = approx_month
            month_start_balance = balance

        if balance < month_start_balance * 0.80:
            if position != 0:
                roi = (cur_price - entry_price) / entry_price if position == 1 else (entry_price - cur_price) / entry_price
                pnl = position_margin * leverage * roi - position_margin * leverage * fee_rate * 2
                balance += pnl
                if pnl > 0:
                    total_profit += pnl
                    wins += 1
                else:
                    total_loss += abs(pnl)
                    losses += 1
                total_trades += 1
                position = 0
            continue

        if position != 0:
            roi = (cur_price - entry_price) / entry_price if position == 1 else (entry_price - cur_price) / entry_price

            if roi <= sl_pct / 100.0:
                pnl = position_margin * leverage * roi - position_margin * leverage * fee_rate * 2
                balance += pnl
                total_loss += abs(pnl)
                losses += 1
                sl_count += 1
                total_trades += 1
                consec_loss += 1
                if consec_loss > max_consec_loss:
                    max_consec_loss = consec_loss
                last_closed_dir = position
                position = 0
                trail_active = False
                continue

            if roi >= trail_act_pct / 100.0:
                trail_active = True
                if roi > highest_roi:
                    highest_roi = roi

            if trail_active and highest_roi > 0:
                drop = highest_roi - roi
                if drop >= trail_pct / 100.0:
                    pnl = position_margin * leverage * roi - position_margin * leverage * fee_rate * 2
                    balance += pnl
                    if pnl > 0:
                        total_profit += pnl
                        wins += 1
                        consec_loss = 0
                    else:
                        total_loss += abs(pnl)
                        losses += 1
                        consec_loss += 1
                        if consec_loss > max_consec_loss:
                            max_consec_loss = consec_loss
                    tsl_count += 1
                    total_trades += 1
                    last_closed_dir = position
                    position = 0
                    trail_active = False
                    continue

            if roi > highest_roi:
                highest_roi = roi

        cross_signal = 0
        if fast_ma[i - 1] <= slow_ma[i - 1] and fast_ma[i] > slow_ma[i]:
            cross_signal = 1
        elif fast_ma[i - 1] >= slow_ma[i - 1] and fast_ma[i] < slow_ma[i]:
            cross_signal = -1

        if cross_signal != 0:
            if position != 0 and position != cross_signal:
                roi = (cur_price - entry_price) / entry_price if position == 1 else (entry_price - cur_price) / entry_price
                pnl = position_margin * leverage * roi - position_margin * leverage * fee_rate * 2
                balance += pnl
                if pnl > 0:
                    total_profit += pnl
                    wins += 1
                    consec_loss = 0
                else:
                    total_loss += abs(pnl)
                    losses += 1
                    consec_loss += 1
                    if consec_loss > max_consec_loss:
                        max_consec_loss = consec_loss
                rev_count += 1
                total_trades += 1
                last_closed_dir = position
                position = 0
                trail_active = False

            if adx[i] >= adx_thresh and rsi_lo <= rsi[i] <= rsi_hi:
                if skip_same_dir and cross_signal == last_closed_dir:
                    pass
                else:
                    pending_signal = cross_signal
                    pending_bar = i
                    pending_price = cur_price

        if pending_signal != 0 and position == 0:
            bars_elapsed = i - pending_bar
            if bars_elapsed >= entry_delay_bars:
                if pending_signal == 1:
                    price_change = (cur_price - pending_price) / pending_price * 100.0
                    ok = price_change >= entry_offset_pct
                else:
                    price_change = (pending_price - cur_price) / pending_price * 100.0
                    ok = price_change >= entry_offset_pct

                if ok and adx[i] >= adx_thresh and rsi_lo <= rsi[i] <= rsi_hi:
                    entry_price = cur_price
                    position_margin = balance * margin_pct / 100.0
                    position = pending_signal
                    highest_roi = 0.0
                    trail_active = False
                    pending_signal = 0
                elif bars_elapsed > entry_delay_bars + 12:
                    pending_signal = 0

        if balance > peak_balance:
            peak_balance = balance
        dd = (peak_balance - balance) / peak_balance * 100.0
        if dd > max_dd:
            max_dd = dd

    if position != 0:
        roi = (close[n - 1] - entry_price) / entry_price if position == 1 else (entry_price - close[n - 1]) / entry_price
        pnl = position_margin * leverage * roi - position_margin * leverage * fee_rate * 2
        balance += pnl
        if pnl > 0:
            total_profit += pnl
            wins += 1
        else:
            total_loss += abs(pnl)
            losses += 1
        total_trades += 1

    pf = total_profit / total_loss if total_loss > 0 else 999.0
    return (balance, total_trades, wins, losses, pf, max_dd, sl_count, tsl_count, rev_count)


# =============================================================================
# Engine 4: NextBar (entry on next bar open)
# =============================================================================
@njit
def backtest_engine4_nextbar(
    close, high, low, open_price, volume, timestamps_epoch,
    fast_ma, slow_ma, adx, rsi,
    adx_thresh, rsi_lo, rsi_hi,
    sl_pct, trail_act_pct, trail_pct,
    margin_pct, leverage,
    entry_delay_bars, entry_offset_pct,
    fee_rate, skip_same_dir, weight_start_idx
):
    """Engine 4: Entry at next bar's open instead of current close"""
    n = len(close)
    balance = 3000.0
    peak_balance = 3000.0
    max_dd = 0.0

    position = 0
    entry_price = 0.0
    position_margin = 0.0
    highest_roi = 0.0
    trail_active = False

    pending_signal = 0
    pending_bar = 0
    pending_price = 0.0
    last_closed_dir = 0

    deferred_entry = 0

    total_trades = 0
    wins = 0
    losses = 0
    total_profit = 0.0
    total_loss = 0.0
    sl_count = 0
    tsl_count = 0
    rev_count = 0
    consec_loss = 0
    max_consec_loss = 0

    month_start_balance = 3000.0
    current_month = -1

    for i in range(1, n):
        if np.isnan(fast_ma[i]) or np.isnan(slow_ma[i]) or np.isnan(adx[i]) or np.isnan(rsi[i]):
            continue

        cur_price = close[i]

        approx_month = i // 8640
        if approx_month != current_month:
            current_month = approx_month
            month_start_balance = balance

        if balance < month_start_balance * 0.80:
            if position != 0:
                roi = (cur_price - entry_price) / entry_price if position == 1 else (entry_price - cur_price) / entry_price
                pnl = position_margin * leverage * roi - position_margin * leverage * fee_rate * 2
                balance += pnl
                if pnl > 0:
                    total_profit += pnl
                    wins += 1
                else:
                    total_loss += abs(pnl)
                    losses += 1
                total_trades += 1
                position = 0
            deferred_entry = 0
            continue

        # Execute deferred entry at this bar's OPEN
        if deferred_entry != 0 and position == 0:
            entry_price = open_price[i]
            position_margin = balance * margin_pct / 100.0
            position = deferred_entry
            highest_roi = 0.0
            trail_active = False
            deferred_entry = 0

        # Position management (on close)
        if position != 0:
            roi = (cur_price - entry_price) / entry_price if position == 1 else (entry_price - cur_price) / entry_price

            if roi <= sl_pct / 100.0:
                pnl = position_margin * leverage * roi - position_margin * leverage * fee_rate * 2
                balance += pnl
                total_loss += abs(pnl)
                losses += 1
                sl_count += 1
                total_trades += 1
                consec_loss += 1
                if consec_loss > max_consec_loss:
                    max_consec_loss = consec_loss
                last_closed_dir = position
                position = 0
                trail_active = False
                continue

            if roi >= trail_act_pct / 100.0:
                trail_active = True
                if roi > highest_roi:
                    highest_roi = roi

            if trail_active and highest_roi > 0:
                drop = highest_roi - roi
                if drop >= trail_pct / 100.0:
                    pnl = position_margin * leverage * roi - position_margin * leverage * fee_rate * 2
                    balance += pnl
                    if pnl > 0:
                        total_profit += pnl
                        wins += 1
                        consec_loss = 0
                    else:
                        total_loss += abs(pnl)
                        losses += 1
                        consec_loss += 1
                        if consec_loss > max_consec_loss:
                            max_consec_loss = consec_loss
                    tsl_count += 1
                    total_trades += 1
                    last_closed_dir = position
                    position = 0
                    trail_active = False
                    continue

            if roi > highest_roi:
                highest_roi = roi

        # Cross detection
        cross_signal = 0
        if fast_ma[i - 1] <= slow_ma[i - 1] and fast_ma[i] > slow_ma[i]:
            cross_signal = 1
        elif fast_ma[i - 1] >= slow_ma[i - 1] and fast_ma[i] < slow_ma[i]:
            cross_signal = -1

        if cross_signal != 0:
            if position != 0 and position != cross_signal:
                roi = (cur_price - entry_price) / entry_price if position == 1 else (entry_price - cur_price) / entry_price
                pnl = position_margin * leverage * roi - position_margin * leverage * fee_rate * 2
                balance += pnl
                if pnl > 0:
                    total_profit += pnl
                    wins += 1
                    consec_loss = 0
                else:
                    total_loss += abs(pnl)
                    losses += 1
                    consec_loss += 1
                    if consec_loss > max_consec_loss:
                        max_consec_loss = consec_loss
                rev_count += 1
                total_trades += 1
                last_closed_dir = position
                position = 0
                trail_active = False

            if adx[i] >= adx_thresh and rsi_lo <= rsi[i] <= rsi_hi:
                if skip_same_dir and cross_signal == last_closed_dir:
                    pass
                else:
                    pending_signal = cross_signal
                    pending_bar = i
                    pending_price = cur_price

        # Delayed entry check - defer actual entry to next bar open
        if pending_signal != 0 and position == 0 and deferred_entry == 0:
            bars_elapsed = i - pending_bar
            if bars_elapsed >= entry_delay_bars:
                if pending_signal == 1:
                    price_change = (cur_price - pending_price) / pending_price * 100.0
                    ok = price_change >= entry_offset_pct
                else:
                    price_change = (pending_price - cur_price) / pending_price * 100.0
                    ok = price_change >= entry_offset_pct

                if ok and adx[i] >= adx_thresh and rsi_lo <= rsi[i] <= rsi_hi:
                    deferred_entry = pending_signal
                    pending_signal = 0
                elif bars_elapsed > entry_delay_bars + 12:
                    pending_signal = 0

        if balance > peak_balance:
            peak_balance = balance
        dd = (peak_balance - balance) / peak_balance * 100.0
        if dd > max_dd:
            max_dd = dd

    if position != 0:
        roi = (close[n - 1] - entry_price) / entry_price if position == 1 else (entry_price - close[n - 1]) / entry_price
        pnl = position_margin * leverage * roi - position_margin * leverage * fee_rate * 2
        balance += pnl
        if pnl > 0:
            total_profit += pnl
            wins += 1
        else:
            total_loss += abs(pnl)
            losses += 1
        total_trades += 1

    pf = total_profit / total_loss if total_loss > 0 else 999.0
    return (balance, total_trades, wins, losses, pf, max_dd,
            sl_count, tsl_count, rev_count, max_consec_loss,
            total_profit, total_loss)


# =============================================================================
# Runner: run a single strategy on a single engine
# =============================================================================
ENGINE_NAMES = {
    1: "Numba-Close",
    2: "Numba-HL",
    3: "Pandas-Close",
    4: "NextBar",
    5: "Wilder-ADX",
    6: "StdEMA-ADX",
}


def run_single_engine(engine_id, strat, tf_data):
    """
    Run one strategy on one engine.
    Returns: (balance, trades, wins, losses, pf, mdd, sl_count, tsl_count, rev_count)
    """
    tf = strat["tf"]
    data = tf_data[tf]

    close = data['close'].values.astype(np.float64)
    high = data['high'].values.astype(np.float64)
    low = data['low'].values.astype(np.float64)
    volume = data['volume'].values.astype(np.float64)
    open_price = data['open'].values.astype(np.float64) if 'open' in data.columns else close.copy()
    n = len(close)
    timestamps_epoch = np.zeros(n, dtype=np.float64)

    # Compute indicators based on engine variant
    if engine_id in (1, 2, 4, 5):
        fast_ma = calc_ma(close, volume, strat["ft"], strat["fp"])
        slow_ma = calc_ma(close, volume, strat["st"], strat["sp"])
        rsi = calc_rsi_wilder(close, 14)
        adx = calc_adx_wilder(high, low, close, strat["ap"])
    elif engine_id == 3:
        fast_ma = calc_ma_numpy(close, volume, strat["ft"], strat["fp"])
        slow_ma = calc_ma_numpy(close, volume, strat["st"], strat["sp"])
        rsi = calc_rsi_numpy(close, 14)
        adx = calc_adx_numpy(high, low, close, strat["ap"])
    elif engine_id == 6:
        fast_ma = calc_ma(close, volume, strat["ft"], strat["fp"])
        slow_ma = calc_ma(close, volume, strat["st"], strat["sp"])
        rsi = calc_rsi_wilder(close, 14)
        adx = calc_adx_standard_ema(high, low, close, strat["ap"])

    # Common parameters
    adx_t = float(strat["at"])
    rsi_lo = float(strat["rl"])
    rsi_hi = float(strat["rh"])
    sl = float(strat["sl"])
    t_act = float(strat["ta"])
    t_pct = float(strat["tp"])
    m_pct = float(strat["m"])
    lev = float(strat["lv"])
    delay = int(strat["d"])
    offset = float(strat["o"])
    skip = int(strat["sk"])
    w_start = n // 3

    if engine_id == 1 or engine_id == 5:
        result = backtest_core(
            close, high, low, volume, timestamps_epoch,
            fast_ma, slow_ma, adx, rsi,
            adx_t, rsi_lo, rsi_hi,
            sl, t_act, t_pct,
            m_pct, lev,
            delay, offset,
            FEE_RATE, skip, w_start
        )
        return (result[0], result[1], result[2], result[3], result[4], result[5],
                result[6], result[7], result[8])

    elif engine_id == 2:
        result = backtest_engine2_hl(
            close, high, low, volume, timestamps_epoch,
            fast_ma, slow_ma, adx, rsi,
            adx_t, rsi_lo, rsi_hi,
            sl, t_act, t_pct,
            m_pct, lev,
            delay, offset,
            FEE_RATE, skip, w_start
        )
        return (result[0], result[1], result[2], result[3], result[4], result[5],
                result[6], result[7], result[8])

    elif engine_id == 3:
        result = backtest_engine3_pandas(
            close, high, low, volume,
            fast_ma, slow_ma, adx, rsi,
            adx_t, rsi_lo, rsi_hi,
            sl, t_act, t_pct,
            m_pct, lev,
            delay, offset,
            FEE_RATE, skip
        )
        return result  # already (bal, trades, wins, losses, pf, mdd, sl, tsl, rev)

    elif engine_id == 4:
        result = backtest_engine4_nextbar(
            close, high, low, open_price, volume, timestamps_epoch,
            fast_ma, slow_ma, adx, rsi,
            adx_t, rsi_lo, rsi_hi,
            sl, t_act, t_pct,
            m_pct, lev,
            delay, offset,
            FEE_RATE, skip, w_start
        )
        return (result[0], result[1], result[2], result[3], result[4], result[5],
                result[6], result[7], result[8])

    elif engine_id == 6:
        result = backtest_core(
            close, high, low, volume, timestamps_epoch,
            fast_ma, slow_ma, adx, rsi,
            adx_t, rsi_lo, rsi_hi,
            sl, t_act, t_pct,
            m_pct, lev,
            delay, offset,
            FEE_RATE, skip, w_start
        )
        return (result[0], result[1], result[2], result[3], result[4], result[5],
                result[6], result[7], result[8])


# =============================================================================
# Detailed backtest runner for a strategy (Engine 1 only, with yearly tracking)
# =============================================================================
def run_detailed_for_strategy(strat, tf_data):
    """
    Run detailed_backtest (from v28_detailed_verify) for one strategy.
    Returns the full detailed result dict with trades, yearly, monthly data.
    """
    tf = strat["tf"]
    data = tf_data[tf]

    timestamps = data['timestamp'].values
    ts_pd = pd.to_datetime(data['timestamp'])
    close = data['close'].values.astype(np.float64)
    high = data['high'].values.astype(np.float64)
    low = data['low'].values.astype(np.float64)
    volume = data['volume'].values.astype(np.float64)

    fast_ma = calc_ma(close, volume, strat["ft"], strat["fp"])
    slow_ma = calc_ma(close, volume, strat["st"], strat["sp"])
    adx_arr = calc_adx_wilder(high, low, close, strat["ap"])
    rsi_arr = calc_rsi_wilder(close, 14)

    result = detailed_backtest(
        ts_pd, close, high, low, volume,
        fast_ma, slow_ma, adx_arr, rsi_arr,
        float(strat["at"]), float(strat["rl"]), float(strat["rh"]),
        float(strat["sl"]), float(strat["ta"]), float(strat["tp"]),
        float(strat["m"]), float(strat["lv"]),
        int(strat["d"]), float(strat["o"]),
        FEE_RATE, int(strat["sk"])
    )

    return result


# =============================================================================
# Compute yearly MDD and PF from detailed trade log
# =============================================================================
def compute_yearly_stats(detailed_result):
    """
    From detailed trades, compute per-year: trade_count, MDD, PF, return%
    Returns dict: { year_str: { trades, mdd, pf, return_pct, start_bal, end_bal } }
    """
    trades = detailed_result['trades']
    yearly_balances = detailed_result['yearly_balances']
    yearly_stats = {}

    # Group trades by year (using exit_time year)
    year_trades = {}
    for t in trades:
        year = t['exit_time'][:4]
        if year not in year_trades:
            year_trades[year] = []
        year_trades[year].append(t)

    for year in sorted(yearly_balances.keys()):
        yb = yearly_balances[year]
        start_bal = yb['start']
        end_bal = yb['end']
        ret_pct = (end_bal - start_bal) / start_bal * 100.0 if start_bal > 0 else 0.0

        yr_trades = year_trades.get(year, [])
        n_trades = len(yr_trades)

        # Compute PF within year
        yr_profit = sum(t['pnl'] for t in yr_trades if t['pnl'] > 0)
        yr_loss = sum(abs(t['pnl']) for t in yr_trades if t['pnl'] <= 0)
        yr_pf = yr_profit / yr_loss if yr_loss > 0 else 999.0

        # Compute MDD within year from trade-by-trade balance progression
        if len(yr_trades) > 0:
            # Reconstruct balance curve within year
            peak_bal = start_bal
            max_dd = 0.0
            bal = start_bal
            for t in yr_trades:
                bal = t['balance_after']
                if bal > peak_bal:
                    peak_bal = bal
                dd = (peak_bal - bal) / peak_bal * 100.0 if peak_bal > 0 else 0.0
                if dd > max_dd:
                    max_dd = dd
            yr_mdd = max_dd
        else:
            yr_mdd = 0.0

        yearly_stats[year] = {
            'trades': n_trades,
            'mdd': round(yr_mdd, 2),
            'pf': round(yr_pf, 2),
            'return_pct': round(ret_pct, 2),
            'start_bal': round(start_bal, 2),
            'end_bal': round(end_bal, 2),
        }

    return yearly_stats


# =============================================================================
# Consistency analysis
# =============================================================================
def check_engine_consistency(engine_results):
    """
    Check if 6 engines agree.
    engine_results: dict { engine_id: (bal, trades, ...) }
    Returns: (is_consistent, deviation_pct, reason_str)
    """
    balances = []
    for eid in range(1, 7):
        if eid in engine_results:
            balances.append(engine_results[eid][0])

    if len(balances) < 2:
        return False, 999.0, "INSUFFICIENT"

    mean_bal = np.mean(balances)
    if mean_bal == 0:
        return False, 999.0, "ZERO_MEAN"

    min_bal = np.min(balances)
    max_bal = np.max(balances)
    deviation = (max_bal - min_bal) / abs(mean_bal) * 100.0

    # E1 vs E5 should be identical (both use backtest_core with Wilder ADX)
    e1_bal = engine_results.get(1, (0,))[0]
    e5_bal = engine_results.get(5, (0,))[0]
    e1_e5_match = abs(e1_bal - e5_bal) < 0.01

    reasons = []
    if not e1_e5_match:
        reasons.append("E1!=E5(Wilder)")

    # E1 vs E3 (same logic, different impl)
    e3_bal = engine_results.get(3, (0,))[0]
    e1_e3_dev = abs(e1_bal - e3_bal) / abs(e1_bal) * 100.0 if e1_bal != 0 else 999.0
    if e1_e3_dev > 5.0:
        reasons.append(f"E1-E3:{e1_e3_dev:.1f}%")

    # E6 deviation
    e6_bal = engine_results.get(6, (0,))[0]
    e1_e6_dev = abs(e1_bal - e6_bal) / abs(e1_bal) * 100.0 if e1_bal != 0 else 999.0
    if e1_e6_dev > 15.0:
        reasons.append(f"E1-E6:{e1_e6_dev:.1f}%")

    if deviation < 10.0:
        return True, round(deviation, 2), "CONSISTENT" if not reasons else "; ".join(reasons)
    else:
        if not reasons:
            reasons.append(f"Dev:{deviation:.1f}%")
        return False, round(deviation, 2), "; ".join(reasons)


# =============================================================================
# Word document generation
# =============================================================================
def generate_word_document(all_results, detailed_results, yearly_stats_all, output_path):
    """
    Generate Word document with 3 tables:
    1. 수익률 BEST 10
    2. 안정형 BEST 10
    3. 폐기형 BEST 10
    """
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.table import WdTableAlignment
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading('v28 전체 기획서 6엔진 교차검증 결과', level=0)
    doc.add_paragraph(f'생성일: {pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'총 전략 수: {len(strategies)}개 | 엔진 수: 6개')
    doc.add_paragraph(f'초기 자본금: $3,000 (일부 $5,000) | 수수료: 0.04% x2 | 월간 손실 제한: -20%')
    doc.add_paragraph('')

    # ----- Build ranking data -----
    ranking_data = []
    for strat in strategies:
        sn = strat["name"]
        cap = strat["cap"]

        if sn not in all_results:
            continue

        eng_res = all_results[sn]
        e1_res = eng_res.get(1, None)
        if e1_res is None:
            continue

        e1_bal = e1_res[0]
        e1_trades = e1_res[1]
        e1_pf = e1_res[4]
        e1_mdd = e1_res[5]

        # Normalize return to percentage (regardless of starting capital)
        # detailed_backtest always starts from $3000
        return_pct = (e1_bal - 3000.0) / 3000.0 * 100.0
        pnl_amount = e1_bal - 3000.0

        # Yearly stats
        ys = yearly_stats_all.get(sn, {})

        # Engine consistency
        is_consistent, dev_pct, consistency_reason = check_engine_consistency(eng_res)

        # Yearly strings
        years = sorted(ys.keys())
        yearly_trade_str = " / ".join([f"{y}:{ys[y]['trades']}" for y in years]) if years else "-"
        yearly_mdd_str = " / ".join([f"{y}:{ys[y]['mdd']:.1f}%" for y in years]) if years else "-"
        yearly_pf_str = " / ".join([f"{y}:{ys[y]['pf']:.2f}" for y in years]) if years else "-"

        ranking_data.append({
            'name': sn,
            'cap': cap,
            'return_pct': return_pct,
            'pnl_amount': pnl_amount,
            'balance': e1_bal,
            'total_trades': e1_trades,
            'pf': e1_pf,
            'mdd': e1_mdd,
            'is_consistent': is_consistent,
            'dev_pct': dev_pct,
            'consistency_reason': consistency_reason,
            'yearly_trades': yearly_trade_str,
            'yearly_mdd': yearly_mdd_str,
            'yearly_pf': yearly_pf_str,
            'yearly_stats': ys,
        })

    # Column headers (shared)
    columns = [
        '순위', '파일명', '손익률', '손익금액', '년도별 거래량',
        '총 거래량', '년도별 MDD', '년도별 PF', '6개 엔진 일치 사유', '비고(사유)'
    ]

    def add_result_table(doc, heading_text, items, heading_level=1):
        """Add a formatted table to the document"""
        doc.add_heading(heading_text, level=heading_level)

        if not items:
            doc.add_paragraph('해당 데이터 없음')
            return

        table = doc.add_table(rows=1 + len(items), cols=len(columns))
        table.style = 'Light Grid Accent 1'
        table.alignment = WdTableAlignment.CENTER

        # Header row
        hdr = table.rows[0]
        for ci, col_name in enumerate(columns):
            cell = hdr.cells[ci]
            cell.text = col_name
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(8)

        # Data rows
        for ri, item in enumerate(items):
            row = table.rows[ri + 1]

            values = [
                str(ri + 1),
                item['name'],
                f"{item['return_pct']:+.1f}%",
                f"${item['pnl_amount']:+,.0f}",
                item['yearly_trades'],
                str(item['total_trades']),
                item['yearly_mdd'],
                item['yearly_pf'],
                "PASS" if item['is_consistent'] else f"FAIL({item['dev_pct']:.1f}%)",
                item.get('remark', item['consistency_reason']),
            ]

            for ci, val in enumerate(values):
                cell = row.cells[ci]
                cell.text = val
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.size = Pt(7)

        # Adjust column widths
        for row in table.rows:
            row.cells[0].width = Cm(1.0)
            row.cells[1].width = Cm(1.8)
            row.cells[2].width = Cm(1.8)
            row.cells[3].width = Cm(2.2)
            row.cells[4].width = Cm(4.0)
            row.cells[5].width = Cm(1.5)
            row.cells[6].width = Cm(4.0)
            row.cells[7].width = Cm(4.0)
            row.cells[8].width = Cm(3.0)
            row.cells[9].width = Cm(3.5)

    # ----- Table 1: 수익률 BEST 10 -----
    # Sort by return, top 10
    sorted_by_return = sorted(ranking_data, key=lambda x: x['return_pct'], reverse=True)
    best_return = sorted_by_return[:10]
    for item in best_return:
        remarks = []
        if item['return_pct'] > 100:
            remarks.append("고수익")
        if item['is_consistent']:
            remarks.append("엔진일치")
        if item['mdd'] < 30:
            remarks.append(f"MDD{item['mdd']:.0f}%")
        item['remark'] = "; ".join(remarks) if remarks else item['consistency_reason']

    add_result_table(doc, '1. 수익률 BEST 10 (Engine 1 기준, 높은 수익률 순)', best_return)

    doc.add_paragraph('')

    # ----- Table 2: 안정형 BEST 10 -----
    # Sort by lowest MDD (only profitable strategies)
    profitable = [d for d in ranking_data if d['return_pct'] > 0]
    sorted_by_mdd = sorted(profitable, key=lambda x: x['mdd'])
    best_stable = sorted_by_mdd[:10]
    for item in best_stable:
        remarks = []
        remarks.append(f"MDD {item['mdd']:.1f}%")
        if item['is_consistent']:
            remarks.append("엔진일치")
        if item['pf'] > 2.0:
            remarks.append(f"PF{item['pf']:.1f}")
        item['remark'] = "; ".join(remarks) if remarks else item['consistency_reason']

    add_result_table(doc, '2. 안정형 BEST 10 (수익 전략 중 낮은 MDD 순)', best_stable)

    doc.add_paragraph('')

    # ----- Table 3: 폐기형 BEST 10 -----
    # Bottom 10: loss, high MDD, or engine inconsistency
    # Score: negative return = bad, high MDD = bad, inconsistency = bad
    for item in ranking_data:
        penalty = 0
        if item['return_pct'] < 0:
            penalty += abs(item['return_pct'])
        if item['mdd'] > 50:
            penalty += item['mdd'] - 50
        if not item['is_consistent']:
            penalty += 20
        item['discard_score'] = penalty

    sorted_by_discard = sorted(ranking_data, key=lambda x: x['discard_score'], reverse=True)
    worst = sorted_by_discard[:10]
    for item in worst:
        remarks = []
        if item['return_pct'] < 0:
            remarks.append(f"손실{item['return_pct']:.1f}%")
        if item['mdd'] > 50:
            remarks.append(f"고MDD{item['mdd']:.1f}%")
        if not item['is_consistent']:
            remarks.append(f"불일치{item['dev_pct']:.1f}%")
        if item['pf'] < 1.0:
            remarks.append(f"PF{item['pf']:.2f}")
        item['remark'] = "; ".join(remarks) if remarks else "폐기 대상"

    add_result_table(doc, '3. 폐기형 BEST 10 (손실/고MDD/엔진불일치 순)', worst)

    doc.add_paragraph('')

    # ----- Summary section -----
    doc.add_heading('4. 전체 요약', level=1)

    total_strategies = len(ranking_data)
    profitable_count = len([d for d in ranking_data if d['return_pct'] > 0])
    consistent_count = len([d for d in ranking_data if d['is_consistent']])
    avg_return = np.mean([d['return_pct'] for d in ranking_data]) if ranking_data else 0
    avg_mdd = np.mean([d['mdd'] for d in ranking_data]) if ranking_data else 0

    summary_text = (
        f"전략 수: {total_strategies}개\n"
        f"수익 전략: {profitable_count}개 ({profitable_count/total_strategies*100:.0f}%)\n"
        f"엔진 일치: {consistent_count}개 ({consistent_count/total_strategies*100:.0f}%)\n"
        f"평균 수익률: {avg_return:+.1f}%\n"
        f"평균 MDD: {avg_mdd:.1f}%"
    )
    doc.add_paragraph(summary_text)

    # ----- Engine comparison table -----
    doc.add_heading('5. 6엔진 교차검증 매트릭스', level=1)

    eng_cols = ['전략', 'E1 잔액', 'E2 잔액', 'E3 잔액', 'E4 잔액', 'E5 잔액', 'E6 잔액', '편차%']
    eng_table = doc.add_table(rows=1 + len(ranking_data), cols=len(eng_cols))
    eng_table.style = 'Light Grid Accent 1'
    eng_table.alignment = WdTableAlignment.CENTER

    hdr = eng_table.rows[0]
    for ci, col_name in enumerate(eng_cols):
        cell = hdr.cells[ci]
        cell.text = col_name
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8)

    for ri, strat in enumerate(strategies):
        sn = strat["name"]
        if sn not in all_results:
            continue

        row = eng_table.rows[ri + 1]
        eng_res = all_results[sn]

        vals = [sn]
        bals = []
        for eid in range(1, 7):
            if eid in eng_res:
                b = eng_res[eid][0]
                vals.append(f"${b:,.0f}")
                bals.append(b)
            else:
                vals.append("ERR")

        if len(bals) >= 2:
            mean_b = np.mean(bals)
            dev = (max(bals) - min(bals)) / abs(mean_b) * 100.0 if mean_b != 0 else 999.0
            vals.append(f"{dev:.1f}%")
        else:
            vals.append("-")

        for ci, val in enumerate(vals):
            cell = row.cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(7)

    # Save
    doc.save(output_path)
    print(f"\n  Word 문서 저장 완료: {output_path}")


# =============================================================================
# MAIN
# =============================================================================
def main():
    print("=" * 90)
    print("  v28 전체 기획서 6엔진 교차검증 결과 리뷰")
    print("  25 Strategies x 6 Engines + Yearly Detailed Analysis")
    print("  Initial Capital: $3,000 | Fee: 0.04% x2 | Monthly Loss Limit: -20%")
    print("=" * 90)

    # =========================================================================
    # PHASE 1: Load & Resample
    # =========================================================================
    print("\n[PHASE 1] 데이터 로드 및 리샘플링...")
    t0 = time.time()
    df_5m = load_5m_data(BASE)
    print(f"  5m 로드 완료: {len(df_5m)}행 ({time.time() - t0:.1f}s)")

    print("  리샘플링 중...")
    tf_data = {'5m': df_5m}
    for mins, label in [(10, '10m'), (15, '15m'), (30, '30m'), (60, '1h')]:
        t1 = time.time()
        tf_data[label] = resample_ohlcv(df_5m, mins)
        print(f"  {label}: {len(tf_data[label])}행 ({time.time() - t1:.1f}s)")

    print(f"  Phase 1 완료: {time.time() - t0:.1f}s\n")

    # =========================================================================
    # PHASE 2: JIT Warmup
    # =========================================================================
    print("[PHASE 2] Numba JIT 워밍업...")
    t0 = time.time()
    _dc = np.random.randn(200).astype(np.float64)
    _dh = _dc + 0.5
    _dl = _dc - 0.5
    _dv = np.abs(np.random.randn(200)).astype(np.float64)
    _do = _dc.copy()
    _dt = np.zeros(200, dtype=np.float64)
    _ = calc_ma(_dc, _dv, 0, 5)
    _ = calc_adx_wilder(_dh, _dl, _dc, 14)
    _ = calc_adx_standard_ema(_dh, _dl, _dc, 14)
    _ = calc_rsi_wilder(_dc, 14)
    _fm = calc_ema(_dc, 5)
    _sm = calc_ema(_dc, 20)
    _adx = calc_adx_wilder(_dh, _dl, _dc, 14)
    _rsi = calc_rsi_wilder(_dc, 14)
    _ = backtest_core(_dc, _dh, _dl, _dv, _dt,
                      _fm, _sm, _adx, _rsi,
                      30.0, 30.0, 70.0, -8.0, 5.0, 3.0, 30.0, 10.0,
                      0, 0.0, 0.0004, 0, 50)
    _ = backtest_engine2_hl(_dc, _dh, _dl, _dv, _dt,
                            _fm, _sm, _adx, _rsi,
                            30.0, 30.0, 70.0, -8.0, 5.0, 3.0, 30.0, 10.0,
                            0, 0.0, 0.0004, 0, 50)
    _ = backtest_engine4_nextbar(_dc, _dh, _dl, _do, _dv, _dt,
                                 _fm, _sm, _adx, _rsi,
                                 30.0, 30.0, 70.0, -8.0, 5.0, 3.0, 30.0, 10.0,
                                 0, 0.0, 0.0004, 0, 50)
    print(f"  JIT 워밍업 완료: {time.time() - t0:.1f}s\n")

    # =========================================================================
    # PHASE 3: 6-Engine Cross-Validation
    # =========================================================================
    n_strats = len(strategies)
    n_engines = 6
    total_runs = n_strats * n_engines

    print(f"[PHASE 3] 6엔진 교차검증: {n_strats} 전략 x {n_engines} 엔진 = {total_runs} 백테스트")
    print("-" * 90)

    all_results = {}  # { strategy_name: { engine_id: (bal, trades, wins, losses, pf, mdd, sl, tsl, rev) } }
    run_count = 0

    for s_idx, strat in enumerate(strategies):
        sn = strat["name"]
        print(f"\n  [{s_idx + 1}/{n_strats}] {sn} ({strat['tf']})", flush=True)
        all_results[sn] = {}

        for eng_id in range(1, 7):
            t0 = time.time()
            try:
                res = run_single_engine(eng_id, strat, tf_data)
                all_results[sn][eng_id] = res
                bal, trades, wins, losses, pf, mdd, sl_c, tsl_c, rev_c = res
                elapsed = time.time() - t0
                run_count += 1
                print(f"    E{eng_id}({ENGINE_NAMES[eng_id]:12s}): "
                      f"${bal:>10,.0f} | PF={pf:>6.2f} | MDD={mdd:>5.1f}% | "
                      f"T={trades:>4d} W={wins:>3d} L={losses:>3d} | "
                      f"SL={sl_c:>3d} TSL={tsl_c:>3d} REV={rev_c:>3d} | "
                      f"{elapsed:.1f}s", flush=True)
            except Exception as e:
                run_count += 1
                print(f"    E{eng_id}({ENGINE_NAMES[eng_id]:12s}): ERROR - {e}", flush=True)
                all_results[sn][eng_id] = (3000.0, 0, 0, 0, 0.0, 0.0, 0, 0, 0)

        print(f"  Progress: {run_count}/{total_runs} ({run_count/total_runs*100:.0f}%)", flush=True)

    print(f"\n  Phase 3 완료: 6엔진 교차검증 {run_count}/{total_runs} 완료\n")

    # =========================================================================
    # PHASE 4: Detailed Backtest (Engine 1 only) for yearly tracking
    # =========================================================================
    print(f"[PHASE 4] 상세 백테스트 (Engine 1, 연도별 분석): {n_strats}개 전략")
    print("-" * 90)

    detailed_results = {}  # { strategy_name: detailed_result_dict }
    yearly_stats_all = {}  # { strategy_name: { year: stats } }

    for s_idx, strat in enumerate(strategies):
        sn = strat["name"]
        t0 = time.time()
        print(f"  [{s_idx + 1}/{n_strats}] {sn} 상세 분석 중...", end="", flush=True)

        try:
            det = run_detailed_for_strategy(strat, tf_data)
            detailed_results[sn] = det
            ys = compute_yearly_stats(det)
            yearly_stats_all[sn] = ys

            elapsed = time.time() - t0
            print(f" ${det['balance']:,.0f} | {det['total_trades']}trades | "
                  f"MDD={det['mdd']:.1f}% | {elapsed:.1f}s", flush=True)

            # Print yearly summary
            for year in sorted(ys.keys()):
                y = ys[year]
                print(f"    {year}: T={y['trades']:>3d} | "
                      f"Ret={y['return_pct']:>+7.1f}% | "
                      f"MDD={y['mdd']:>5.1f}% | "
                      f"PF={y['pf']:>6.2f}", flush=True)

        except Exception as e:
            elapsed = time.time() - t0
            print(f" ERROR: {e} ({elapsed:.1f}s)", flush=True)
            detailed_results[sn] = None
            yearly_stats_all[sn] = {}

    print(f"\n  Phase 4 완료\n")

    # =========================================================================
    # PHASE 5: Generate Word Document
    # =========================================================================
    print("[PHASE 5] Word 문서 생성 중...")
    output_path = os.path.join(BASE, "v28_전체기획서_6엔진_교차검증_결과.docx")

    try:
        generate_word_document(all_results, detailed_results, yearly_stats_all, output_path)
        print(f"  성공: {output_path}")
    except ImportError:
        print("  ERROR: python-docx 패키지 미설치. 'pip install python-docx' 실행 필요.")
    except Exception as e:
        print(f"  ERROR: {e}")

    # =========================================================================
    # PHASE 6: Console Summary
    # =========================================================================
    print("\n")
    print("=" * 90)
    print("  최종 결과 요약")
    print("=" * 90)

    # Sort by return
    summary_data = []
    for strat in strategies:
        sn = strat["name"]
        if sn in all_results and 1 in all_results[sn]:
            e1 = all_results[sn][1]
            ret = (e1[0] - 3000.0) / 3000.0 * 100.0
            is_c, dev, reason = check_engine_consistency(all_results[sn])
            summary_data.append((sn, e1[0], ret, e1[1], e1[4], e1[5], is_c, dev))

    summary_data.sort(key=lambda x: x[2], reverse=True)

    print(f"\n{'순위':>4} {'전략':>10} {'잔액':>12} {'수익률':>10} {'거래':>5} {'PF':>7} {'MDD':>7} {'일치':>6} {'편차':>7}")
    print("-" * 80)
    for rank, (sn, bal, ret, trades, pf, mdd, is_c, dev) in enumerate(summary_data, 1):
        status = "O" if is_c else "X"
        print(f"{rank:>4} {sn:>10} ${bal:>10,.0f} {ret:>+9.1f}% {trades:>5d} {pf:>6.2f} {mdd:>6.1f}% {status:>6} {dev:>6.1f}%")

    print("-" * 80)
    print(f"\n  총 전략: {len(summary_data)}개")
    print(f"  수익: {len([d for d in summary_data if d[2] > 0])}개")
    print(f"  손실: {len([d for d in summary_data if d[2] <= 0])}개")
    print(f"  엔진 일치: {len([d for d in summary_data if d[6]])}개")
    print(f"\n  완료!")


if __name__ == '__main__':
    main()
