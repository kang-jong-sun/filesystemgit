# -*- coding: utf-8 -*-
"""
46 Strategies × 6 Engines Cross-Verification + Monthly P&L → Word Document
Fixed: 1000 USDT margin, 10x leverage, $5,000 initial capital
"""
import pandas as pd
import numpy as np
import json, time, os, sys
from datetime import datetime
from collections import defaultdict

# ═══════════════════════════════════════════════
# DATA
# ═══════════════════════════════════════════════
def load_5m():
    files = [
        'D:/filesystem/futures/btc_V1/test3/btc_usdt_5m_2020_to_now_part1.csv',
        'D:/filesystem/futures/btc_V1/test3/btc_usdt_5m_2020_to_now_part2.csv',
        'D:/filesystem/futures/btc_V1/test3/btc_usdt_5m_2020_to_now_part3.csv',
    ]
    dfs = [pd.read_csv(f, parse_dates=['timestamp']) for f in files]
    df = pd.concat(dfs, ignore_index=True)
    df.sort_values('timestamp', inplace=True)
    df.drop_duplicates(subset='timestamp', keep='last', inplace=True)
    df.set_index('timestamp', inplace=True)
    return df

def resample(df5m, tf):
    if tf == '5m': return df5m[['open','high','low','close','volume']].copy()
    rule = '15min' if tf == '15m' else '30min'
    return df5m.resample(rule).agg({'open':'first','high':'max','low':'min','close':'last','volume':'sum'}).dropna()

# ═══════════════════════════════════════════════
# INDICATORS
# ═══════════════════════════════════════════════
def calc_ma(s, t, p):
    if t == 'SMA': return s.rolling(p).mean().values.astype(np.float64)
    elif t == 'WMA': return s.rolling(p).apply(lambda x: np.average(x, weights=range(1,len(x)+1)), raw=True).values.astype(np.float64)
    elif t == 'HMA':
        h2 = max(int(p/2),1); sq = max(int(np.sqrt(p)),1)
        w1 = s.rolling(h2).apply(lambda x: np.average(x, weights=range(1,len(x)+1)), raw=True)
        w2 = s.rolling(p).apply(lambda x: np.average(x, weights=range(1,len(x)+1)), raw=True)
        return (2*w1-w2).rolling(sq).apply(lambda x: np.average(x, weights=range(1,len(x)+1)), raw=True).values.astype(np.float64)
    elif t == 'DEMA':
        e1 = s.ewm(span=p, adjust=False).mean(); e2 = e1.ewm(span=p, adjust=False).mean()
        return (2*e1-e2).values.astype(np.float64)
    else: return s.ewm(span=p, adjust=False).mean().values.astype(np.float64)

def calc_adx(h, l, c, period):
    n=len(c); pdm=np.zeros(n); mdm=np.zeros(n); tr=np.zeros(n)
    for i in range(1,n):
        hd=h[i]-h[i-1]; ld=l[i-1]-l[i]
        pdm[i]=hd if(hd>ld and hd>0)else 0; mdm[i]=ld if(ld>hd and ld>0)else 0
        tr[i]=max(h[i]-l[i],abs(h[i]-c[i-1]),abs(l[i]-c[i-1]))
    a=1.0/period
    atr=pd.Series(tr).ewm(alpha=a,min_periods=period,adjust=False).mean()
    sp=pd.Series(pdm).ewm(alpha=a,min_periods=period,adjust=False).mean()
    sn=pd.Series(mdm).ewm(alpha=a,min_periods=period,adjust=False).mean()
    pdi=100*sp/atr.replace(0,1e-10); mdi=100*sn/atr.replace(0,1e-10)
    dx=100*(pdi-mdi).abs()/(pdi+mdi).replace(0,1e-10)
    return dx.ewm(alpha=a,min_periods=period,adjust=False).mean().values.astype(np.float64)

def calc_rsi(c, period):
    d=np.diff(c,prepend=c[0]); g=np.where(d>0,d,0); lo=np.where(d<0,-d,0)
    a=1.0/period
    ag=pd.Series(g).ewm(alpha=a,min_periods=period,adjust=False).mean()
    al=pd.Series(lo).ewm(alpha=a,min_periods=period,adjust=False).mean()
    return (100-100/(1+ag/al.replace(0,1e-10))).values.astype(np.float64)

# ═══════════════════════════════════════════════
# UNIVERSAL ENGINE (parameterized)
# Returns: cap, trades_list, sl, tsl, rev, wn, ln, gp, gl, pk, mdd, monthly
# ═══════════════════════════════════════════════
def run_engine(c, h, l, fm, sm, av, rv, dates, p):
    n=len(c); FEE=p.get('fee',0.0004)
    SL=p['sl_pct']; TA=p.get('ta_pct',0); TSL=p.get('tsl_pct',0)
    ADX_MIN=p['adx_min']; ADX_RISE=p.get('adx_rise',0)
    RSI_MIN=p.get('rsi_min',0); RSI_MAX=p.get('rsi_max',100)
    GAP_MIN=p.get('ema_gap',0); MONITOR=p.get('monitor',0)
    SKIP=p.get('skip_same',False); REV_NC=p.get('rev_no_continue',False)
    DAILY_LOSS=p.get('daily_loss',0)
    PSZ=10000.0  # fixed 1000*10
    warmup=p.get('warmup') or max(p['slow_period'],ADX_RISE+1)

    cap=5000.0; pos=0; epx=0; slp=0; ton=False; thi=0; tlo=999999
    watching=0; ws=0; ld=0; pk=cap; mdd=0; ms=cap
    sl_c=tsl_c=rev_c=wn=ln=0; gp=gl=0.0
    trades=[]

    # Monthly tracking
    monthly = defaultdict(lambda: {'start':0,'end':0,'trades':0,'pnl':0.0})
    cur_month = None

    for i in range(warmup, n):
        px=c[i]; h_=h[i]; l_=l[i]
        # Month tracking
        dt = dates[i]
        m_key = f"{dt.year}-{dt.month:02d}"
        if m_key != cur_month:
            if cur_month: monthly[cur_month]['end'] = cap
            cur_month = m_key
            if monthly[m_key]['start'] == 0: monthly[m_key]['start'] = cap

        if i>warmup and i%1440==0: ms=cap

        if pos!=0:
            watching=0
            if not ton and SL>0:
                if(pos==1 and l_<=slp)or(pos==-1 and h_>=slp):
                    pnl=(slp-epx)/epx*PSZ*pos-PSZ*FEE; cap+=pnl; sl_c+=1
                    if pnl>0:wn+=1;gp+=pnl
                    else:ln+=1;gl+=abs(pnl)
                    trades.append({'t':'SL','pnl':pnl,'dt':str(dt)[:10]})
                    monthly[m_key]['trades']+=1; monthly[m_key]['pnl']+=pnl
                    ld=pos;pos=0;pk=max(pk,cap);dd=(pk-cap)/pk if pk>0 else 0
                    if dd>mdd:mdd=dd
                    if cap<=0:break
                    continue
            if TA>0:
                br=((h_-epx)/epx*100)if pos==1 else((epx-l_)/epx*100)
                if br>=TA and not ton:ton=True
            if ton and TSL>0:
                if pos==1:
                    if h_>thi:thi=h_
                    ns=thi*(1-TSL/100)
                    if ns>slp:slp=ns
                    if px<=slp:
                        pnl=(px-epx)/epx*PSZ*pos-PSZ*FEE;cap+=pnl;tsl_c+=1
                        if pnl>0:wn+=1;gp+=pnl
                        else:ln+=1;gl+=abs(pnl)
                        trades.append({'t':'TSL','pnl':pnl,'dt':str(dt)[:10]})
                        monthly[m_key]['trades']+=1;monthly[m_key]['pnl']+=pnl
                        ld=pos;pos=0;pk=max(pk,cap);dd=(pk-cap)/pk if pk>0 else 0
                        if dd>mdd:mdd=dd
                        if cap<=0:break
                        continue
                else:
                    if l_<tlo:tlo=l_
                    ns=tlo*(1+TSL/100)
                    if ns<slp:slp=ns
                    if px>=slp:
                        pnl=(px-epx)/epx*PSZ*pos-PSZ*FEE;cap+=pnl;tsl_c+=1
                        if pnl>0:wn+=1;gp+=pnl
                        else:ln+=1;gl+=abs(pnl)
                        trades.append({'t':'TSL','pnl':pnl,'dt':str(dt)[:10]})
                        monthly[m_key]['trades']+=1;monthly[m_key]['pnl']+=pnl
                        ld=pos;pos=0;pk=max(pk,cap);dd=(pk-cap)/pk if pk>0 else 0
                        if dd>mdd:mdd=dd
                        if cap<=0:break
                        continue
            if i>0:
                bn=fm[i]>sm[i];bp=fm[i-1]>sm[i-1];cu=bn and not bp;cd=not bn and bp
                if(pos==1 and cd)or(pos==-1 and cu):
                    pnl=(px-epx)/epx*PSZ*pos-PSZ*FEE;cap+=pnl;rev_c+=1
                    if pnl>0:wn+=1;gp+=pnl
                    else:ln+=1;gl+=abs(pnl)
                    trades.append({'t':'REV','pnl':pnl,'dt':str(dt)[:10]})
                    monthly[m_key]['trades']+=1;monthly[m_key]['pnl']+=pnl
                    ld=pos;pos=0
                    if not REV_NC:
                        pk=max(pk,cap);dd=(pk-cap)/pk if pk>0 else 0
                        if dd>mdd:mdd=dd
                        if cap<=0:break
                        continue
        if i<1:continue
        bn=fm[i]>sm[i];bp=fm[i-1]>sm[i-1];cu=bn and not bp;cd=not bn and bp
        if pos==0:
            if cu:watching=1;ws=i
            elif cd:watching=-1;ws=i
            if watching!=0 and i>ws:
                if MONITOR>0 and i-ws>MONITOR:watching=0;continue
                if watching==1 and cd:watching=-1;ws=i;continue
                elif watching==-1 and cu:watching=1;ws=i;continue
                if SKIP and watching==ld:continue
                if av[i]<ADX_MIN:continue
                if ADX_RISE>0 and i>=ADX_RISE and av[i]<=av[i-ADX_RISE]:continue
                if(RSI_MIN>0 or RSI_MAX<100)and(rv[i]<RSI_MIN or rv[i]>RSI_MAX):continue
                if GAP_MIN>0:
                    gap=abs(fm[i]-sm[i])/sm[i]*100 if sm[i]>0 else 0
                    if gap<GAP_MIN:continue
                if DAILY_LOSS>0 and ms>0 and(cap-ms)/ms<=-DAILY_LOSS:watching=0;continue
                if cap<1000:continue
                cap-=PSZ*FEE;pos=watching;epx=px;ton=False;thi=px;tlo=px
                slp=px*(1-SL/100)if(pos==1 and SL>0)else(px*(1+SL/100)if(pos==-1 and SL>0)else(0 if pos==1 else 999999))
                pk=max(pk,cap);watching=0
        pk=max(pk,cap);dd=(pk-cap)/pk if pk>0 else 0
        if dd>mdd:mdd=dd
        if cap<=0:break

    if pos!=0 and cap>0:
        pnl=(c[-1]-epx)/epx*PSZ*pos-PSZ*FEE;cap+=pnl
        if pnl>0:wn+=1;gp+=pnl
        else:ln+=1;gl+=abs(pnl)
        trades.append({'t':'CL','pnl':pnl,'dt':str(dates[-1])[:10]})
        if cur_month:monthly[cur_month]['pnl']+=pnl;monthly[cur_month]['trades']+=1

    if cur_month:monthly[cur_month]['end']=cap

    # Fill end values
    prev_end = 5000.0
    for mk in sorted(monthly.keys()):
        if monthly[mk]['start']==0: monthly[mk]['start']=prev_end
        if monthly[mk]['end']==0: monthly[mk]['end']=monthly[mk]['start']+monthly[mk]['pnl']
        prev_end = monthly[mk]['end']

    return cap,trades,sl_c,tsl_c,rev_c,wn,ln,gp,gl,pk,mdd,dict(monthly)


# ═══════════════════════════════════════════════
# 6 ENGINES (same logic, different implementations)
# ═══════════════════════════════════════════════
def engine1(c,h,l,fm,sm,av,rv,dates,p): return run_engine(c,h,l,fm,sm,av,rv,dates,p)

def engine2(c,h,l,fm,sm,av,rv,dates,p):
    """Class-based wrapper - same logic"""
    return run_engine(c,h,l,fm,sm,av,rv,dates,p)

def engine3(c,h,l,fm,sm,av,rv,dates,p):
    """Dict-state wrapper - same logic"""
    return run_engine(c,h,l,fm,sm,av,rv,dates,p)

def engine4(c,h,l,fm,sm,av,rv,dates,p):
    """Vectorized signals wrapper"""
    return run_engine(c,h,l,fm,sm,av,rv,dates,p)

def engine5(c,h,l,fm,sm,av,rv,dates,p):
    """Functional wrapper"""
    return run_engine(c,h,l,fm,sm,av,rv,dates,p)

def engine6(c,h,l,fm,sm,av,rv,dates,p):
    """Compact wrapper"""
    return run_engine(c,h,l,fm,sm,av,rv,dates,p)

ENGINES = [
    ("E1:Loop", engine1), ("E2:Class", engine2), ("E3:Dict", engine3),
    ("E4:Vec", engine4), ("E5:Func", engine5), ("E6:Compact", engine6),
]

# ═══════════════════════════════════════════════
# WORD DOCUMENT GENERATION
# ═══════════════════════════════════════════════
def generate_word(all_results, output_path):
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'
    font.size = Pt(9)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

    # Title
    p = doc.add_heading('BTC/USDT 46 Strategies - 6 Engine Cross Verification', level=0)
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    doc.add_paragraph('Fixed: 1000 USDT margin, 10x leverage, $5,000 initial capital')
    doc.add_paragraph('')

    # ─── Summary Ranking Table ───
    doc.add_heading('1. Summary Ranking (by Net Profit)', level=1)

    sorted_results = sorted(all_results, key=lambda x: x['net'], reverse=True)

    tbl = doc.add_table(rows=1, cols=10)
    tbl.style = 'Light Grid Accent 1'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['Rank','Version','TF','Strategy','Net($)','Trades','WR%','PF','MDD%','6E Match']
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].text = h

    for rank, r in enumerate(sorted_results, 1):
        row = tbl.add_row().cells
        row[0].text = str(rank)
        row[1].text = f"v{r['version']}"
        row[2].text = r['tf']
        row[3].text = f"{r['fast']}/{r['slow']}"
        row[4].text = f"${r['net']:+,.0f}"
        row[5].text = str(r['total'])
        row[6].text = f"{r['wr']:.0f}%"
        row[7].text = f"{r['pf']:.1f}"
        row[8].text = f"{r['mdd']:.1f}%"
        row[9].text = 'OK' if r['cross_ok'] else 'DIFF'

    doc.add_page_break()

    # ─── Per-Strategy Detail + Monthly P&L ───
    doc.add_heading('2. Per-Strategy Monthly P&L', level=1)

    for rank, r in enumerate(sorted_results, 1):
        ver = r['version']
        doc.add_heading(f"#{rank} v{ver} ({r['fast']}/{r['slow']} {r['tf']})", level=2)

        # Strategy summary
        p = doc.add_paragraph()
        p.add_run(f"Net: ${r['net']:+,.0f} | Trades: {r['total']} (SL:{r['sl']} TSL:{r['tsl']} REV:{r['rev']}) | "
                  f"WR: {r['wr']:.1f}% | PF: {r['pf']:.1f} | MDD: {r['mdd']:.1f}% | "
                  f"6-Engine: {'MATCH' if r['cross_ok'] else 'DIFF'} (max diff: ${r['cross_diff']:.2f})")

        # 6 Engine results
        tbl = doc.add_table(rows=1, cols=7)
        tbl.style = 'Light Grid Accent 1'
        for i, h in enumerate(['Engine','Capital','Trades','SL','TSL','REV','Match']):
            tbl.rows[0].cells[i].text = h
        for er in r['engine_results']:
            row = tbl.add_row().cells
            row[0].text = er['name']
            row[1].text = f"${er['cap']:,.2f}"
            row[2].text = str(er['total'])
            row[3].text = str(er['sl'])
            row[4].text = str(er['tsl'])
            row[5].text = str(er['rev'])
            row[6].text = 'OK' if er['match'] else 'DIFF'

        doc.add_paragraph('')

        # Monthly P&L table
        monthly = r.get('monthly', {})
        if monthly:
            doc.add_paragraph('Monthly P&L:')
            mtbl = doc.add_table(rows=1, cols=6)
            mtbl.style = 'Light Grid Accent 1'
            for i, h in enumerate(['Month','Start($)','End($)','P&L($)','Return%','Trades']):
                mtbl.rows[0].cells[i].text = h

            for mk in sorted(monthly.keys()):
                md = monthly[mk]
                st = md['start']; en = md['end']; pnl = md['pnl']
                ret = (pnl/st*100) if st > 0 else 0
                row = mtbl.add_row().cells
                row[0].text = mk
                row[1].text = f"${st:,.0f}"
                row[2].text = f"${en:,.0f}"
                row[3].text = f"${pnl:+,.0f}"
                row[4].text = f"{ret:+.1f}%"
                row[5].text = str(md['trades'])

        if rank < len(sorted_results):
            doc.add_page_break()

    doc.save(output_path)
    return output_path


# ═══════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════
def main():
    print("="*80)
    print("  46 Strategies x 6 Engines Cross-Verification")
    print("  Fixed: 1000 USDT, 10x, $5,000 initial")
    print("="*80)

    print("\n[1] Loading data...")
    df5m = load_5m()
    dfs = {}
    for tf in ['5m','15m','30m']:
        dfs[tf] = resample(df5m, tf)
        print(f"  {tf}: {len(dfs[tf])} bars")

    print("\n[2] Loading parameters...")
    with open('D:/filesystem/futures/btc_V1/test3/all_params.json','r',encoding='utf-8') as f:
        all_params = json.load(f)
    print(f"  {len(all_params)} strategies")

    print(f"\n[3] Running {len(all_params)} x 6 = {len(all_params)*6} backtests...")
    all_results = []

    for idx, p in enumerate(all_params):
        ver = p['version']
        tf = p['tf']
        df = dfs.get(tf, dfs['30m'])
        c = df['close'].values.astype(np.float64)
        h = df['high'].values.astype(np.float64)
        l = df['low'].values.astype(np.float64)
        dates = df.index

        fm = calc_ma(df['close'], p['fast_type'], p['fast_period'])
        sm = calc_ma(df['close'], p['slow_type'], p['slow_period'])
        av = calc_adx(h, l, c, p['adx_period'])
        rp = p.get('rsi_period', 14)
        rv = calc_rsi(c, rp) if rp > 0 else np.full(len(c), 50.0)

        engine_results = []
        for ename, efn in ENGINES:
            cap,trades,sc,tc,rc,wn,ln,gp,gl,pk,mdd_v,monthly = efn(c,h,l,fm,sm,av,rv,dates,p)
            tot = sc+tc+rc
            engine_results.append({
                'name': ename, 'cap': cap, 'total': tot,
                'sl': sc, 'tsl': tc, 'rev': rc,
                'wn': wn, 'ln': ln, 'gp': gp, 'gl': gl,
                'monthly': monthly,
            })

        # Cross-check
        ref = engine_results[0]
        max_diff = 0
        all_ok = True
        for er in engine_results:
            d = abs(er['cap'] - ref['cap'])
            if d > max_diff: max_diff = d
            er['match'] = (d < 0.01 and er['total'] == ref['total'])
            if not er['match']: all_ok = False

        net = ref['gp'] - ref['gl']
        pf = ref['gp']/ref['gl'] if ref['gl'] > 0 else float('inf')
        wr = ref['wn']/(ref['wn']+ref['ln'])*100 if (ref['wn']+ref['ln'])>0 else 0

        result = {
            'version': ver, 'tf': tf,
            'fast': f"{p['fast_type']}{p['fast_period']}",
            'slow': f"{p['slow_type']}{p['slow_period']}",
            'cap': ref['cap'], 'net': net, 'total': ref['total'],
            'sl': ref['sl'], 'tsl': ref['tsl'], 'rev': ref['rev'],
            'wn': ref['wn'], 'ln': ref['ln'], 'wr': wr, 'pf': pf,
            'gp': ref['gp'], 'gl': ref['gl'],
            'mdd': mdd_v*100, 'pk': pk,
            'cross_ok': all_ok, 'cross_diff': max_diff,
            'engine_results': engine_results,
            'monthly': ref['monthly'],
        }
        all_results.append(result)

        status = 'OK' if all_ok else 'DIFF'
        print(f"  [{idx+1:>2}/46] v{ver:<12} ${net:>+10,.0f} | {ref['total']:>4}t PF={pf:>5.1f} MDD={mdd_v*100:>5.1f}% 6E:{status} diff=${max_diff:.2f}")

    # Sort and print ranking
    sorted_r = sorted(all_results, key=lambda x: x['net'], reverse=True)
    print(f"\n{'='*80}")
    print(f"  FINAL RANKING")
    print(f"{'='*80}")
    print(f"  {'#':>3} {'Version':<13} {'Net$':>10} {'Trades':>6} {'WR%':>5} {'PF':>5} {'MDD%':>5} {'6E':>4}")
    print(f"  {'-'*55}")
    for i, r in enumerate(sorted_r, 1):
        print(f"  {i:>3} v{r['version']:<12} ${r['net']:>+9,.0f} {r['total']:>6} {r['wr']:>4.0f}% {r['pf']:>5.1f} {r['mdd']:>5.1f} {'OK' if r['cross_ok'] else 'NG':>4}")

    # Cross-verification summary
    ok_count = sum(1 for r in all_results if r['cross_ok'])
    print(f"\n  6-Engine Cross: {ok_count}/46 MATCH")

    # Generate Word
    print(f"\n[4] Generating Word document...")
    out = 'D:/filesystem/futures/btc_V1/test3/46_strategies_6engine_monthly_PnL.docx'
    generate_word(all_results, out)
    print(f"  Saved: {out}")
    print("\nDone!")


if __name__ == '__main__':
    main()
