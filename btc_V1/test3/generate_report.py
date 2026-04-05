# -*- coding: utf-8 -*-
"""Generate Word Report: 21 Versions Verification & Analysis"""
import sys, json, os
sys.stdout.reconfigure(encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

DATA_DIR = os.path.dirname(os.path.abspath(__file__))

def set_cell(cell, text, bold=False, size=8, align='center', color=None):
    p = cell.paragraphs[0]
    p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER,
                    'right': WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = 'Malgun Gothic'
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for j, h in enumerate(headers):
        set_cell(table.rows[0].cells[j], h, bold=True, size=8)
    # Data
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            color = None
            if isinstance(val, str) and val.startswith('+'):
                color = (0, 100, 0)
            elif isinstance(val, str) and val.startswith('-') and '%' in val:
                color = (200, 0, 0)
            set_cell(table.rows[i+1].cells[j], val, size=8, color=color)
    return table

def main():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # ============================================================
    # TITLE
    # ============================================================
    title = doc.add_heading('BTC/USDT 선물 자동매매 시스템', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('21개 버전 통합 검증 레포트 v15.3', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('작성일: 2026-03-27 | 데이터: BTC/USDT 5분봉 75개월 (655,399캔들) | 수수료: 0.04% x 2')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # ============================================================
    # 1. 개요
    # ============================================================
    doc.add_heading('1. 검증 개요', level=1)
    doc.add_paragraph(
        '본 레포트는 v10.1부터 v15.5까지 총 21개 기획서의 전략 설정값을 추출하고, '
        '동일한 백테스트 엔진(자체 구축)으로 각 30회씩 반복 실행하여 기획서 주장값의 정확성과 '
        '전략 성능을 객관적으로 검증한 결과입니다.'
    )
    doc.add_paragraph('검증 조건:', style='List Bullet')
    items = [
        '초기 자본: $3,000 USDT',
        '데이터: 2019-12 ~ 2026-03 (약 75개월, 655,399캔들)',
        '수수료: 테이커 0.04% (진입+청산 모두 반영)',
        '마진 모드: 격리마진 (ISOLATED)',
        '백테스트 엔진: Wilder RSI/ADX, EMA/HMA/WMA/SMA 지원, 트레일링 스톱(high/low 추적, 종가 판단)',
        '반복 횟수: 각 버전 30회 (결정적 시스템이므로 일관성 검증 목적)',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ============================================================
    # 2. 전체 랭킹
    # ============================================================
    doc.add_heading('2. 21개 버전 전체 랭킹 (30회 검증 결과)', level=1)

    # Load results
    with open(os.path.join(DATA_DIR, 'verify_all_results.json'), encoding='utf-8') as f:
        results = json.load(f)

    headers = ['순위', '버전', 'TF', '실측 잔액', '실측 수익률', '실측 PF', '실측 MDD',
               '거래', 'FL', '30x일관', '기획서잔액', '잔액차이', '상태', '점수']

    rows = []
    for i, r in enumerate(results):
        rows.append([
            str(i+1),
            r['ver'],
            r.get('tf', '?'),
            f"${r['actual_bal']:,.0f}",
            f"{r['actual_ret']:+,.1f}%",
            f"{r['actual_pf']:.2f}",
            f"{r['actual_mdd']:.1f}%",
            str(r.get('actual_tr', 0)),
            str(r.get('actual_fl', 0)),
            'O' if r.get('consistent', True) else 'X',
            f"${r['claimed_bal']:,.0f}",
            f"{r.get('bal_diff', 0):.1f}%",
            r['status'],
            f"{r.get('score', -999):.1f}",
        ])

    add_table(doc, headers, rows)

    p = doc.add_paragraph()
    run = p.add_run('\n* MATCH: 기획서 대비 25% 이내 | CLOSE: 50% 이내 | DIFFER: 50% 초과')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

    # ============================================================
    # 3. BEST 3
    # ============================================================
    doc.add_heading('3. BEST 3 선정', level=1)

    # BEST 1: v14.3
    doc.add_heading('3-1. 1위: v14.3 (고수익형)', level=2)
    doc.add_paragraph(
        '5분봉 EMA(7/100) + ADX(14)>=30 + RSI(14) 30~58\n'
        'SL: -10% | Trail: +8%/-6% | 10x 30% | ML-25% CP3/288 DD-30%'
    )
    best1 = [
        ['최종 잔액', '$3,210,275', '수익률', '+106,909%'],
        ['PF', '71.60', 'MDD', '73.7%'],
        ['거래', '274회', 'FL', '27회'],
        ['SL 발동', '0회', 'TSL', '60회'],
        ['기획서 일치', 'MATCH (0%)', '30x 일관', 'YES'],
    ]
    t = doc.add_table(rows=len(best1), cols=4)
    t.style = 'Light Grid Accent 1'
    for i, row in enumerate(best1):
        for j, val in enumerate(row):
            set_cell(t.rows[i].cells[j], val, bold=(j % 2 == 0), size=9)

    doc.add_paragraph()
    doc.add_paragraph('선정 이유:', style='Heading 3')
    reasons1 = [
        '절대 수익 $3.2M으로 전 버전 중 압도적 1위 (2위 대비 6배)',
        'PF 71.60으로 손실 $1당 수익 $71.60 회수하는 극단적 손익비 구조',
        '기획서 주장값과 실측이 완벽히 일치 (0% 차이) — 가장 신뢰할 수 있는 기획서',
        '2024~2026년 연속 대형 수익 (+223%, +4,091%, +194%)',
        'SL -10%가 강제청산과 동일 거리 → SL 발동 0회, 역신호/트레일링이 핵심 청산 수단',
    ]
    for r in reasons1:
        doc.add_paragraph(r, style='List Bullet')

    doc.add_paragraph('리스크:', style='Heading 3')
    risks1 = [
        'FL 27회 (격리마진이므로 증거금 30%만 손실, 전체 계좌 파산 아님)',
        'MDD 73.7% (2020-2021 코로나 구간, 2023 이후 MDD 40% 미만)',
        '2020년 -64.5% 대형 손실 구간 존재',
    ]
    for r in risks1:
        doc.add_paragraph(r, style='List Bullet')

    # BEST 2: v13.3/v13.5
    doc.add_heading('3-2. 2위: v13.3/v13.5 (균형형)', level=2)
    doc.add_paragraph(
        '5분봉 EMA(7/100) + ADX(14)>=30 + RSI(14) 30~58\n'
        'SL: -7% | Trail: +8%/-6% | 10x 20% | ML-20% CP3/288 DD-50%'
    )
    best2 = [
        ['최종 잔액', '$509,639', '수익률', '+16,888%'],
        ['PF', '15.83', 'MDD', '69.3%'],
        ['거래', '309회', 'FL', '1회'],
        ['기획서 일치', 'MATCH (8.8%)', '30x 일관', 'YES'],
    ]
    t = doc.add_table(rows=len(best2), cols=4)
    t.style = 'Light Grid Accent 1'
    for i, row in enumerate(best2):
        for j, val in enumerate(row):
            set_cell(t.rows[i].cells[j], val, bold=(j % 2 == 0), size=9)

    doc.add_paragraph()
    doc.add_paragraph('선정 이유:', style='Heading 3')
    reasons2 = [
        'PF 15.83으로 높은 손익비 유지하면서 FL 1회로 안정적',
        'SL -7%가 강제청산 여유 3%p 확보 → FL을 1회로 억제',
        'v13.3과 v13.5가 동일 결과 = 전략 설정의 일관성 확인',
        '보호 메커니즘(ML-20%, CP3, DD-50%)이 2020-2021 손실 제한',
        'v14.3 대비 MDD 4.4%p 낮고 FL 26회 적음',
    ]
    for r in reasons2:
        doc.add_paragraph(r, style='List Bullet')

    # BEST 3: v14.4
    doc.add_heading('3-3. 3위: v14.4 (안전형)', level=2)
    doc.add_paragraph(
        '30분봉 EMA(3/200) + ADX(14)>=35 + RSI(14) 30~65\n'
        'SL: -7% | Trail: +6%/-3% | 10x 25% | ML-20%'
    )
    best3 = [
        ['최종 잔액', '$670,088', '수익률', '+22,236%'],
        ['PF', '2.47', 'MDD', '32.3%'],
        ['거래', '105회', 'FL', '0회'],
        ['기획서 일치', 'MATCH (20%)', '30x 일관', 'YES'],
    ]
    t = doc.add_table(rows=len(best3), cols=4)
    t.style = 'Light Grid Accent 1'
    for i, row in enumerate(best3):
        for j, val in enumerate(row):
            set_cell(t.rows[i].cells[j], val, bold=(j % 2 == 0), size=9)

    doc.add_paragraph()
    doc.add_paragraph('선정 이유:', style='Heading 3')
    reasons3 = [
        'MDD 32.3%로 전 버전 중 최저 — 실전 운용에 가장 적합',
        'FL 0회 — 강제청산 완전 제거',
        '2020~2025 전 기간 연속 수익 (2020: +189%, 2021: +286%)',
        '5분봉 전략들이 -57~-85% 손실인 2020-2021에서 대형 수익 달성',
        '거래 월 1.3회로 낮은 수수료 부담',
    ]
    for r in reasons3:
        doc.add_paragraph(r, style='List Bullet')

    doc.add_page_break()

    # ============================================================
    # 4. BEST 3 연도별 비교
    # ============================================================
    doc.add_heading('4. BEST 3 연도별 수익률 비교', level=1)

    yr_headers = ['연도', 'v14.3 (1위)', 'v13.5 (2위)', 'v14.4 (3위)']
    yr_data = {
        'v14.3': {'2020': -64, '2021': +4, '2022': +183, '2023': +157, '2024': +223, '2025': +4091, '2026': +194},
        'v13.5': {'2020': -57, '2021': -2, '2022': +192, '2023': +56, '2024': +331, '2025': +859, '2026': +112},
        'v14.4': {'2020': +189, '2021': +286, '2022': +43, '2023': +328, '2024': +100, '2025': +86, '2026': -12},
    }
    yr_rows = []
    for y in ['2020','2021','2022','2023','2024','2025','2026']:
        yr_rows.append([
            y,
            f"{yr_data['v14.3'].get(y, 0):+,.0f}%",
            f"{yr_data['v13.5'].get(y, 0):+,.0f}%",
            f"{yr_data['v14.4'].get(y, 0):+,.0f}%",
        ])
    add_table(doc, yr_headers, yr_rows)

    doc.add_paragraph()
    doc.add_paragraph(
        '핵심 인사이트: v14.3과 v13.5(5분봉)는 2024-2025년에 폭발적 성장을 보이나 2020-2021에서 손실. '
        'v14.4(30분봉)는 전 기간 안정적이나 2025년 +86%로 상대적 저성장. '
        '실전에서는 v14.4로 시작하여 시스템 검증 후 v13.5 또는 v14.3으로 전환 권장.'
    )

    doc.add_page_break()

    # ============================================================
    # 5. 폐기 대상 10개
    # ============================================================
    doc.add_heading('5. 폐기 대상 10개 기획서', level=1)

    discard = [
        ('v11.1', '$10 (-99.7%)', '99.7%', 'ADX 20 과소 → 923회 과다거래, 계좌 파산. 기획서 $35,492 vs 실측 $10 (100% 불일치)'),
        ('v12.5', '$3,529 (+18%)', '98.9%', '75개월간 +18% (연 2.8%). MDD 98.9% 파산 직전. 기획서 $233,881 vs 실측 $3,529 (98.5% 불일치)'),
        ('v10.1', '$4,957 (+65%)', '98.0%', 'MDD 98% ($3,000→$60 하락). FL 10회. 보호 메커니즘 없음. 실전 불가'),
        ('v12.2', '$8,195 (+173%)', '99.0%', '트레일링 1% 과도 타이트. 기획서 MDD 14.4% vs 실측 99.0%. 기획서 신뢰 불가'),
        ('v12.0', '$19,101 (+537%)', '99.1%', 'MDD 99.1%, FL 11회. 보호 없음. v12.3과 동일 결과 (중복)'),
        ('v12.3', '$19,101 (+537%)', '99.1%', '기획서 $252,892 vs 실측 $19,101 (92.4% 불일치). v12.0과 중복'),
        ('v13.2', '$12,494 (+317%)', '49.6%', 'EMA(30/200) 저빈도 크로스. PF 1.80 저조. v14.4가 상위 호환'),
        ('v14.2F', '$67,365 (+2,146%)', '61.1%', '기획서 $798,358 vs 실측 $67,365 (91.6% 불일치). PF 1.22 최저'),
        ('v13.4', '$31,253 (+942%)', '39.0%', '기획서 $161,046 vs 실측 $31,253 (80.6% 불일치). v14.4가 상위 호환'),
        ('v14.1', '$225,828 (+7,428%)', '76.4%', 'FL 7회. 기획서 $124,882 vs 실측 $225,828 (80.8% 불일치). v13.0과 유사하나 열위'),
    ]

    disc_headers = ['#', '버전', '실측 잔액', 'MDD', '폐기 사유']
    disc_rows = []
    for i, (ver, bal, mdd, reason) in enumerate(discard):
        disc_rows.append([str(i+1), ver, bal, mdd, reason])
    add_table(doc, disc_headers, disc_rows)

    doc.add_paragraph()
    doc.add_heading('폐기 판단 기준', level=2)
    criteria = [
        '기획서 주장값 대비 실측 70% 이상 차이 → 기획서 신뢰 불가',
        'MDD 95% 이상 → 사실상 파산 수준, 실전 운용 불가능',
        'PF 1.5 미만 → 수익 구조 불충분',
        'FL 10회 이상 → 잦은 강제청산으로 복리 침식',
        '상위 호환 존재 → 동일 구조의 더 나은 버전이 있음',
    ]
    for c in criteria:
        doc.add_paragraph(c, style='List Bullet')

    doc.add_page_break()

    # ============================================================
    # 6. 모델별 월별 손익 (v14.3)
    # ============================================================
    doc.add_heading('6. 최고수익 모델 (v14.3) 월별 손익 상세', level=1)

    try:
        with open(os.path.join(DATA_DIR, 'v143_monthly_data.json'), encoding='utf-8') as f:
            monthly = json.load(f)
        ma = monthly.get('A', {})

        m_headers = ['월', '손익금', '손익률', '누적잔액', '거래', 'TSL', 'REV', 'FL']
        m_rows = []
        year_pnl = 0; year_key = ''; year_start = 3000; year_tr = 0

        for m in sorted(ma.keys()):
            d = ma[m]
            y = m[:4]
            if y != year_key and year_key != '':
                yr = year_pnl / year_start * 100 if year_start > 0 else 0
                m_rows.append([f'{year_key} 합계', f'${year_pnl:+,.0f}', f'{yr:+.1f}%', '', str(year_tr), '', '', ''])
                year_start = d['bal'] - d['pnl']; year_pnl = 0; year_tr = 0
            if year_key == '': year_start = 3000
            year_key = y; year_pnl += d['pnl']; year_tr += d['tr']

            m_rows.append([
                m, f"${d['pnl']:+,.0f}", f"{d['pct']:+.1f}%", f"${d['bal']:,.0f}",
                str(d['tr']), str(d['tsl']), str(d['rev']), str(d['fl'])
            ])

        if year_key:
            yr = year_pnl / year_start * 100 if year_start > 0 else 0
            m_rows.append([f'{year_key} 합계', f'${year_pnl:+,.0f}', f'{yr:+.1f}%', '', str(year_tr), '', '', ''])

        add_table(doc, m_headers, m_rows)
    except Exception as e:
        doc.add_paragraph(f'월별 데이터 로드 실패: {e}')

    doc.add_page_break()

    # ============================================================
    # 7. 핵심 설정값
    # ============================================================
    doc.add_heading('7. BEST 3 핵심 설정값', level=1)

    cfg_headers = ['항목', 'v14.3 (1위)', 'v13.5 (2위)', 'v14.4 (3위)']
    cfg_rows = [
        ['타임프레임', '5분봉', '5분봉', '30분봉'],
        ['Fast MA', 'EMA(7)', 'EMA(7)', 'EMA(3)'],
        ['Slow MA', 'EMA(100)', 'EMA(100)', 'EMA(200)'],
        ['ADX', '14 >= 30', '14 >= 30', '14 >= 35'],
        ['RSI', '14 [30~58]', '14 [30~58]', '14 [30~65]'],
        ['SL', '-10%', '-7%', '-7%'],
        ['트레일링 활성화', '+8%', '+8%', '+6%'],
        ['트레일링 폭', '-6%', '-6%', '-3%'],
        ['레버리지', '10x', '10x', '10x'],
        ['증거금', '30%', '20%', '25%'],
        ['월간한도', '-25%', '-20%', '-20%'],
        ['연패정지', '3회/288캔들', '3회/288캔들', '없음'],
        ['낙폭축소', '-30%', '-50%', '없음'],
    ]
    add_table(doc, cfg_headers, cfg_rows)

    doc.add_page_break()

    # ============================================================
    # 8. 결론
    # ============================================================
    doc.add_heading('8. 결론 및 권장사항', level=1)

    doc.add_heading('8-1. 핵심 발견', level=2)
    findings = [
        '5분봉 EMA(7/100) 크로스 전략이 v10.1~v15.5까지 일관되게 유효 — 구조적 강건성 확인',
        '보호 메커니즘(ML/CP/DD)이 PF를 4.74 → 15.83~71.60으로 극적 향상 — 핵심 수익 요인',
        'SL -10%는 SL 발동 0회를 달성하지만 FL 27회 발생 — 트레이드오프 존재',
        '30분봉(v14.4)은 MDD 32.3%로 실전 안정성 최고이나 PF 2.47로 5분봉 대비 낮음',
        '21개 기획서 중 10개(48%)가 기획서 주장값과 50% 이상 차이 — 검증의 중요성 확인',
    ]
    for f in findings:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_heading('8-2. 실전 적용 권장', level=2)
    recs = [
        '초보자: v14.4 (30m, MDD 32.3%, FL 0회)로 시작',
        '3개월 검증 후: v13.5 (5m, PF 15.83, FL 1회)로 전환',
        '고위험 감수: v14.3 (5m, PF 71.60, FL 27회) — 격리마진 필수 확인',
        '격리마진 반드시 확인: 교차마진 사용 시 전체 계좌 파산 위험',
        'MDD 50% 초과 시: 마진 절반으로 수동 전환',
        '월 1회: 최근 3개월 성과 점검',
    ]
    for r in recs:
        doc.add_paragraph(r, style='List Bullet')

    doc.add_heading('8-3. 주의사항', level=2)
    warnings = [
        '백테스트 수익률은 실전과 다를 수 있음 (슬리피지, 네트워크 지연 등)',
        '과최적화 리스크: 85,000+ 조합 중 최적 선택. Walk-forward 검증 권장',
        '실전 예상 수익률: 백테스트의 30~50% 수준으로 보수적 추정',
    ]
    for w in warnings:
        doc.add_paragraph(w, style='List Bullet')

    # ============================================================
    # SAVE
    # ============================================================
    out_path = os.path.join(DATA_DIR, 'BTC_선물_자동매매_21버전_검증레포트_v15.3.docx')
    doc.save(out_path)
    print(f'Report saved: {out_path}')
    print(f'File size: {os.path.getsize(out_path):,} bytes')


if __name__ == '__main__':
    main()
