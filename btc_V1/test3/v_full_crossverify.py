#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
==========================================================================
v_full_crossverify.py  --  6-Engine Full Cross-Verification for ALL Strategies
==========================================================================
Tests ALL unique strategies from planning docs v12.3 ~ v32.3
across 6 calculation engines with yearly breakdowns.

Engines:
  1. Wilder Manual     - manual Wilder smoothing
  2. Pandas EWM alpha  - ewm(alpha=1/period, adjust=False)
  3. Pandas EWM span   - ewm(span=period, adjust=False)
  4. SMA-seed Wilder   - seed=SMA(first N), then Wilder
  5. No ADX filter     - remove ADX, keep MA cross + RSI
  6. No RSI filter     - remove RSI, keep MA cross + ADX

Output:
  - cross_verify_full_results.csv
  - 전체_기획서_6엔진_교차검증.docx
"""

import numpy as np
import pandas as pd
import time
import os
import sys
import warnings
from multiprocessing import Pool, cpu_count

warnings.filterwarnings('ignore')

# ============================================================
# CONSTANTS
# ============================================================
INITIAL_CAPITAL = 5000.0
FEE_RATE = 0.0004  # 0.04%
DATA_DIR = os.path.dirname(os.path.abspath(__file__))

# ============================================================
# ALL UNIQUE STRATEGIES extracted from planning docs
# ============================================================
# Format: (name, fast_type, fast_len, slow_type, slow_len, tf,
#          adx_period, adx_threshold, rsi_period, rsi_lo, rsi_hi,
#          delay, sl_pct, trail_act, trail_w, margin_pct, leverage,
#          is_v32, adx_rise_bars, ema_gap_min, monitor_window)

def make_strat(name, ft, fl, st, sl_ma, tf, ap, at, rp, rlo, rhi,
               delay, sl, ta, tw, m, lev,
               is_v32=False, arb=0, egm=0.0, mw=0):
    return (name, ft, fl, st, sl_ma, tf,
            ap, at, rp, rlo, rhi,
            delay, sl, ta, tw, m, lev,
            is_v32, arb, egm, mw)

STRATEGIES = [
    # v12.3: EMA(7)/EMA(100) 5m, ADX(14)>=30, RSI(14) 30-58, SL-9%, Trail+8/-6, M20% L10x
    make_strat("v12.3", "EMA", 7, "EMA", 100, "30m", 14, 30, 14, 30, 58,
               0, 0.09, 0.08, 0.06, 0.20, 10),

    # v13.5: EMA(7)/EMA(100) 5m, ADX(14)>=30, RSI(14) 30-58, SL-7%, Trail+8/-6, M20% L10x
    make_strat("v13.5", "EMA", 7, "EMA", 100, "30m", 14, 30, 14, 30, 58,
               0, 0.07, 0.08, 0.06, 0.20, 10),

    # v14.2F: HMA(7)/EMA(200) 30m, ADX(20)>=25, RSI(14) 25-65, delay 3, SL-7%, Trail+10/-1, M30% L10x
    make_strat("v14.2F", "HMA", 7, "EMA", 200, "30m", 20, 25, 14, 25, 65,
               3, 0.07, 0.10, 0.01, 0.30, 10),

    # v14.4: EMA(3)/EMA(200) 30m, ADX(14)>=35, RSI(14) 30-65, SL-7%, Trail+6/-3, M25% L10x
    make_strat("v14.4", "EMA", 3, "EMA", 200, "30m", 14, 35, 14, 30, 65,
               0, 0.07, 0.06, 0.03, 0.25, 10),

    # v15.2: EMA(3)/EMA(200) 30m, ADX(14)>=35, RSI(14) 30-65, delay 6, SL-5%, Trail+6/-5, M30% L10x
    make_strat("v15.2", "EMA", 3, "EMA", 200, "30m", 14, 35, 14, 30, 65,
               6, 0.05, 0.06, 0.05, 0.30, 10),

    # v15.4: EMA(3)/EMA(200) 30m, ADX(14)>=35, RSI(14) 30-65, SL-7%, Trail+6/-3, M40% L10x
    make_strat("v15.4", "EMA", 3, "EMA", 200, "30m", 14, 35, 14, 30, 65,
               0, 0.07, 0.06, 0.03, 0.40, 10),

    # v15.5: EMA(3)/EMA(200) 30m, ADX(14)>=35, RSI(14) 35-65, SL-7%, Trail+6/-5, M35% L10x
    make_strat("v15.5", "EMA", 3, "EMA", 200, "30m", 14, 35, 14, 35, 65,
               0, 0.07, 0.06, 0.05, 0.35, 10),

    # v15.6: EMA(3)/EMA(200) 30m, ADX(14)>=35, RSI(14) 35-65, SL-7%, Trail+6/-5, M35% L10x
    make_strat("v15.6", "EMA", 3, "EMA", 200, "30m", 14, 35, 14, 35, 65,
               0, 0.07, 0.06, 0.05, 0.35, 10),

    # v16.0: WMA(3)/EMA(200) 30m, ADX(20)>=35, RSI(14) 35-65, SL-8%, Trail+4/-3, M50% L10x
    make_strat("v16.0", "WMA", 3, "EMA", 200, "30m", 20, 35, 14, 35, 65,
               0, 0.08, 0.04, 0.03, 0.50, 10),

    # v16.2: WMA(3)/EMA(200) 30m, ADX(20)>=35, RSI(14) 30-65, SL-7%, Trail+6/-3, M30% L10x (Engine A)
    make_strat("v16.2", "WMA", 3, "EMA", 200, "30m", 20, 35, 14, 30, 65,
               0, 0.07, 0.06, 0.03, 0.30, 10),

    # v16.2F: EMA(3)/SMA(300) 30m, ADX(20)>=45, RSI(14) 30-65, SL-5%, Trail+6/-2.5, M20% L10x
    make_strat("v16.2F", "EMA", 3, "SMA", 300, "30m", 20, 45, 14, 30, 65,
               0, 0.05, 0.06, 0.025, 0.20, 10),

    # v16.4: WMA(3)/EMA(200) 30m, ADX(20)>=35, RSI(14) 35-65, SL-8%, Trail+4/-3, M30% L10x
    make_strat("v16.4", "WMA", 3, "EMA", 200, "30m", 20, 35, 14, 35, 65,
               0, 0.08, 0.04, 0.03, 0.30, 10),

    # v16.5: HMA(21)/EMA(250) 10m, ADX(20)>=35, RSI(14) 40-75, SL-6%, Trail+7/-3, M40% L10x
    make_strat("v16.5", "HMA", 21, "EMA", 250, "10m", 20, 35, 14, 40, 75,
               0, 0.06, 0.07, 0.03, 0.40, 10),

    # v16.6: WMA(3)/EMA(200) 30m, ADX(20)>=35, RSI(14) 30-70, delay 5, SL-8%, Trail+3/-2, M50% L10x
    make_strat("v16.6", "WMA", 3, "EMA", 200, "30m", 20, 35, 14, 30, 70,
               5, 0.08, 0.03, 0.02, 0.50, 10),

    # v17.0: EMA(3)/EMA(200) 30m, ADX(20)>=16, RSI(14) 30-70, SL-8%, Trail+3/-2, M50% L10x
    make_strat("v17.0", "EMA", 3, "EMA", 200, "30m", 20, 16, 14, 30, 70,
               0, 0.08, 0.03, 0.02, 0.50, 10),

    # v22.0: WMA(3)/EMA(200) 30m, ADX(20)>=35, RSI(14) 30-60, SL-8%, Trail+3/-2, M50% L10x
    make_strat("v22.0", "WMA", 3, "EMA", 200, "30m", 20, 35, 14, 30, 60,
               0, 0.08, 0.03, 0.02, 0.50, 10),

    # v22.1: WMA(3)/EMA(200) 30m, ADX(20)>=35, RSI(14) 35-60, SL-8%, Trail+4/-1, M50% L10x
    make_strat("v22.1", "WMA", 3, "EMA", 200, "30m", 20, 35, 14, 35, 60,
               0, 0.08, 0.04, 0.01, 0.50, 10),

    # v22.2: EMA(3)/EMA(200) 30m, ADX(14)>=35, RSI(14) 30-65, SL-7%, Trail+6/-3, M25% L10x
    make_strat("v22.2", "EMA", 3, "EMA", 200, "30m", 14, 35, 14, 30, 65,
               0, 0.07, 0.06, 0.03, 0.25, 10),

    # v22.3: EMA(3)/EMA(250) 30m, ADX(20)>=25, RSI(14) 35-65, SL-7%, Trail+4/-2, M40% L10x
    make_strat("v22.3", "EMA", 3, "EMA", 250, "30m", 20, 25, 14, 35, 65,
               0, 0.07, 0.04, 0.02, 0.40, 10),

    # v22.6: WMA(3)/EMA(200) 30m, ADX(20)>=35, RSI(14) 35-60, SL-7%, Trail+5/-3, M40% L10x
    make_strat("v22.6", "WMA", 3, "EMA", 200, "30m", 20, 35, 14, 35, 60,
               0, 0.07, 0.05, 0.03, 0.40, 10),

    # v22.8: EMA(100)/EMA(600) 30m, ADX(20)>=30, RSI(14) 35-75, SL-8%, Trail+6/-5, M35% L10x
    make_strat("v22.8", "EMA", 100, "EMA", 600, "30m", 20, 30, 14, 35, 75,
               0, 0.08, 0.06, 0.05, 0.35, 10),

    # v23.2: WMA(3)/EMA(200) 10m, ADX(14)>=35, RSI(14) 40-75, delay 5, SL-8%, Trail+15/-1, M60% L20x
    make_strat("v23.2", "WMA", 3, "EMA", 200, "10m", 14, 35, 14, 40, 75,
               5, 0.08, 0.15, 0.01, 0.60, 20),

    # v23.4: EMA(3)/EMA(200) 30m, ADX(14)>=35, RSI(14) 30-65, SL-7%, Trail+6/-3, M30% L10x
    make_strat("v23.4", "EMA", 3, "EMA", 200, "30m", 14, 35, 14, 30, 65,
               0, 0.07, 0.06, 0.03, 0.30, 10),

    # v23.5: EMA(3)/SMA(200) 10m, ADX(14)>=35, RSI(14) 40-75, delay 5, SL-10%, Trail+8/-4, M25% L3x
    make_strat("v23.5", "EMA", 3, "SMA", 200, "10m", 14, 35, 14, 40, 75,
               5, 0.10, 0.08, 0.04, 0.25, 3),

    # v23.5b: HMA(5)/EMA(150) 30m, ADX(20)>=25, RSI(14) 30-65, delay 3, SL-10%, Trail+10/-1, M25% L10x
    make_strat("v23.5b", "HMA", 5, "EMA", 150, "30m", 20, 25, 14, 30, 65,
               3, 0.10, 0.10, 0.01, 0.25, 10),

    # v24.2: EMA(3)/EMA(100) 30m, ADX(20)>=30, RSI(14) 30-70, SL-8%, Trail+6/-5, M70% L10x
    make_strat("v24.2", "EMA", 3, "EMA", 100, "30m", 20, 30, 14, 30, 70,
               0, 0.08, 0.06, 0.05, 0.70, 10),

    # v25.0: EMA(5)/EMA(100) 15m, ADX(14)>=30, RSI(14) 40-60, SL-4%, Trail+5/-3, M30% L10x
    make_strat("v25.0", "EMA", 5, "EMA", 100, "15m", 14, 30, 14, 40, 60,
               0, 0.04, 0.05, 0.03, 0.30, 10),

    # v25.1A: HMA(3)/EMA(250) 15m, ADX(14)>=45, RSI(14) 35-70, delay 3, SL-8%, Trail+3/-2, M40% L5x
    make_strat("v25.1A", "HMA", 3, "EMA", 250, "15m", 14, 45, 14, 35, 70,
               3, 0.08, 0.03, 0.02, 0.40, 5),

    # v25.1C: EMA(5)/EMA(100) 15m, ADX(14)>=30, RSI(14) 40-75, SL-5%, Trail+5/-3, M30% L10x
    make_strat("v25.1C", "EMA", 5, "EMA", 100, "15m", 14, 30, 14, 40, 75,
               0, 0.05, 0.05, 0.03, 0.30, 10),

    # v25.2: EMA(3)/EMA(200) 30m, ADX(20)>=35, RSI(14) 40-75, SL-7%, Trail+7/-3, M40% L10x
    make_strat("v25.2", "EMA", 3, "EMA", 200, "30m", 20, 35, 14, 40, 75,
               0, 0.07, 0.07, 0.03, 0.40, 10),

    # v26.0: EMA(3)/SMA(300) 30m, ADX(14)>=40, RSI(14) 30-70, SL-8%, Trail+4/-3, M50% L10x
    make_strat("v26.0", "EMA", 3, "SMA", 300, "30m", 14, 40, 14, 30, 70,
               0, 0.08, 0.04, 0.03, 0.50, 10),

    # v27: EMA(5)/SMA(50) 30m, ADX(14)>=35, RSI(14) 40-80, SL-12%, Trail+5/-3, M40% L10x
    make_strat("v27", "EMA", 5, "SMA", 50, "30m", 14, 35, 14, 40, 80,
               0, 0.12, 0.05, 0.03, 0.40, 10),

    # v28.0: HMA(14)/VWMA(300) 15m, ADX(20)>=35, RSI(14) 35-70, delay 3, SL-7%, Trail+10/-3, M50% L10x
    make_strat("v28.0", "HMA", 14, "EMA", 300, "15m", 20, 35, 14, 35, 70,
               3, 0.07, 0.10, 0.03, 0.50, 10),

    # v32.2: EMA(100)/EMA(600) 30m, ADX(20)>=30, RSI(10) 40-80, SL-3%, Trail+12/-9, M35% L10x
    # SPECIAL: adx_rise=6, ema_gap=0.2%, monitor=24
    make_strat("v32.2", "EMA", 100, "EMA", 600, "30m", 20, 30, 10, 40, 80,
               0, 0.03, 0.12, 0.09, 0.35, 10,
               True, 6, 0.002, 24),

    # v32.3: EMA(75)/SMA(750) 30m, ADX(20)>=30, RSI(11) 40-80, SL-3%, Trail+12/-9, M35% L10x
    # SPECIAL: adx_rise=6, ema_gap=0.2%, monitor=24
    make_strat("v32.3", "EMA", 75, "SMA", 750, "30m", 20, 30, 11, 40, 80,
               0, 0.03, 0.12, 0.09, 0.35, 10,
               True, 6, 0.002, 24),
]

ENGINE_NAMES = [
    "Wilder Manual",
    "EWM alpha",
    "EWM span",
    "SMA-seed Wilder",
    "No ADX",
    "No RSI",
]

YEARS = [2020, 2021, 2022, 2023, 2024, 2025, 2026]


# ============================================================
# DATA LOADING
# ============================================================
def load_5m_data():
    parts = []
    for i in range(1, 4):
        fp = os.path.join(DATA_DIR, f"btc_usdt_5m_2020_to_now_part{i}.csv")
        df = pd.read_csv(fp, parse_dates=['timestamp'])
        parts.append(df)
    df = pd.concat(parts, ignore_index=True)
    df.drop_duplicates(subset='timestamp', keep='first', inplace=True)
    df.sort_values('timestamp', inplace=True)
    df.set_index('timestamp', inplace=True)
    return df


def resample_ohlcv(df_5m, tf):
    rule_map = {'5m': '5min', '10m': '10min', '15m': '15min', '30m': '30min', '1h': '60min'}
    rule = rule_map.get(tf, tf)
    ohlcv = df_5m.resample(rule).agg({
        'open': 'first', 'high': 'max', 'low': 'min',
        'close': 'last', 'volume': 'sum'
    }).dropna()
    return ohlcv


# ============================================================
# INDICATOR CALCULATIONS
# ============================================================
def calc_ema(series, period):
    return series.ewm(span=period, adjust=False).mean()


def calc_sma(series, period):
    return series.rolling(period).mean()


def calc_wma(series, period):
    weights = np.arange(1, period + 1, dtype=float)
    def _wma(x):
        return np.dot(x, weights) / weights.sum()
    return series.rolling(period).apply(_wma, raw=True)


def calc_hma(series, period):
    half_len = max(int(period / 2), 1)
    sqrt_len = max(int(np.sqrt(period)), 1)
    wma_half = calc_wma(series, half_len)
    wma_full = calc_wma(series, period)
    diff = 2 * wma_half - wma_full
    return calc_wma(diff, sqrt_len)


def calc_ma(series, ma_type, period):
    if ma_type == 'EMA':
        return calc_ema(series, period)
    elif ma_type == 'SMA':
        return calc_sma(series, period)
    elif ma_type == 'WMA':
        return calc_wma(series, period)
    elif ma_type == 'HMA':
        return calc_hma(series, period)
    else:
        raise ValueError(f"Unknown MA type: {ma_type}")


# ============================================================
# ADX CALCULATIONS - 6 ENGINES
# ============================================================
def _true_range(high, low, close):
    n = len(high)
    tr = np.empty(n)
    tr[0] = high[0] - low[0]
    for i in range(1, n):
        tr[i] = max(high[i] - low[i],
                     abs(high[i] - close[i-1]),
                     abs(low[i] - close[i-1]))
    return tr


def _dm_plus_minus(high, low):
    n = len(high)
    dm_p = np.zeros(n)
    dm_m = np.zeros(n)
    for i in range(1, n):
        up = high[i] - high[i-1]
        dn = low[i-1] - low[i]
        if up > dn and up > 0:
            dm_p[i] = up
        if dn > up and dn > 0:
            dm_m[i] = dn
    return dm_p, dm_m


def wilder_smooth_manual(values, period):
    n = len(values)
    result = np.full(n, np.nan)
    seed_vals = values[1:period+1]
    if len(seed_vals) < period:
        return result
    result[period] = np.nansum(seed_vals)
    for i in range(period + 1, n):
        result[i] = result[i-1] - result[i-1] / period + values[i]
    return result


def calc_adx_wilder(high, low, close, period):
    """Engine 1: Wilder Manual."""
    tr = _true_range(high, low, close)
    dm_p, dm_m = _dm_plus_minus(high, low)
    n = len(high)

    atr = wilder_smooth_manual(tr, period)
    sdm_p = wilder_smooth_manual(dm_p, period)
    sdm_m = wilder_smooth_manual(dm_m, period)

    di_p = np.full(n, np.nan)
    di_m = np.full(n, np.nan)
    dx = np.full(n, np.nan)

    for i in range(period, n):
        if not np.isnan(atr[i]) and atr[i] > 0:
            di_p[i] = 100.0 * sdm_p[i] / atr[i]
            di_m[i] = 100.0 * sdm_m[i] / atr[i]
            s = di_p[i] + di_m[i]
            if s > 0:
                dx[i] = 100.0 * abs(di_p[i] - di_m[i]) / s

    adx = np.full(n, np.nan)
    valid_dx = []
    for i in range(period, n):
        if not np.isnan(dx[i]):
            valid_dx.append((i, dx[i]))
            if len(valid_dx) == period:
                break
    if len(valid_dx) == period:
        seed_idx = valid_dx[-1][0]
        adx[seed_idx] = np.mean([v for _, v in valid_dx])
        for i in range(seed_idx + 1, n):
            if not np.isnan(dx[i]) and not np.isnan(adx[i-1]):
                adx[i] = (adx[i-1] * (period - 1) + dx[i]) / period
    return adx


def calc_adx_ewm_alpha(high, low, close, period):
    """Engine 2: Pandas EWM alpha=1/period."""
    h, l, c = pd.Series(high), pd.Series(low), pd.Series(close)
    tr = pd.concat([h-l, (h-c.shift(1)).abs(), (l-c.shift(1)).abs()], axis=1).max(axis=1)
    dm_p = pd.Series(np.where((h-h.shift(1) > l.shift(1)-l) & (h-h.shift(1) > 0), h-h.shift(1), 0.0))
    dm_m = pd.Series(np.where((l.shift(1)-l > h-h.shift(1)) & (l.shift(1)-l > 0), l.shift(1)-l, 0.0))
    alpha = 1.0 / period
    atr = tr.ewm(alpha=alpha, adjust=False).mean()
    di_p = 100 * dm_p.ewm(alpha=alpha, adjust=False).mean() / atr
    di_m = 100 * dm_m.ewm(alpha=alpha, adjust=False).mean() / atr
    dx = 100 * (di_p - di_m).abs() / (di_p + di_m + 1e-10)
    adx = dx.ewm(alpha=alpha, adjust=False).mean()
    return adx.values


def calc_adx_ewm_span(high, low, close, period):
    """Engine 3: Pandas EWM span=period."""
    h, l, c = pd.Series(high), pd.Series(low), pd.Series(close)
    tr = pd.concat([h-l, (h-c.shift(1)).abs(), (l-c.shift(1)).abs()], axis=1).max(axis=1)
    dm_p = pd.Series(np.where((h-h.shift(1) > l.shift(1)-l) & (h-h.shift(1) > 0), h-h.shift(1), 0.0))
    dm_m = pd.Series(np.where((l.shift(1)-l > h-h.shift(1)) & (l.shift(1)-l > 0), l.shift(1)-l, 0.0))
    atr = tr.ewm(span=period, adjust=False).mean()
    di_p = 100 * dm_p.ewm(span=period, adjust=False).mean() / atr
    di_m = 100 * dm_m.ewm(span=period, adjust=False).mean() / atr
    dx = 100 * (di_p - di_m).abs() / (di_p + di_m + 1e-10)
    adx = dx.ewm(span=period, adjust=False).mean()
    return adx.values


def calc_adx_sma_seed(high, low, close, period):
    """Engine 4: SMA-seed Wilder."""
    tr = _true_range(high, low, close)
    dm_p, dm_m = _dm_plus_minus(high, low)
    n = len(high)

    def sma_seed_wilder(vals, per):
        res = np.full(n, np.nan)
        s, cnt = 0.0, 0
        for j in range(1, min(per+1, n)):
            if not np.isnan(vals[j]):
                s += vals[j]
                cnt += 1
        if cnt == per:
            res[per] = s
            for i in range(per+1, n):
                res[i] = res[i-1] - res[i-1]/per + vals[i]
        return res

    atr = sma_seed_wilder(tr, period)
    sdm_p = sma_seed_wilder(dm_p, period)
    sdm_m = sma_seed_wilder(dm_m, period)

    di_p, di_m, dx = np.full(n, np.nan), np.full(n, np.nan), np.full(n, np.nan)
    for i in range(period, n):
        if not np.isnan(atr[i]) and atr[i] > 0:
            di_p[i] = 100.0 * sdm_p[i] / atr[i]
            di_m[i] = 100.0 * sdm_m[i] / atr[i]
            s = di_p[i] + di_m[i]
            if s > 0:
                dx[i] = 100.0 * abs(di_p[i] - di_m[i]) / s

    adx = np.full(n, np.nan)
    valid_dx = []
    for i in range(period, n):
        if not np.isnan(dx[i]):
            valid_dx.append((i, dx[i]))
            if len(valid_dx) == period:
                break
    if len(valid_dx) == period:
        seed_idx = valid_dx[-1][0]
        adx[seed_idx] = np.mean([v for _, v in valid_dx])
        for i in range(seed_idx+1, n):
            if not np.isnan(dx[i]) and not np.isnan(adx[i-1]):
                adx[i] = (adx[i-1] * (period-1) + dx[i]) / period
    return adx


def calc_adx(high, low, close, period, engine_id):
    if engine_id == 1:
        return calc_adx_wilder(high, low, close, period)
    elif engine_id == 2:
        return calc_adx_ewm_alpha(high, low, close, period)
    elif engine_id == 3:
        return calc_adx_ewm_span(high, low, close, period)
    elif engine_id == 4:
        return calc_adx_sma_seed(high, low, close, period)
    else:
        return calc_adx_wilder(high, low, close, period)


# ============================================================
# RSI CALCULATIONS
# ============================================================
def calc_rsi_wilder(close, period):
    n = len(close)
    rsi = np.full(n, np.nan)
    delta = np.diff(close, prepend=close[0])
    gain = np.where(delta > 0, delta, 0.0)
    loss = np.where(delta < 0, -delta, 0.0)
    avg_gain, avg_loss = np.full(n, np.nan), np.full(n, np.nan)
    if period < n:
        avg_gain[period] = np.mean(gain[1:period+1])
        avg_loss[period] = np.mean(loss[1:period+1])
        for i in range(period+1, n):
            avg_gain[i] = (avg_gain[i-1]*(period-1) + gain[i]) / period
            avg_loss[i] = (avg_loss[i-1]*(period-1) + loss[i]) / period
        for i in range(period, n):
            if not np.isnan(avg_gain[i]) and not np.isnan(avg_loss[i]):
                if avg_loss[i] == 0:
                    rsi[i] = 100.0
                else:
                    rsi[i] = 100.0 - 100.0 / (1.0 + avg_gain[i]/avg_loss[i])
    return rsi


def calc_rsi_ewm_alpha(close, period):
    s = pd.Series(close)
    delta = s.diff()
    gain, loss = delta.clip(lower=0), (-delta).clip(lower=0)
    alpha = 1.0 / period
    rs = gain.ewm(alpha=alpha, adjust=False).mean() / (loss.ewm(alpha=alpha, adjust=False).mean() + 1e-10)
    return (100.0 - 100.0 / (1.0 + rs)).values


def calc_rsi_ewm_span(close, period):
    s = pd.Series(close)
    delta = s.diff()
    gain, loss = delta.clip(lower=0), (-delta).clip(lower=0)
    rs = gain.ewm(span=period, adjust=False).mean() / (loss.ewm(span=period, adjust=False).mean() + 1e-10)
    return (100.0 - 100.0 / (1.0 + rs)).values


def calc_rsi(close, period, engine_id):
    if engine_id == 1:
        return calc_rsi_wilder(close, period)
    elif engine_id == 2:
        return calc_rsi_ewm_alpha(close, period)
    elif engine_id == 3:
        return calc_rsi_ewm_span(close, period)
    elif engine_id == 4:
        return calc_rsi_wilder(close, period)
    else:
        return calc_rsi_wilder(close, period)


# ============================================================
# BACKTEST ENGINE (with yearly tracking)
# ============================================================
def run_backtest(ohlcv_df, strat, engine_id):
    (name, fast_type, fast_len, slow_type, slow_len, tf,
     adx_period, adx_th, rsi_period, rsi_lo, rsi_hi,
     delay, sl_pct, trail_act, trail_w, margin_pct, leverage,
     is_v32, adx_rise_bars, ema_gap_min, monitor_window) = strat

    close = ohlcv_df['close'].values
    high = ohlcv_df['high'].values
    low = ohlcv_df['low'].values
    timestamps = ohlcv_df.index
    n = len(close)

    # MA
    fast_ma = calc_ma(ohlcv_df['close'], fast_type, fast_len).values
    slow_ma = calc_ma(ohlcv_df['close'], slow_type, slow_len).values

    # ADX / RSI
    use_adx = (engine_id != 5)
    use_rsi = (engine_id != 6)

    if use_adx:
        adx = calc_adx(high, low, close, adx_period, engine_id)
    else:
        adx = np.full(n, 50.0)

    if use_rsi:
        rsi = calc_rsi(close, rsi_period, engine_id)
    else:
        rsi = np.full(n, 50.0)

    # Warmup
    warmup = max(slow_len, fast_len, 2 * adx_period + 10, 100)
    if is_v32:
        warmup = max(warmup, 600)

    # Yearly tracking
    yearly_pnl = {y: 0.0 for y in YEARS}
    yearly_trades = {y: 0 for y in YEARS}
    yearly_peak = {y: 0.0 for y in YEARS}
    yearly_mdd = {y: 0.0 for y in YEARS}
    yearly_gp = {y: 0.0 for y in YEARS}
    yearly_gl = {y: 0.0 for y in YEARS}

    # State
    balance = INITIAL_CAPITAL
    peak_balance = INITIAL_CAPITAL
    max_dd = 0.0

    pos = 0
    entry_price = 0.0
    pos_size_usd = 0.0
    trail_high = 0.0
    trail_low = 999999.0
    trail_active = False

    trades = 0
    wins = 0
    sl_hits = 0
    gross_profit = 0.0
    gross_loss = 0.0

    cross_signal = 0
    cross_bar = -9999
    last_close_dir = 0

    def get_year(idx):
        try:
            return timestamps[idx].year
        except:
            return 2020

    def update_yearly_mdd(yr, bal, peak_yr):
        if bal > peak_yr:
            peak_yr = bal
        dd = (peak_yr - bal) / peak_yr * 100 if peak_yr > 0 else 0
        if dd > yearly_mdd[yr]:
            yearly_mdd[yr] = dd
        return peak_yr

    # Set yearly starting peaks
    for y in YEARS:
        yearly_peak[y] = INITIAL_CAPITAL

    for i in range(warmup, n):
        if np.isnan(fast_ma[i]) or np.isnan(slow_ma[i]):
            continue

        yr = get_year(i)
        if yr not in yearly_pnl:
            yr = max(y for y in YEARS if y <= yr) if yr > 2026 else min(y for y in YEARS if y >= yr)

        # Detect cross
        if i >= 1 and not np.isnan(fast_ma[i-1]) and not np.isnan(slow_ma[i-1]):
            if fast_ma[i-1] <= slow_ma[i-1] and fast_ma[i] > slow_ma[i]:
                cross_signal = 1
                cross_bar = i
            elif fast_ma[i-1] >= slow_ma[i-1] and fast_ma[i] < slow_ma[i]:
                cross_signal = -1
                cross_bar = i

        # Signal with delay
        cur_signal = 0
        if cross_signal != 0:
            bars_since = i - cross_bar
            if delay == 0:
                if bars_since == 0:
                    cur_signal = cross_signal
            else:
                if bars_since == delay:
                    cur_signal = cross_signal

            # Monitor window for v32
            if monitor_window > 0 and bars_since > monitor_window:
                cross_signal = 0

        # Sustained signal when no position
        if cur_signal == 0 and pos == 0:
            if fast_ma[i] > slow_ma[i]:
                cur_signal = 1
            elif fast_ma[i] < slow_ma[i]:
                cur_signal = -1

        # Filters
        adx_ok = True
        rsi_ok = True

        if use_adx and not np.isnan(adx[i]):
            adx_ok = adx[i] >= adx_th
            if is_v32 and adx_rise_bars > 0 and i >= adx_rise_bars:
                if not np.isnan(adx[i - adx_rise_bars]):
                    adx_ok = adx_ok and (adx[i] > adx[i - adx_rise_bars])
        elif use_adx and np.isnan(adx[i]):
            adx_ok = False

        if use_rsi and not np.isnan(rsi[i]):
            rsi_ok = rsi_lo <= rsi[i] <= rsi_hi
        elif use_rsi and np.isnan(rsi[i]):
            rsi_ok = False

        # EMA gap for v32
        if is_v32 and ema_gap_min > 0 and slow_ma[i] != 0:
            gap = abs(fast_ma[i] - slow_ma[i]) / abs(slow_ma[i])
            if gap < ema_gap_min:
                adx_ok = False

        # Skip same direction re-entry for v32
        if is_v32 and cur_signal != 0 and cur_signal == last_close_dir and pos == 0:
            cur_signal = 0

        # --- Position management ---
        if pos != 0:
            if pos == 1:
                roi = (close[i] - entry_price) / entry_price * leverage
                if close[i] > trail_high:
                    trail_high = close[i]
                sl_roi = (low[i] - entry_price) / entry_price * leverage
            else:
                roi = (entry_price - close[i]) / entry_price * leverage
                if close[i] < trail_low:
                    trail_low = close[i]
                sl_roi = (entry_price - high[i]) / entry_price * leverage

            # SL check
            if not trail_active and sl_roi <= -sl_pct:
                pnl = -sl_pct * pos_size_usd
                balance += pnl
                gross_loss += abs(pnl)
                trades += 1
                sl_hits += 1
                yearly_pnl[yr] += pnl
                yearly_trades[yr] += 1
                yearly_gl[yr] += abs(pnl)
                yearly_peak[yr] = update_yearly_mdd(yr, balance, yearly_peak[yr])
                last_close_dir = pos
                pos = 0
                trail_active = False
                if balance <= 0:
                    balance = 0
                    break
                # MDD
                if balance > peak_balance:
                    peak_balance = balance
                dd = (peak_balance - balance) / peak_balance * 100
                if dd > max_dd:
                    max_dd = dd
                continue

            # Trail activation
            if roi >= trail_act:
                trail_active = True

            # TSL check
            if trail_active:
                if pos == 1:
                    trail_roi = (close[i] - trail_high) / trail_high * leverage
                else:
                    trail_roi = (trail_low - close[i]) / trail_low * leverage

                if trail_roi <= -trail_w:
                    pnl = roi * pos_size_usd
                    fee = pos_size_usd * FEE_RATE
                    balance += pnl - fee
                    if pnl > 0:
                        gross_profit += pnl
                        wins += 1
                        yearly_gp[yr] += pnl
                    else:
                        gross_loss += abs(pnl)
                        yearly_gl[yr] += abs(pnl)
                    trades += 1
                    yearly_pnl[yr] += pnl - fee
                    yearly_trades[yr] += 1
                    yearly_peak[yr] = update_yearly_mdd(yr, balance, yearly_peak[yr])
                    last_close_dir = pos
                    pos = 0
                    trail_active = False
                    if balance > peak_balance:
                        peak_balance = balance
                    dd = (peak_balance - balance) / peak_balance * 100
                    if dd > max_dd:
                        max_dd = dd
                    continue

            # REV signal
            if cur_signal != 0 and cur_signal != pos and adx_ok and rsi_ok:
                pnl = roi * pos_size_usd
                fee = pos_size_usd * FEE_RATE * 2
                balance += pnl - fee
                if pnl > 0:
                    gross_profit += pnl
                    wins += 1
                    yearly_gp[yr] += pnl
                else:
                    gross_loss += abs(pnl)
                    yearly_gl[yr] += abs(pnl)
                trades += 1
                yearly_pnl[yr] += pnl - fee
                yearly_trades[yr] += 1
                yearly_peak[yr] = update_yearly_mdd(yr, balance, yearly_peak[yr])
                last_close_dir = pos
                pos = 0
                trail_active = False

                if balance <= 0:
                    balance = 0
                    break

                # Open reverse
                pos = cur_signal
                pos_size_usd = balance * margin_pct
                entry_price = close[i]
                trail_high = close[i]
                trail_low = close[i]
                trail_active = False
                fee_entry = pos_size_usd * FEE_RATE
                balance -= fee_entry

                if balance > peak_balance:
                    peak_balance = balance
                dd = (peak_balance - balance) / peak_balance * 100
                if dd > max_dd:
                    max_dd = dd
                continue

        # --- Entry ---
        if pos == 0 and cur_signal != 0 and adx_ok and rsi_ok:
            pos = cur_signal
            pos_size_usd = balance * margin_pct
            entry_price = close[i]
            trail_high = close[i]
            trail_low = close[i]
            trail_active = False
            fee_entry = pos_size_usd * FEE_RATE
            balance -= fee_entry

        # MDD tracking
        if balance > peak_balance:
            peak_balance = balance
        dd = (peak_balance - balance) / peak_balance * 100
        if dd > max_dd:
            max_dd = dd

    # Close remaining
    if pos != 0:
        yr = get_year(n-1)
        if yr not in yearly_pnl:
            yr = 2026
        if pos == 1:
            roi = (close[-1] - entry_price) / entry_price * leverage
        else:
            roi = (entry_price - close[-1]) / entry_price * leverage
        pnl = roi * pos_size_usd
        fee = pos_size_usd * FEE_RATE
        balance += pnl - fee
        if pnl > 0:
            gross_profit += pnl
            wins += 1
            yearly_gp[yr] += pnl
        else:
            gross_loss += abs(pnl)
            yearly_gl[yr] += abs(pnl)
        trades += 1
        yearly_pnl[yr] += pnl - fee
        yearly_trades[yr] += 1

    if balance > peak_balance:
        peak_balance = balance
    dd = (peak_balance - balance) / peak_balance * 100
    if dd > max_dd:
        max_dd = dd

    pf = gross_profit / gross_loss if gross_loss > 0 else (999.0 if gross_profit > 0 else 0.0)
    ret_pct = (balance - INITIAL_CAPITAL) / INITIAL_CAPITAL * 100

    # Yearly PF
    yearly_pf = {}
    for y in YEARS:
        if yearly_gl[y] > 0:
            yearly_pf[y] = round(yearly_gp[y] / yearly_gl[y], 2)
        elif yearly_gp[y] > 0:
            yearly_pf[y] = 999.0
        else:
            yearly_pf[y] = 0.0

    return {
        'strategy': name,
        'engine_id': engine_id,
        'engine_name': ENGINE_NAMES[engine_id - 1],
        'final_balance': round(balance, 2),
        'return_pct': round(ret_pct, 2),
        'pnl_usd': round(balance - INITIAL_CAPITAL, 2),
        'pf': round(pf, 2),
        'mdd_pct': round(max_dd, 2),
        'trades': trades,
        'wins': wins,
        'sl_hits': sl_hits,
        **{f'pnl_{y}': round(yearly_pnl[y], 2) for y in YEARS},
        **{f'trades_{y}': yearly_trades[y] for y in YEARS},
        **{f'mdd_{y}': round(yearly_mdd[y], 2) for y in YEARS},
        **{f'pf_{y}': yearly_pf[y] for y in YEARS},
    }


# ============================================================
# MULTIPROCESSING
# ============================================================
_SHARED_DATA = {}


def _pool_init(shared_dict):
    global _SHARED_DATA
    _SHARED_DATA = shared_dict


def worker_task(args):
    strat, engine_id = args
    tf = strat[5]
    ohlcv_df = _SHARED_DATA[tf]
    return run_backtest(ohlcv_df, strat, engine_id)


# ============================================================
# DOCX GENERATION
# ============================================================
def generate_docx(df_all, output_path):
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        import subprocess
        subprocess.run([sys.executable, '-m', 'pip', 'install', 'python-docx'], check=True)
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'
    font.size = Pt(8)

    doc.add_heading('BTC/USDT 전체 기획서 6엔진 교차검증 리포트', level=0)
    doc.add_paragraph(f'생성일: {pd.Timestamp.now().strftime("%Y-%m-%d %H:%M")}')
    doc.add_paragraph(f'전략 수: {len(STRATEGIES)}개 | 엔진 수: 6개 | 총 백테스트: {len(STRATEGIES)*6}회')
    doc.add_paragraph(f'데이터: BTC/USDT 5분봉 (2020-01 ~ 2026-03) -> 10m, 15m, 30m 리샘플링')
    doc.add_paragraph('')

    # --- Aggregate per strategy ---
    strat_summary = []
    for sname in [s[0] for s in STRATEGIES]:
        sdf = df_all[df_all['strategy'] == sname]
        avg_ret = sdf['return_pct'].mean()
        avg_pf = sdf['pf'].mean()
        avg_mdd = sdf['mdd_pct'].mean()
        avg_pnl = sdf['pnl_usd'].mean()
        profitable_engines = (sdf['return_pct'] > 0).sum()
        total_trades = sdf['trades'].mean()

        # Yearly data (averaged across engines)
        yr_trades = {y: sdf[f'trades_{y}'].mean() for y in YEARS}
        yr_mdd = {y: sdf[f'mdd_{y}'].mean() for y in YEARS}
        yr_pf = {y: sdf[f'pf_{y}'].mean() for y in YEARS}

        # Engine consistency reason
        engine_results = []
        for _, row in sdf.iterrows():
            engine_results.append((row['engine_name'], row['return_pct']))
        consistent = all(r[1] > 0 for r in engine_results)
        if consistent:
            reason = "6/6 전체 수익"
        elif profitable_engines >= 4:
            reason = f"{profitable_engines}/6 수익"
        else:
            losing = [(r[0], r[1]) for r in engine_results if r[1] <= 0]
            reason = f"{profitable_engines}/6 수익, " + ", ".join(f"{r[0]}:{r[1]:.0f}%" for r in losing[:2])

        strat_summary.append({
            'name': sname,
            'avg_ret': avg_ret,
            'avg_pnl': avg_pnl,
            'avg_pf': avg_pf,
            'avg_mdd': avg_mdd,
            'profitable_engines': profitable_engines,
            'total_trades': total_trades,
            'yr_trades': yr_trades,
            'yr_mdd': yr_mdd,
            'yr_pf': yr_pf,
            'reason': reason,
        })

    # --- TABLE 1: Return BEST 10 ---
    # Profitable in 4+ engines, sorted by avg return
    ret_candidates = [s for s in strat_summary if s['profitable_engines'] >= 4]
    ret_candidates.sort(key=lambda x: x['avg_ret'], reverse=True)
    ret_best10 = ret_candidates[:10]

    # --- TABLE 2: Stability BEST 10 ---
    stab_candidates = [s for s in strat_summary if s['profitable_engines'] >= 4]
    stab_candidates.sort(key=lambda x: (x['avg_mdd'], -x['avg_pf']))
    stab_best10 = stab_candidates[:10]

    # --- TABLE 3: Discard BEST 10 ---
    discard_candidates = [s for s in strat_summary if s['profitable_engines'] < 4 or s['avg_ret'] < 0]
    if len(discard_candidates) < 10:
        remaining = [s for s in strat_summary if s not in discard_candidates]
        remaining.sort(key=lambda x: x['avg_ret'])
        discard_candidates.extend(remaining[:10-len(discard_candidates)])
    discard_candidates.sort(key=lambda x: x['avg_ret'])
    discard_best10 = discard_candidates[:10]

    def add_table(doc, title, data_list, color_label):
        doc.add_heading(title, level=1)

        cols = ['순위', '파일명', '손익률', '손익금액',
                '2020', '2021', '2022', '2023', '2024', '2025', '2026',
                '총거래',
                'MDD_20', 'MDD_21', 'MDD_22', 'MDD_23', 'MDD_24', 'MDD_25', 'MDD_26',
                'PF_20', 'PF_21', 'PF_22', 'PF_23', 'PF_24', 'PF_25', 'PF_26',
                '6엔진 일치', '비고']

        table = doc.add_table(rows=1, cols=len(cols))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        # Header
        hdr = table.rows[0].cells
        for j, col in enumerate(cols):
            hdr[j].text = col
            for p in hdr[j].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(6)
                    run.font.bold = True

        for rank, item in enumerate(data_list, 1):
            row = table.add_row().cells
            vals = [
                str(rank),
                item['name'],
                f"{item['avg_ret']:.1f}%",
                f"${item['avg_pnl']:,.0f}",
            ]
            # Yearly trades
            for y in YEARS:
                vals.append(f"{item['yr_trades'][y]:.0f}")
            vals.append(f"{item['total_trades']:.0f}")
            # Yearly MDD
            for y in YEARS:
                vals.append(f"{item['yr_mdd'][y]:.1f}%")
            # Yearly PF
            for y in YEARS:
                pf_v = item['yr_pf'][y]
                vals.append(f"{pf_v:.1f}" if pf_v < 100 else "INF")
            vals.append(item['reason'])
            # Remarks
            pe = item['profitable_engines']
            if pe == 6:
                vals.append("전 엔진 수익")
            elif pe >= 4:
                vals.append(f"{pe}개 엔진 수익")
            else:
                vals.append(f"수익 {pe}개만")

            for j, v in enumerate(vals):
                row[j].text = v
                for p in row[j].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.size = Pt(6)

    add_table(doc, 'Table 1: 수익률 BEST 10', ret_best10, 'green')
    add_table(doc, 'Table 2: 안정형 BEST 10 (최저 MDD)', stab_best10, 'blue')
    add_table(doc, 'Table 3: 폐기형 BEST 10', discard_best10, 'red')

    # --- Engine comparison matrix ---
    doc.add_heading('6엔진 수익률 매트릭스', level=1)
    matrix_cols = ['전략'] + ENGINE_NAMES + ['AVG', 'STD']
    mt = doc.add_table(rows=1, cols=len(matrix_cols))
    mt.style = 'Table Grid'
    hdr = mt.rows[0].cells
    for j, c in enumerate(matrix_cols):
        hdr[j].text = c
        for p in hdr[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(6)
                run.font.bold = True

    for sname in [s[0] for s in STRATEGIES]:
        sdf = df_all[df_all['strategy'] == sname]
        row = mt.add_row().cells
        row[0].text = sname
        for eid in range(1, 7):
            val = sdf[sdf['engine_id'] == eid]['return_pct'].values
            row[eid].text = f"{val[0]:.0f}%" if len(val) > 0 else "N/A"
        avg = sdf['return_pct'].mean()
        std = sdf['return_pct'].std()
        row[7].text = f"{avg:.0f}%"
        row[8].text = f"{std:.0f}"
        for j in range(len(matrix_cols)):
            for p in row[j].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(6)

    doc.save(output_path)
    print(f"  DOCX saved: {output_path}")


# ============================================================
# MAIN
# ============================================================
def main():
    global _SHARED_DATA

    print("=" * 80)
    print("  FULL 6-ENGINE CROSS-VERIFICATION")
    print(f"  {len(STRATEGIES)} Strategies x 6 Engines = {len(STRATEGIES)*6} Backtests")
    print("=" * 80)
    print()

    t_start = time.time()

    # Load data
    print("[1/5] Loading 5m data...")
    df_5m = load_5m_data()
    print(f"      Loaded {len(df_5m):,} rows")

    # Resample
    print("[2/5] Resampling...")
    needed_tfs = set(s[5] for s in STRATEGIES)
    print(f"      TFs needed: {sorted(needed_tfs)}")

    for tf in needed_tfs:
        _SHARED_DATA[tf] = resample_ohlcv(df_5m, tf)
        print(f"      {tf}: {len(_SHARED_DATA[tf]):,} bars")
    del df_5m

    # Build tasks
    print(f"[3/5] Running {len(STRATEGIES)*6} backtests...")
    tasks = []
    for strat in STRATEGIES:
        for eng_id in range(1, 7):
            tasks.append((strat, eng_id))

    n_workers = min(cpu_count(), 8)
    print(f"      Workers: {n_workers}")

    results = []
    with Pool(processes=n_workers, initializer=_pool_init, initargs=(_SHARED_DATA,)) as pool:
        for i, res in enumerate(pool.imap_unordered(worker_task, tasks)):
            results.append(res)
            done = i + 1
            if done % 30 == 0 or done == len(tasks):
                elapsed = time.time() - t_start
                pct = done / len(tasks) * 100
                print(f"      {done}/{len(tasks)} ({pct:.0f}%) - {elapsed:.1f}s")

    # Results
    print("[4/5] Generating CSV...")
    df_all = pd.DataFrame(results)
    csv_path = os.path.join(DATA_DIR, "cross_verify_full_results.csv")
    df_all.to_csv(csv_path, index=False, encoding='utf-8-sig')
    print(f"      CSV saved: {csv_path}")

    # Summary printout
    print()
    print("=" * 120)
    print("  RETURN (%) MATRIX")
    print("=" * 120)
    for sname in [s[0] for s in STRATEGIES]:
        sdf = df_all[df_all['strategy'] == sname]
        vals = []
        for eid in range(1, 7):
            r = sdf[sdf['engine_id'] == eid]['return_pct'].values
            vals.append(f"{r[0]:>10,.0f}%" if len(r) > 0 else f"{'N/A':>10}")
        avg = sdf['return_pct'].mean()
        std = sdf['return_pct'].std()
        print(f"  {sname:>10s} | {'|'.join(vals)} | AVG {avg:>10,.0f}% | STD {std:>8,.0f}")

    # DOCX
    print()
    print("[5/5] Generating DOCX...")
    docx_path = os.path.join(DATA_DIR, "전체_기획서_6엔진_교차검증.docx")
    generate_docx(df_all, docx_path)

    elapsed = time.time() - t_start
    print()
    print(f"  COMPLETE in {elapsed:.1f}s")
    print(f"  CSV:  {csv_path}")
    print(f"  DOCX: {docx_path}")


if __name__ == '__main__':
    main()
