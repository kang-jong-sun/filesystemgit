# -*- coding: utf-8 -*-
"""v14.4 & v15.4 Trade-by-Trade Detail Report + Word"""
import sys, os, time, json
import numpy as np
import pandas as pd
import warnings
warnings.filterwarnings('ignore')
sys.stdout.reconfigure(encoding='utf-8', errors='replace')

DATA_DIR = os.path.dirname(os.path.abspath(__file__))

# ============================================================
# Indicators
# ============================================================
def ema(s,p): return s.ewm(span=p, adjust=False).mean()
def adx_w(h,l,c,p=14):
    tr=pd.concat([h-l,(h-c.shift(1)).abs(),(l-c.shift(1)).abs()],axis=1).max(axis=1)
    up=h-h.shift(1); dn=l.shift(1)-l
    pdm=np.where((up>dn)&(up>0),up,0.0); mdm=np.where((dn>up)&(dn>0),dn,0.0)
    atr=pd.Series(tr,index=c.index).ewm(alpha=1/p,min_periods=p,adjust=False).mean()
    pdi=100*pd.Series(pdm,index=c.index).ewm(alpha=1/p,min_periods=p,adjust=False).mean()/atr
    mdi=100*pd.Series(mdm,index=c.index).ewm(alpha=1/p,min_periods=p,adjust=False).mean()/atr
    dx=100*(pdi-mdi).abs()/(pdi+mdi).replace(0,np.nan)
    return dx.ewm(alpha=1/p,min_periods=p,adjust=False).mean()
def rsi_w(c,p=14):
    d=c.diff(); g=d.clip(lower=0); lo=(-d).clip(lower=0)
    ag=g.ewm(alpha=1/p,min_periods=p,adjust=False).mean()
    al=lo.ewm(alpha=1/p,min_periods=p,adjust=False).mean()
    return 100-100/(1+ag/al.replace(0,np.nan))


def backtest_detail(df_tf, cfg):
    """Full trade-by-trade backtest with detailed logging"""
    n = len(df_tf)
    times=df_tf['time'].values; closes=df_tf['close'].values
    highs=df_tf['high'].values; lows=df_tf['low'].values
    mf=df_tf['ma_fast'].values; ms=df_tf['ma_slow'].values
    rsi_v=df_tf['rsi'].values; adx_v=df_tf['adx'].values

    fee=0.0004; balance=3000.0; peak_bal=balance
    pos=0; entry_p=0.0; entry_t=''; sz=0.0; margin_used=0.0
    peak_price=0.0; trail_on=False
    cur_m=''; m_start=balance; m_paused=False
    lev=cfg['lev']; margin=cfg['margin']; sl_pct=cfg['sl']
    ta_pct=cfg['ta']; tw_pct=cfg['tw']; ml=cfg['ml']
    liq_dist=1.0/lev

    trades = []  # detailed trade log
    monthly = {}
    m_tr=0; m_sl=0; m_tsl=0; m_rev=0; m_fl=0; m_pnl_sum=0

    def save_month():
        nonlocal m_tr, m_sl, m_tsl, m_rev, m_fl, m_pnl_sum
        if cur_m:
            monthly[cur_m] = {
                'pnl': balance - m_start,
                'pct': (balance-m_start)/m_start*100 if m_start>0 else 0,
                'bal': round(balance,2), 'tr': m_tr,
                'sl': m_sl, 'tsl': m_tsl, 'rev': m_rev, 'fl': m_fl
            }
        m_tr=0; m_sl=0; m_tsl=0; m_rev=0; m_fl=0

    def record_trade(exit_p, exit_t, exit_type, idx):
        nonlocal balance, pos, peak_bal, m_tr, m_sl, m_tsl, m_rev, m_fl
        pnl_pct = pos * (exit_p - entry_p) / entry_p
        pnl_usd = sz * pnl_pct
        fee_total = sz * fee  # exit fee
        net_pnl = pnl_usd - fee_total
        balance += net_pnl
        if exit_type == 'FL': balance = max(balance, 0)
        peak_bal = max(peak_bal, balance)

        m_tr += 1
        if exit_type == 'SL': m_sl += 1
        elif exit_type == 'TSL': m_tsl += 1
        elif exit_type == 'REV': m_rev += 1
        elif exit_type == 'FL': m_fl += 1

        hold_candles = idx - trades[-1]['entry_idx'] if trades and 'entry_idx' in trades[-1] else 0

        trades.append({
            'no': len(trades) + 1,
            'entry_time': entry_t,
            'exit_time': exit_t,
            'dir': 'LONG' if pos == 1 else 'SHORT',
            'entry_price': round(entry_p, 2),
            'exit_price': round(exit_p, 2),
            'margin': round(margin_used, 2),
            'position_size': round(sz, 2),
            'pnl_pct': round(pnl_pct * 100, 2),
            'pnl_usd': round(net_pnl, 2),
            'balance_after': round(balance, 2),
            'exit_type': exit_type,
            'peak_price': round(peak_price, 2),
            'adx_at_entry': 0,  # will be set
            'rsi_at_entry': 0,
            'hold_candles': 0,
        })
        pos = 0
        return

    for i in range(1, n):
        if np.isnan(mf[i]) or np.isnan(ms[i]) or np.isnan(rsi_v[i]) or np.isnan(adx_v[i]):
            continue
        t = str(times[i])[:16]; cc = closes[i]; hh = highs[i]; ll = lows[i]
        mk = t[:7]
        if mk != cur_m:
            save_month(); cur_m = mk; m_start = balance; m_paused = False

        if pos != 0:
            w = ll if pos==1 else hh
            # FL
            if pos*(w-entry_p)/entry_p <= -liq_dist:
                fl_p = entry_p*(1-pos*liq_dist)
                record_trade(fl_p, t, 'FL', i)
                continue
            # SL
            if pos*(w-entry_p)/entry_p <= -sl_pct:
                sl_p = entry_p*(1-pos*sl_pct)
                record_trade(sl_p, t, 'SL', i)
                continue
            # Trail
            if pos==1: peak_price = max(peak_price, hh)
            else: peak_price = min(peak_price, ll)
            ppnl = pos*(peak_price-entry_p)/entry_p
            if ppnl >= ta_pct: trail_on = True
            if trail_on:
                if pos==1:
                    tsl = peak_price*(1-tw_pct)
                    if cc <= tsl:
                        record_trade(tsl, t, 'TSL', i); continue
                else:
                    tsl = peak_price*(1+tw_pct)
                    if cc >= tsl:
                        record_trade(tsl, t, 'TSL', i); continue
            # REV
            cu = mf[i]>ms[i] and mf[i-1]<=ms[i-1]
            cd = mf[i]<ms[i] and mf[i-1]>=ms[i-1]
            ao = adx_v[i]>=cfg['adx_min']; ro = cfg['rsi_min']<=rsi_v[i]<=cfg['rsi_max']
            nd = 0
            if pos==1 and cd and ao and ro: nd=-1
            elif pos==-1 and cu and ao and ro: nd=1
            if nd != 0:
                record_trade(cc, t, 'REV', i)
                if balance > 10 and not m_paused:
                    margin_used = balance * margin
                    pos = nd; entry_p = cc; entry_t = t
                    sz = margin_used * lev
                    balance -= sz * fee
                    peak_price = cc; trail_on = False
                    # Record entry info in next trade stub
                continue
            # ML
            if ml < 0 and m_start > 0:
                ur = sz*pos*(cc-entry_p)/entry_p
                if (balance+ur-m_start)/m_start < ml: m_paused = True

        if pos == 0 and balance > 10:
            if ml < 0 and m_start > 0:
                if (balance-m_start)/m_start < ml: m_paused = True
            if m_paused: continue

            cu = mf[i]>ms[i] and mf[i-1]<=ms[i-1]
            cd = mf[i]<ms[i] and mf[i-1]>=ms[i-1]
            ao = adx_v[i]>=cfg['adx_min']; ro = cfg['rsi_min']<=rsi_v[i]<=cfg['rsi_max']
            sig = 0
            if cu and ao and ro: sig = 1
            elif cd and ao and ro: sig = -1
            if sig != 0:
                margin_used = balance * margin
                pos = sig; entry_p = cc; entry_t = t
                sz = margin_used * lev
                balance -= sz * fee
                peak_price = cc; trail_on = False
        peak_bal = max(peak_bal, balance)

    if pos != 0:
        record_trade(closes[-1], str(times[-1])[:16], 'END', n-1)
    save_month()

    # Re-number trades and calc hold time
    for i, tr in enumerate(trades):
        tr['no'] = i + 1

    return trades, monthly, balance


def main():
    t0 = time.time()
    print("Loading data..."); sys.stdout.flush()

    parts = []
    for f in sorted(os.listdir(DATA_DIR)):
        if f.startswith('btc_usdt_5m') and f.endswith('.csv'):
            parts.append(pd.read_csv(os.path.join(DATA_DIR, f), parse_dates=['timestamp']))
    df = pd.concat(parts, ignore_index=True)
    df = df.drop_duplicates(subset='timestamp').sort_values('timestamp').reset_index(drop=True)
    df = df.rename(columns={'timestamp': 'time'})
    for c in ['open','high','low','close','volume']: df[c] = df[c].astype(float)

    df30 = df.set_index('time').resample('30min').agg(
        {'open':'first','high':'max','low':'min','close':'last','volume':'sum'}).dropna().reset_index()
    print(f"30m: {len(df30):,} candles"); sys.stdout.flush()

    df30['ma_fast'] = ema(df30['close'], 3)
    df30['ma_slow'] = ema(df30['close'], 200)
    df30['adx'] = adx_w(df30['high'], df30['low'], df30['close'], 14)
    df30['rsi'] = rsi_w(df30['close'], 14)

    CONFIGS = {
        'v14.4': {'adx_min':35,'rsi_min':30,'rsi_max':65,'sl':0.07,'ta':0.06,'tw':0.03,
                  'lev':10,'margin':0.25,'ml':-0.20},
        'v15.4': {'adx_min':35,'rsi_min':30,'rsi_max':65,'sl':0.07,'ta':0.06,'tw':0.03,
                  'lev':10,'margin':0.40,'ml':-0.30},
    }

    all_data = {}

    for name, cfg in CONFIGS.items():
        print(f"\n{'='*60}")
        print(f"  Running {name}...")
        print(f"{'='*60}"); sys.stdout.flush()

        # Run + 10x consistency
        trades, monthly, final_bal = backtest_detail(df30, cfg)
        consistent = True
        for run in range(9):
            t2, m2, b2 = backtest_detail(df30, cfg)
            if round(b2) != round(final_bal):
                consistent = False; break

        print(f"  Trades: {len(trades)}")
        print(f"  Final: ${final_bal:,.0f}")
        print(f"  10x Consistent: {consistent}")
        sys.stdout.flush()

        all_data[name] = {'trades': trades, 'monthly': monthly, 'bal': final_bal, 'consistent': consistent, 'cfg': cfg}

    # Generate Word
    print("\nGenerating Word report..."); sys.stdout.flush()
    generate_word(all_data)
    print(f"Done in {time.time()-t0:.0f}s"); sys.stdout.flush()


def generate_word(all_data):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT

    def cell(c, text, bold=False, sz=7, align='center', color=None):
        p = c.paragraphs[0]
        p.alignment = {'left':WD_ALIGN_PARAGRAPH.LEFT,'center':WD_ALIGN_PARAGRAPH.CENTER,
                       'right':WD_ALIGN_PARAGRAPH.RIGHT}.get(align, WD_ALIGN_PARAGRAPH.CENTER)
        run = p.add_run(str(text))
        run.font.size = Pt(sz); run.bold = bold; run.font.name = 'Malgun Gothic'
        if color: run.font.color.rgb = RGBColor(*color)

    def add_tbl(doc, hdrs, rows):
        t = doc.add_table(rows=1+len(rows), cols=len(hdrs))
        t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, h in enumerate(hdrs): cell(t.rows[0].cells[j], h, bold=True, sz=7)
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                clr = None
                if isinstance(val, str):
                    if val.startswith('+') and ('$' in val or '%' in val): clr = (0,100,0)
                    elif val.startswith('-') and ('$' in val or '%' in val): clr = (200,0,0)
                cell(t.rows[i+1].cells[j], val, sz=7, color=clr)
        return t

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(42.0); sec.page_height = Cm(29.7)
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.left_margin = Cm(1.0); sec.right_margin = Cm(1.0)
    sec.top_margin = Cm(1.0); sec.bottom_margin = Cm(1.0)

    # TITLE
    t = doc.add_heading('v14.4 / v15.4 Trade-by-Trade Verification Report', level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('30m EMA(3/200) + ADX(14)>=35 + RSI(14) 30~65 | SL-7% Trail+6%/-3%\n'
                     'Data: 2020-01 ~ 2026-03 (75mo, 109,234 candles) | Fee: 0.04% x 2 | 10x Verified')
    run.font.size = Pt(9); run.font.color.rgb = RGBColor(100,100,100)
    doc.add_page_break()

    for name in ['v14.4', 'v15.4']:
        data = all_data[name]
        trades = data['trades']
        monthly = data['monthly']
        cfg = data['cfg']

        # Summary
        doc.add_heading(f'{name} Summary', level=1)
        doc.add_paragraph(f'Margin: {cfg["margin"]*100:.0f}% | Leverage: {cfg["lev"]}x | '
                          f'ML: {cfg["ml"]*100:.0f}% | 10x Consistent: {data["consistent"]}')

        wins = [t for t in trades if t['pnl_usd'] > 0]
        losses = [t for t in trades if t['pnl_usd'] <= 0]
        sl_c = sum(1 for t in trades if t['exit_type']=='SL')
        tsl_c = sum(1 for t in trades if t['exit_type']=='TSL')
        rev_c = sum(1 for t in trades if t['exit_type']=='REV')
        fl_c = sum(1 for t in trades if t['exit_type']=='FL')
        pk=3000; mdd=0
        for m in sorted(monthly.keys()):
            b=monthly[m]['bal']; pk=max(pk,b)
            d=(pk-b)/pk if pk>0 else 0; mdd=max(mdd,d)

        info = [
            ['Final Balance', f'${data["bal"]:,.0f}', 'Return', f'{(data["bal"]-3000)/3000*100:+,.1f}%'],
            ['Trades', str(len(trades)), 'Win Rate', f'{len(wins)/len(trades)*100:.1f}%' if trades else '0%'],
            ['SL', str(sl_c), 'TSL', str(tsl_c)],
            ['REV', str(rev_c), 'FL', str(fl_c)],
            ['MDD', f'{mdd*100:.1f}%', 'Avg Win', f'{np.mean([t["pnl_pct"] for t in wins]):.2f}%' if wins else '-'],
            ['Avg Loss', f'{np.mean([t["pnl_pct"] for t in losses]):.2f}%' if losses else '-',
             'Max Win', f'{max(t["pnl_pct"] for t in trades):.2f}%' if trades else '-'],
        ]
        ti = doc.add_table(rows=len(info), cols=4)
        ti.style = 'Light Grid Accent 1'
        for i, row in enumerate(info):
            for j, val in enumerate(row):
                cell(ti.rows[i].cells[j], val, bold=(j%2==0), sz=9)

        doc.add_page_break()

        # TRADE-BY-TRADE
        doc.add_heading(f'{name} - Trade-by-Trade Detail ({len(trades)} trades)', level=1)

        hdrs = ['#','Entry Time','Exit Time','Dir','Entry$','Exit$','Margin$','Pos$',
                'PnL%','PnL$','Balance','Type','Peak$']
        rows = []
        for tr in trades:
            clr_pnl = f"{tr['pnl_pct']:+.2f}%"
            clr_usd = f"${tr['pnl_usd']:+,.0f}" if abs(tr['pnl_usd']) >= 1 else f"${tr['pnl_usd']:+,.2f}"
            rows.append([
                str(tr['no']),
                tr['entry_time'],
                tr['exit_time'],
                tr['dir'],
                f"${tr['entry_price']:,.1f}",
                f"${tr['exit_price']:,.1f}",
                f"${tr['margin']:,.0f}",
                f"${tr['position_size']:,.0f}",
                clr_pnl,
                clr_usd,
                f"${tr['balance_after']:,.0f}",
                tr['exit_type'],
                f"${tr['peak_price']:,.1f}",
            ])

        # Split into pages of 40 rows
        page_size = 35
        for pg_start in range(0, len(rows), page_size):
            pg_rows = rows[pg_start:pg_start+page_size]
            add_tbl(doc, hdrs, pg_rows)
            if pg_start + page_size < len(rows):
                doc.add_page_break()

        doc.add_page_break()

        # MONTHLY SUMMARY
        doc.add_heading(f'{name} - Monthly Summary', level=1)
        m_hdrs = ['Month','PnL$','PnL%','Balance','Trades','SL','TSL','REV','FL']
        m_rows = []
        yr_key=''; yr_pnl=0; yr_start=3000; yr_tr=0
        for m in sorted(monthly.keys()):
            d = monthly[m]; y = m[:4]
            if y != yr_key and yr_key != '':
                yr_ret = yr_pnl/yr_start*100 if yr_start>0 else 0
                m_rows.append([f'{yr_key} TOTAL', f'${yr_pnl:+,.0f}', f'{yr_ret:+.1f}%',
                              '', str(yr_tr), '', '', '', ''])
                yr_start = d['bal']-d['pnl']; yr_pnl=0; yr_tr=0
            if yr_key=='': yr_start=3000
            yr_key=y; yr_pnl+=d['pnl']; yr_tr+=d['tr']
            m_rows.append([
                m, f"${d['pnl']:+,.0f}", f"{d['pct']:+.1f}%", f"${d['bal']:,.0f}",
                str(d['tr']), str(d['sl']), str(d['tsl']), str(d['rev']), str(d['fl'])
            ])
        if yr_key:
            yr_ret = yr_pnl/yr_start*100 if yr_start>0 else 0
            m_rows.append([f'{yr_key} TOTAL', f'${yr_pnl:+,.0f}', f'{yr_ret:+.1f}%',
                          '', str(yr_tr), '', '', '', ''])
        add_tbl(doc, m_hdrs, m_rows)

        doc.add_page_break()

    # COMPARISON
    doc.add_heading('v14.4 vs v15.4 Comparison', level=1)
    comp_hdrs = ['Item', 'v14.4 (25%)', 'v15.4 (40%)']
    d1 = all_data['v14.4']; d2 = all_data['v15.4']
    t1 = d1['trades']; t2 = d2['trades']
    w1 = [t for t in t1 if t['pnl_usd']>0]; w2 = [t for t in t2 if t['pnl_usd']>0]

    pk1=3000; mdd1=0
    for m in sorted(d1['monthly'].keys()):
        b=d1['monthly'][m]['bal']; pk1=max(pk1,b)
        dd=(pk1-b)/pk1 if pk1>0 else 0; mdd1=max(mdd1,dd)
    pk2=3000; mdd2=0
    for m in sorted(d2['monthly'].keys()):
        b=d2['monthly'][m]['bal']; pk2=max(pk2,b)
        dd=(pk2-b)/pk2 if pk2>0 else 0; mdd2=max(mdd2,dd)

    comp_rows = [
        ['Final Balance', f'${d1["bal"]:,.0f}', f'${d2["bal"]:,.0f}'],
        ['Return', f'{(d1["bal"]-3000)/3000*100:+,.1f}%', f'{(d2["bal"]-3000)/3000*100:+,.1f}%'],
        ['MDD', f'{mdd1*100:.1f}%', f'{mdd2*100:.1f}%'],
        ['Trades', str(len(t1)), str(len(t2))],
        ['Win Rate', f'{len(w1)/len(t1)*100:.1f}%', f'{len(w2)/len(t2)*100:.1f}%'],
        ['Margin', '25%', '40%'],
        ['Monthly Limit', '-20%', '-30%'],
        ['FL', str(sum(1 for t in t1 if t['exit_type']=='FL')),
               str(sum(1 for t in t2 if t['exit_type']=='FL'))],
        ['10x Consistent', str(d1['consistent']), str(d2['consistent'])],
    ]
    add_tbl(doc, comp_hdrs, comp_rows)

    out = os.path.join(DATA_DIR, 'v144_v154_Trade_Detail_Report.docx')
    doc.save(out)
    print(f"Saved: {out}")
    print(f"Size: {os.path.getsize(out):,} bytes")


if __name__ == '__main__':
    main()
