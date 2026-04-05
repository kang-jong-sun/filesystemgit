"""
전체 기획서 재검토: 45개 × 6엔진 교차검증 + 연도별 상세 + Word 파일 생성
"""
import sys,time,json,os
import pandas as pd,numpy as np
import warnings;warnings.filterwarnings('ignore')
sys.stdout.reconfigure(line_buffering=True)
from docx import Document
from docx.shared import Inches,Pt,Cm,RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

DATA_PATH=r"D:\filesystem\futures\btc_V1\test1\btc_usdt_5m_merged.csv"
OUT=r"D:\filesystem\futures\btc_V1\test1"

# ============ INDICATORS ============
def _ema(d,p):return pd.Series(d).ewm(span=p,adjust=False).mean().values
def _sma(d,p):return pd.Series(d).rolling(p,min_periods=p).mean().values
def _wma(d,p):
    w=np.arange(1,p+1,dtype=float)
    return pd.Series(d).rolling(p).apply(lambda x:np.dot(x,w)/w.sum(),raw=True).values
def _hma(d,p):
    h=max(int(p/2),1);sq=max(int(np.sqrt(p)),1)
    wh=_wma(d,h);wf=_wma(d,p);diff=2*wh-wf
    m=~np.isnan(diff);r=np.full_like(d,np.nan,dtype=float)
    if m.sum()>sq:r[m]=_wma(diff[m],sq)
    return r
def _dema(d,p):e1=_ema(d,p);e2=_ema(e1,p);return 2*e1-e2
def _vwma(cl,vo,p):
    cv=pd.Series(cl*vo).rolling(p,min_periods=p).sum().values
    v=pd.Series(vo).rolling(p,min_periods=p).sum().values
    with np.errstate(divide='ignore',invalid='ignore'):return np.where(v>0,cv/v,np.nan)
def calc_ma(cl,vo,t,p):
    if t=='WMA':return _wma(cl,p)
    elif t=='SMA':return _sma(cl,p)
    elif t=='HMA':return _hma(cl,p)
    elif t=='DEMA':return _dema(cl,p)
    elif t=='VWMA':return _vwma(cl,vo,p)
    return _ema(cl,p)

def adx_wilder(hi,lo,cl,p):
    n=len(cl);tr=np.zeros(n);pdm=np.zeros(n);mdm=np.zeros(n)
    for i in range(1,n):
        tr[i]=max(hi[i]-lo[i],abs(hi[i]-cl[i-1]),abs(lo[i]-cl[i-1]))
        up=hi[i]-hi[i-1];dn=lo[i-1]-lo[i]
        pdm[i]=up if up>dn and up>0 else 0
        mdm[i]=dn if dn>up and dn>0 else 0
    atr=np.zeros(n);pds=np.zeros(n);mds=np.zeros(n)
    if n>p:
        atr[p]=np.mean(tr[1:p+1]);pds[p]=np.mean(pdm[1:p+1]);mds[p]=np.mean(mdm[1:p+1])
        for i in range(p+1,n):
            atr[i]=atr[i-1]*(1-1/p)+tr[i]/p
            pds[i]=pds[i-1]*(1-1/p)+pdm[i]/p
            mds[i]=mds[i-1]*(1-1/p)+mdm[i]/p
    dx=np.zeros(n)
    for i in range(p,n):
        if atr[i]>0:
            pdi=100*pds[i]/atr[i];mdi=100*mds[i]/atr[i];s=pdi+mdi
            dx[i]=100*abs(pdi-mdi)/s if s>0 else 0
    r=np.zeros(n);st=2*p
    if n>st:
        r[st]=np.mean(dx[p:st+1])
        for i in range(st+1,n):r[i]=r[i-1]*(1-1/p)+dx[i]/p
    return r

def adx_ewm(hi,lo,cl,p):
    tr1=hi-lo;tr2=np.abs(hi-np.roll(cl,1));tr3=np.abs(lo-np.roll(cl,1))
    tr1[0]=tr2[0]=tr3[0]=0;tr=np.maximum(np.maximum(tr1,tr2),tr3)
    up=hi-np.roll(hi,1);dn=np.roll(lo,1)-lo;up[0]=dn[0]=0
    pdm=np.where((up>dn)&(up>0),up,0.0);mdm=np.where((dn>up)&(dn>0),dn,0.0)
    a=1.0/p
    av=pd.Series(tr).ewm(alpha=a,min_periods=p,adjust=False).mean().values
    ps=pd.Series(pdm).ewm(alpha=a,min_periods=p,adjust=False).mean().values
    ms=pd.Series(mdm).ewm(alpha=a,min_periods=p,adjust=False).mean().values
    with np.errstate(divide='ignore',invalid='ignore'):
        pdi=np.where(av>0,100*ps/av,0);mdi=np.where(av>0,100*ms/av,0)
        ds=pdi+mdi;dx=np.where(ds>0,100*np.abs(pdi-mdi)/ds,0)
    return pd.Series(dx).ewm(alpha=a,min_periods=p,adjust=False).mean().values

# ============ BACKTEST WITH YEARLY TRACKING ============
def bt(cl,hi,lo,yr,mk,mf,ms,adx_arr,adx_min,delay,sl_pct,ta,tp,lev,mpct,init,
       fee=0.0004,skip_same=True,strict=False,wu=300):
    n=len(cl);bal=init;peak=init;mdd=0.0
    trades=[];in_pos=False;p_dir=0;p_entry=0.0;p_size=0.0
    p_sl=0.0;p_high=0.0;p_low=0.0;t_act=False;t_sl=0.0
    cm=-1;ms2=bal;ml=False;pend=0;pcnt=0
    inv=1.0/lev
    valid=~(np.isnan(mf)|np.isnan(ms));above=np.where(valid,mf>ms,False)

    for i in range(wu,n):
        c=cl[i];h=hi[i];l=lo[i]
        m=mk[i]
        if m!=cm:cm=m;ms2=bal;ml=False
        if not ml and ms2>0:
            if (bal-ms2)/ms2<=-0.20:ml=True
        if in_pos:
            if p_dir==1 and l<=p_entry*(1-inv):
                bal-=p_size*inv;trades.append({'p':-p_size*inv,'r':'FL','y':yr[i]});in_pos=False
                if bal>peak:peak=bal
                dd=(peak-bal)/peak if peak>0 else 0
                if dd>mdd:mdd=dd;continue
            if p_dir==-1 and h>=p_entry*(1+inv):
                bal-=p_size*inv;trades.append({'p':-p_size*inv,'r':'FL','y':yr[i]});in_pos=False
                if bal>peak:peak=bal
                dd=(peak-bal)/peak if peak>0 else 0
                if dd>mdd:mdd=dd;continue
            if p_dir==1 and l<=p_sl:
                pnl=p_size*(p_sl-p_entry)/p_entry-p_size*fee;bal+=pnl
                trades.append({'p':pnl,'r':'SL','y':yr[i]});in_pos=False
            elif p_dir==-1 and h>=p_sl:
                pnl=p_size*(p_entry-p_sl)/p_entry-p_size*fee;bal+=pnl
                trades.append({'p':pnl,'r':'SL','y':yr[i]});in_pos=False
            elif in_pos:
                if p_dir==1:
                    if h>p_high:p_high=h
                    if(p_high-p_entry)/p_entry*lev>=ta*lev:
                        t_act=True;ns=p_high*(1-tp)
                        if ns>t_sl:t_sl=ns
                    if t_act and c<=t_sl:
                        pnl=p_size*(c-p_entry)/p_entry-p_size*fee;bal+=pnl
                        trades.append({'p':pnl,'r':'TSL','y':yr[i]});in_pos=False
                else:
                    if l<p_low:p_low=l
                    if(p_entry-p_low)/p_entry*lev>=ta*lev:
                        t_act=True;ns=p_low*(1+tp)
                        if t_sl==0 or ns<t_sl:t_sl=ns
                    if t_act and c>=t_sl:
                        pnl=p_size*(p_entry-c)/p_entry-p_size*fee;bal+=pnl
                        trades.append({'p':pnl,'r':'TSL','y':yr[i]});in_pos=False
        sig=0
        if i>=wu+1 and valid[i] and valid[i-1]:
            ao=adx_arr[i]>=adx_min if not np.isnan(adx_arr[i]) else False
            if above[i] and not above[i-1] and ao:
                if strict:
                    if i>=2 and above[i]:sig=1
                else:sig=1
            elif not above[i] and above[i-1] and ao:
                if strict:
                    if i>=2 and not above[i]:sig=-1
                else:sig=-1
        if sig!=0 and delay>0:pend=sig;pcnt=delay;sig=0
        if pcnt>0:
            pcnt-=1
            if pcnt==0:
                if pend==1 and valid[i] and mf[i]>ms[i]:sig=pend
                elif pend==-1 and valid[i] and mf[i]<ms[i]:sig=pend
                pend=0
        if sig!=0:
            if in_pos:
                if p_dir!=sig:
                    pnl=p_size*(c-p_entry)/p_entry*p_dir-p_size*fee;bal+=pnl
                    trades.append({'p':pnl,'r':'REV','y':yr[i]});in_pos=False
                elif not skip_same:
                    pnl=p_size*(c-p_entry)/p_entry*p_dir-p_size*fee;bal+=pnl
                    trades.append({'p':pnl,'r':'REV','y':yr[i]});in_pos=False
            if not in_pos and not ml and bal>10:
                mg=bal*mpct;sz=mg*lev;bal-=sz*fee
                p_dir=sig;p_entry=c;p_size=sz
                p_sl=c*(1-sl_pct) if sig==1 else c*(1+sl_pct)
                p_high=c;p_low=c;t_act=False;t_sl=0;in_pos=True
        if bal>peak:peak=bal
        if peak>0:
            dd=(peak-bal)/peak
            if dd>mdd:mdd=dd
    if in_pos:
        c=cl[-1];pnl=p_size*(c-p_entry)/p_entry*p_dir-p_size*fee;bal+=pnl
        trades.append({'p':pnl,'r':'END','y':yr[-1]})

    # Compile with yearly detail
    tc=len(trades)
    if tc==0:return{'bal':bal,'ret':0,'tr':0,'pf':0,'mdd':0,'wr':0,'sl':0,'tsl':0,'rev':0,'fl':0,'yearly':{}}
    pnls=np.array([t['p'] for t in trades])
    w=pnls>0;lo2=pnls<=0
    gp=pnls[w].sum() if w.any() else 0;gl=abs(pnls[lo2].sum()) if lo2.any() else 0.001
    yearly={}
    for t in trades:
        y=int(t['y'])
        if y not in yearly:yearly[y]={'tr':0,'pnl':0.0,'w':0,'l':0,'sl':0,'fl':0}
        yearly[y]['tr']+=1;yearly[y]['pnl']+=t['p']
        if t['p']>0:yearly[y]['w']+=1
        else:yearly[y]['l']+=1
        if t['r']=='SL':yearly[y]['sl']+=1
        if t['r']=='FL':yearly[y]['fl']+=1
    # Per-year PF and MDD (simplified)
    for y in yearly:
        yp=sum(t['p'] for t in trades if int(t['y'])==y and t['p']>0)
        yl=abs(sum(t['p'] for t in trades if int(t['y'])==y and t['p']<=0))
        yearly[y]['pf']=round(min(yp/(yl if yl>0 else 0.001),999.99),2)
    return{
        'bal':round(bal,0),'ret':round((bal-init)/init*100,1),'tr':tc,
        'pf':round(min(gp/gl,999.99),2),'mdd':round(mdd*100,1),
        'wr':round(w.sum()/tc*100,1),
        'sl':sum(1 for t in trades if t['r']=='SL'),
        'tsl':sum(1 for t in trades if t['r']=='TSL'),
        'rev':sum(1 for t in trades if t['r']=='REV'),
        'fl':sum(1 for t in trades if t['r']=='FL'),
        'yearly':yearly,'pnl_usd':round(bal-init,0),
    }

# ============ DATA ============
def load():
    print('Loading...')
    df=pd.read_csv(DATA_PATH,parse_dates=['timestamp']).sort_values('timestamp').reset_index(drop=True)
    data={}
    for nm,mi in[('5min',5),('10min',10),('15min',15),('30min',30),('1h',60)]:
        d=df.set_index('timestamp')
        rule='%dmin'%mi if mi<60 else'%dh'%(mi//60)
        r=df.copy() if mi==5 else d.resample(rule).agg({'open':'first','high':'max','low':'min','close':'last','volume':'sum'}).dropna().reset_index()
        ts=pd.to_datetime(r['timestamp'].values)
        data[nm]={'cl':r['close'].values.astype(float),'hi':r['high'].values.astype(float),
                  'lo':r['low'].values.astype(float),'vo':r['volume'].values.astype(float),
                  'yr':ts.year.values.astype(np.int32),'mk':(ts.year.values*100+ts.month.values).astype(np.int32),'n':len(r)}
        print('  %s:%d'%(nm,data[nm]['n']))
    return data

# ============ STRATEGIES (45 files) ============
S=[
 {'f':'v12.3','mf':'EMA','fp':7,'ms':'EMA','sp':100,'tf':'5min','ap':14,'at':30,'rm':30,'rx':58,'d':0,'sl':0.09,'ta':0.08,'tp':0.06,'l':10,'m':0.20,'i':3000},
 {'f':'v13.5','mf':'EMA','fp':7,'ms':'EMA','sp':100,'tf':'5min','ap':14,'at':30,'rm':30,'rx':58,'d':0,'sl':0.07,'ta':0.08,'tp':0.06,'l':10,'m':0.20,'i':3000},
 {'f':'v14.2_FINAL','mf':'HMA','fp':7,'ms':'EMA','sp':200,'tf':'30min','ap':14,'at':25,'rm':30,'rx':65,'d':3,'sl':0.07,'ta':0.10,'tp':0.01,'l':10,'m':0.30,'i':3000},
 {'f':'v14.4','mf':'EMA','fp':3,'ms':'EMA','sp':200,'tf':'30min','ap':14,'at':35,'rm':30,'rx':65,'d':0,'sl':0.07,'ta':0.06,'tp':0.03,'l':10,'m':0.25,'i':3000},
 {'f':'v15.2','mf':'EMA','fp':3,'ms':'EMA','sp':200,'tf':'30min','ap':14,'at':35,'rm':30,'rx':65,'d':6,'sl':0.05,'ta':0.06,'tp':0.05,'l':10,'m':0.30,'i':3000},
 {'f':'v15.4','mf':'EMA','fp':3,'ms':'EMA','sp':200,'tf':'30min','ap':14,'at':35,'rm':30,'rx':65,'d':0,'sl':0.07,'ta':0.06,'tp':0.03,'l':10,'m':0.40,'i':3000},
 {'f':'v15.5','mf':'EMA','fp':3,'ms':'EMA','sp':200,'tf':'30min','ap':14,'at':35,'rm':35,'rx':65,'d':0,'sl':0.07,'ta':0.06,'tp':0.05,'l':10,'m':0.35,'i':3000},
 {'f':'v15.6_2','mf':'EMA','fp':3,'ms':'EMA','sp':150,'tf':'15min','ap':14,'at':45,'rm':30,'rx':65,'d':0,'sl':0.04,'ta':0.10,'tp':0.03,'l':15,'m':0.40,'i':3000},
 {'f':'v16.0','mf':'WMA','fp':3,'ms':'EMA','sp':200,'tf':'30min','ap':20,'at':35,'rm':35,'rx':65,'d':0,'sl':0.08,'ta':0.04,'tp':0.03,'l':10,'m':0.50,'i':3000},
 {'f':'v16.2_1','mf':'WMA','fp':3,'ms':'EMA','sp':200,'tf':'30min','ap':20,'at':35,'rm':35,'rx':65,'d':0,'sl':0.08,'ta':0.03,'tp':0.02,'l':10,'m':0.50,'i':3000},
 {'f':'v16.2_2','mf':'EMA','fp':3,'ms':'EMA','sp':200,'tf':'30min','ap':14,'at':40,'rm':30,'rx':70,'d':0,'sl':0.08,'ta':0.03,'tp':0.02,'l':10,'m':0.50,'i':3000},
 {'f':'v16.2_3','mf':'DEMA','fp':3,'ms':'EMA','sp':300,'tf':'15min','ap':20,'at':40,'rm':30,'rx':70,'d':2,'sl':0.08,'ta':0.07,'tp':0.02,'l':10,'m':0.25,'i':3000},
 {'f':'v16.2_FINAL_1','mf':'WMA','fp':3,'ms':'EMA','sp':200,'tf':'30min','ap':20,'at':35,'rm':35,'rx':70,'d':0,'sl':0.08,'ta':0.04,'tp':0.03,'l':10,'m':0.50,'i':3000},
 {'f':'v16.2_FINAL_2','mf':'HMA','fp':21,'ms':'EMA','sp':250,'tf':'10min','ap':20,'at':35,'rm':40,'rx':75,'d':0,'sl':0.06,'ta':0.07,'tp':0.03,'l':10,'m':0.40,'i':3000},
 {'f':'v16.4','mf':'WMA','fp':3,'ms':'EMA','sp':200,'tf':'30min','ap':20,'at':35,'rm':35,'rx':65,'d':0,'sl':0.08,'ta':0.04,'tp':0.03,'l':10,'m':0.30,'i':3000},
 {'f':'v16.5','mf':'HMA','fp':21,'ms':'EMA','sp':250,'tf':'10min','ap':20,'at':35,'rm':40,'rx':75,'d':0,'sl':0.06,'ta':0.07,'tp':0.03,'l':10,'m':0.40,'i':3000},
 {'f':'v16.6','mf':'WMA','fp':3,'ms':'EMA','sp':200,'tf':'30min','ap':20,'at':35,'rm':30,'rx':70,'d':5,'sl':0.08,'ta':0.03,'tp':0.02,'l':10,'m':0.50,'i':3000},
 {'f':'v17.0_2','mf':'HMA','fp':21,'ms':'EMA','sp':250,'tf':'10min','ap':20,'at':35,'rm':40,'rx':75,'d':0,'sl':0.06,'ta':0.07,'tp':0.03,'l':10,'m':0.40,'i':3000},
 {'f':'v17.0_3','mf':'WMA','fp':3,'ms':'SMA','sp':300,'tf':'30min','ap':14,'at':40,'rm':30,'rx':70,'d':0,'sl':0.08,'ta':0.04,'tp':0.03,'l':10,'m':0.50,'i':3000},
 {'f':'v17.0_4','mf':'EMA','fp':5,'ms':'EMA','sp':300,'tf':'10min','ap':20,'at':35,'rm':40,'rx':75,'d':0,'sl':0.07,'ta':0.08,'tp':0.03,'l':10,'m':0.40,'i':3000},
 {'f':'v22.0_2','mf':'WMA','fp':3,'ms':'EMA','sp':200,'tf':'30min','ap':20,'at':35,'rm':35,'rx':70,'d':0,'sl':0.08,'ta':0.03,'tp':0.02,'l':10,'m':0.50,'i':3000},
 {'f':'v22.1_2','mf':'HMA','fp':21,'ms':'EMA','sp':250,'tf':'10min','ap':20,'at':35,'rm':40,'rx':75,'d':0,'sl':0.06,'ta':0.07,'tp':0.03,'l':10,'m':0.40,'i':3000},
 {'f':'v22.2','mf':'EMA','fp':3,'ms':'EMA','sp':200,'tf':'30min','ap':14,'at':35,'rm':30,'rx':65,'d':0,'sl':0.07,'ta':0.06,'tp':0.03,'l':10,'m':0.25,'i':3000},
 {'f':'v22.3','mf':'EMA','fp':3,'ms':'EMA','sp':200,'tf':'30min','ap':20,'at':35,'rm':35,'rx':65,'d':0,'sl':0.07,'ta':0.06,'tp':0.03,'l':10,'m':0.35,'i':3000},
 {'f':'v22.6_2','mf':'WMA','fp':3,'ms':'EMA','sp':200,'tf':'30min','ap':20,'at':35,'rm':35,'rx':65,'d':0,'sl':0.08,'ta':0.04,'tp':0.03,'l':10,'m':0.50,'i':3000},
 {'f':'v22.6_3','mf':'HMA','fp':21,'ms':'EMA','sp':250,'tf':'10min','ap':20,'at':35,'rm':40,'rx':75,'d':0,'sl':0.06,'ta':0.07,'tp':0.03,'l':10,'m':0.40,'i':3000},
 {'f':'v22.8','mf':'EMA','fp':5,'ms':'EMA','sp':100,'tf':'5min','ap':14,'at':30,'rm':40,'rx':60,'d':0,'sl':0.04,'ta':0.05,'tp':0.03,'l':10,'m':0.30,'i':3000},
 {'f':'v23.2','mf':'EMA','fp':5,'ms':'EMA','sp':100,'tf':'5min','ap':14,'at':30,'rm':40,'rx':60,'d':0,'sl':0.04,'ta':0.05,'tp':0.03,'l':10,'m':0.30,'i':3000},
 {'f':'v23.4','mf':'HMA','fp':21,'ms':'EMA','sp':250,'tf':'10min','ap':20,'at':35,'rm':40,'rx':75,'d':0,'sl':0.06,'ta':0.07,'tp':0.03,'l':10,'m':0.40,'i':3000},
 {'f':'v23.5','mf':'EMA','fp':5,'ms':'EMA','sp':100,'tf':'5min','ap':14,'at':30,'rm':40,'rx':60,'d':0,'sl':0.04,'ta':0.05,'tp':0.03,'l':10,'m':0.30,'i':3000},
 {'f':'v23.5b','mf':'EMA','fp':5,'ms':'EMA','sp':100,'tf':'5min','ap':14,'at':30,'rm':40,'rx':60,'d':0,'sl':0.04,'ta':0.05,'tp':0.03,'l':10,'m':0.30,'i':3000},
 {'f':'v24.2','mf':'EMA','fp':5,'ms':'EMA','sp':100,'tf':'5min','ap':14,'at':30,'rm':40,'rx':60,'d':0,'sl':0.04,'ta':0.05,'tp':0.03,'l':10,'m':0.30,'i':3000},
 {'f':'v25.0','mf':'EMA','fp':5,'ms':'EMA','sp':100,'tf':'5min','ap':14,'at':30,'rm':40,'rx':60,'d':0,'sl':0.04,'ta':0.05,'tp':0.03,'l':10,'m':0.30,'i':3000},
 {'f':'v25.1','mf':'HMA','fp':21,'ms':'EMA','sp':250,'tf':'10min','ap':20,'at':35,'rm':40,'rx':75,'d':0,'sl':0.06,'ta':0.07,'tp':0.03,'l':10,'m':0.50,'i':3000},
 {'f':'v25.1A','mf':'HMA','fp':21,'ms':'EMA','sp':250,'tf':'10min','ap':20,'at':35,'rm':40,'rx':75,'d':0,'sl':0.06,'ta':0.07,'tp':0.03,'l':10,'m':0.50,'i':3000},
 {'f':'v25.1C','mf':'EMA','fp':5,'ms':'EMA','sp':100,'tf':'5min','ap':14,'at':30,'rm':40,'rx':60,'d':0,'sl':0.04,'ta':0.05,'tp':0.03,'l':10,'m':0.30,'i':3000},
 {'f':'v25.2_1','mf':'EMA','fp':3,'ms':'EMA','sp':200,'tf':'30min','ap':20,'at':35,'rm':40,'rx':75,'d':0,'sl':0.07,'ta':0.07,'tp':0.03,'l':10,'m':0.40,'i':3000},
 {'f':'v25.2_3','mf':'EMA','fp':5,'ms':'EMA','sp':300,'tf':'10min','ap':20,'at':35,'rm':40,'rx':75,'d':0,'sl':0.07,'ta':0.08,'tp':0.03,'l':10,'m':0.40,'i':3000},
 {'f':'v26.0_1','mf':'EMA','fp':3,'ms':'SMA','sp':300,'tf':'30min','ap':14,'at':40,'rm':30,'rx':70,'d':0,'sl':0.08,'ta':0.04,'tp':0.03,'l':10,'m':0.50,'i':3000},
 {'f':'v26.0_2','mf':'HMA','fp':21,'ms':'EMA','sp':250,'tf':'10min','ap':20,'at':35,'rm':40,'rx':75,'d':0,'sl':0.06,'ta':0.07,'tp':0.03,'l':10,'m':0.40,'i':3000},
 {'f':'v27_1','mf':'WMA','fp':3,'ms':'EMA','sp':200,'tf':'30min','ap':20,'at':35,'rm':30,'rx':70,'d':5,'sl':0.08,'ta':0.03,'tp':0.02,'l':10,'m':0.50,'i':3000},
 {'f':'v27_2','mf':'HMA','fp':21,'ms':'EMA','sp':250,'tf':'10min','ap':20,'at':35,'rm':40,'rx':75,'d':0,'sl':0.06,'ta':0.07,'tp':0.03,'l':10,'m':0.40,'i':3000},
 {'f':'v28.0_2','mf':'EMA','fp':5,'ms':'EMA','sp':100,'tf':'5min','ap':14,'at':30,'rm':40,'rx':60,'d':0,'sl':0.04,'ta':0.05,'tp':0.03,'l':10,'m':0.30,'i':3000},
 {'f':'v28.0_3','mf':'HMA','fp':21,'ms':'EMA','sp':250,'tf':'10min','ap':20,'at':35,'rm':40,'rx':75,'d':0,'sl':0.06,'ta':0.07,'tp':0.03,'l':10,'m':0.40,'i':3000},
 {'f':'v32.2','mf':'WMA','fp':3,'ms':'EMA','sp':200,'tf':'30min','ap':20,'at':35,'rm':30,'rx':70,'d':5,'sl':0.08,'ta':0.03,'tp':0.02,'l':10,'m':0.50,'i':3000},
 {'f':'v32.3','mf':'HMA','fp':21,'ms':'EMA','sp':250,'tf':'10min','ap':20,'at':35,'rm':40,'rx':75,'d':0,'sl':0.06,'ta':0.07,'tp':0.03,'l':10,'m':0.40,'i':3000},
]

ENGINES=[
 {'id':'E1','n':'Wilder','adx':'w','fee':0.0004,'slip':0,'strict':False,'skip':True},
 {'id':'E2','n':'EWM','adx':'e','fee':0.0004,'slip':0,'strict':False,'skip':True},
 {'id':'E3','n':'Strict','adx':'w','fee':0.0004,'slip':0,'strict':True,'skip':True},
 {'id':'E4','n':'Slip','adx':'w','fee':0.0004,'slip':0.0005,'strict':False,'skip':True},
 {'id':'E5','n':'HiFee','adx':'w','fee':0.0006,'slip':0,'strict':False,'skip':True},
 {'id':'E6','n':'Reentry','adx':'w','fee':0.0004,'slip':0,'strict':False,'skip':False},
]

def main():
    t0=time.time()
    print('='*70)
    print('  FULL REVIEW: %d files x %d engines = %d backtests'%(len(S),len(ENGINES),len(S)*len(ENGINES)))
    print('='*70)
    data=load()
    results=[]
    cnt=0;total=len(S)*len(ENGINES)

    for s in S:
        tf=s['tf'];d=data[tf]
        cl=d['cl'];hi=d['hi'];lo=d['lo'];vo=d['vo'];yr=d['yr'];mk=d['mk']
        mf=calc_ma(cl,vo,s['mf'],s['fp']);ms=calc_ma(cl,vo,s['ms'],s['sp'])
        wu=max(s['sp']+50,300)
        row={'f':s['f'],'desc':'%s(%d)/%s(%d) %s L%dx M%d%%'%(s['mf'],s['fp'],s['ms'],s['sp'],s['tf'],s['l'],s['m']*100)}
        eng_rets=[];eng_pfs=[];eng_mdds=[];eng_trs=[]

        for eng in ENGINES:
            ax=adx_wilder(hi,lo,cl,s['ap']) if eng['adx']=='w' else adx_ewm(hi,lo,cl,s['ap'])
            if eng['slip']>0:
                np.random.seed(42)
                sl2=1+np.random.uniform(-eng['slip'],eng['slip'],len(cl))
                cl2=cl*sl2;hi2=hi*sl2;lo2=lo*sl2
            else:cl2=cl;hi2=hi;lo2=lo
            r=bt(cl2,hi2,lo2,yr,mk,mf,ms,ax,s['at'],s['d'],s['sl'],s['ta'],s['tp'],s['l'],s['m'],s['i'],
                 fee=eng['fee'],skip_same=eng['skip'],strict=eng['strict'],wu=wu)
            row[eng['id']]=r
            eng_rets.append(r['ret']);eng_pfs.append(r['pf']);eng_mdds.append(r['mdd']);eng_trs.append(r['tr'])
            cnt+=1
            if cnt%30==0:print('  %d/%d (%.0fs)'%(cnt,total,time.time()-t0))

        row['avg_ret']=round(np.mean(eng_rets),1)
        row['min_ret']=round(np.min(eng_rets),1)
        row['avg_pf']=round(np.mean(eng_pfs),2)
        row['min_pf']=round(np.min(eng_pfs),2)
        row['avg_mdd']=round(np.mean(eng_mdds),1)
        row['max_mdd']=round(np.max(eng_mdds),1)
        row['avg_tr']=round(np.mean(eng_trs),0)
        row['init']=s['i']
        results.append(row)

    # Rankings
    profitable=[r for r in results if r['min_ret']>0]
    by_ret=sorted(profitable,key=lambda x:x['avg_ret'],reverse=True)
    by_stable=sorted(profitable,key=lambda x:x['max_mdd'])
    discard=sorted([r for r in results if r['min_ret']<=0 or r['max_mdd']>50 or r['min_pf']<1.0],key=lambda x:x['min_ret'])

    # Print summary
    print('\n=== RETURN BEST 10 ===')
    for i,r in enumerate(by_ret[:10]):
        print('%2d. %s | %s | Avg:%+.1f%% Min:%+.1f%% PF:%.2f MDD:%.1f%% Tr:%.0f'%(
            i+1,r['f'],r['desc'],r['avg_ret'],r['min_ret'],r['avg_pf'],r['avg_mdd'],r['avg_tr']))
    print('\n=== STABILITY BEST 10 ===')
    for i,r in enumerate(by_stable[:10]):
        print('%2d. %s | MaxMDD:%.1f%% | Avg:%+.1f%% PF:%.2f Tr:%.0f'%(
            i+1,r['f'],r['max_mdd'],r['avg_ret'],r['avg_pf'],r['avg_tr']))
    print('\n=== DISCARD 10 ===')
    for i,r in enumerate(discard[:10]):
        print('%2d. %s | MinRet:%+.1f%% MinPF:%.2f MaxMDD:%.1f%%'%(
            i+1,r['f'],r['min_ret'],r['min_pf'],r['max_mdd']))

    # ============ WORD DOCUMENT ============
    print('\nGenerating Word document...')
    doc=Document()
    doc.add_heading('BTC/USDT 선물 자동매매 기획서 전체 재검토 보고서',0)
    doc.add_paragraph('검증일: 2026-04-01 | 데이터: 75개월 (655,399개 5분봉) | 6엔진 교차검증')
    doc.add_paragraph('엔진: Wilder ADX / EWM ADX / Strict Entry / Slippage 0.05% / High Fee 0.06% / Same-Dir Reentry')

    def add_ranking_table(doc,title,ranked,category):
        doc.add_heading(title,1)
        years=[2020,2021,2022,2023,2024,2025,2026]
        cols=['순위','파일명','손익률','손익금액']
        for y in years:cols.append('%d거래'%y)
        cols+=['총거래']
        for y in years:cols.append('%dMDD'%y)
        for y in years:cols.append('%dPF'%y)
        cols+=['6엔진일치사유','비고']

        tbl=doc.add_table(rows=1+len(ranked),cols=len(cols),style='Table Grid')
        tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
        # Header
        for j,c in enumerate(cols):
            cell=tbl.rows[0].cells[j]
            cell.text=c
            for p in cell.paragraphs:
                p.alignment=WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:run.font.size=Pt(7);run.font.bold=True

        for i,r in enumerate(ranked):
            row_cells=tbl.rows[i+1].cells
            # Use E1 (Wilder) as primary for yearly data
            e1=r.get('E1',{})
            yrly=e1.get('yearly',{})

            row_cells[0].text=str(i+1)
            row_cells[1].text=r['f']
            row_cells[2].text='%+.1f%%'%r['avg_ret']
            row_cells[3].text='$%s'%'{:,.0f}'.format(r['avg_ret']/100*r['init'])

            idx=4
            total_tr=0
            for y in years:
                yd=yrly.get(y,{})
                tr=yd.get('tr',0)
                total_tr+=tr
                row_cells[idx].text=str(tr) if tr>0 else '-'
                idx+=1
            row_cells[idx].text=str(total_tr);idx+=1

            for y in years:
                row_cells[idx].text='-';idx+=1  # MDD per year (simplified)
            for y in years:
                yd=yrly.get(y,{})
                pf=yd.get('pf',0)
                row_cells[idx].text='%.1f'%pf if pf>0 else '-'
                idx+=1

            # 6엔진 일치 사유
            all_profit=r['min_ret']>0
            reasons=[]
            if all_profit:reasons.append('6엔진전수익')
            else:reasons.append('일부손실')
            if r['max_mdd']<=30:reasons.append('MDD<=30%%')
            if r['avg_pf']>=5:reasons.append('PF>=5')
            row_cells[idx].text=', '.join(reasons);idx+=1

            # 비고
            notes=[]
            if category=='discard':
                if r['min_ret']<=0:notes.append('손실발생')
                if r['max_mdd']>50:notes.append('MDD>50%%')
                if r['min_pf']<1:notes.append('PF<1')
            else:
                notes.append(r['desc'][:30])
            row_cells[idx].text=', '.join(notes)

            # Font size
            for j in range(len(cols)):
                for p in row_cells[j].paragraphs:
                    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:run.font.size=Pt(6)

    add_ranking_table(doc,'수익률 BEST 10',by_ret[:10],'return')
    add_ranking_table(doc,'안정형 BEST 10',by_stable[:10],'stable')
    add_ranking_table(doc,'폐기형 BEST 10',discard[:10],'discard')

    # Summary page
    doc.add_heading('전체 45개 기획서 6엔진 교차검증 요약',1)
    stbl=doc.add_table(rows=1+len(results),cols=10,style='Table Grid')
    hdr=['파일명','전략','평균수익','최소수익','평균PF','최소PF','평균MDD','최대MDD','평균거래','판정']
    for j,h in enumerate(hdr):
        stbl.rows[0].cells[j].text=h
        for p in stbl.rows[0].cells[j].paragraphs:
            for run in p.runs:run.font.size=Pt(7);run.font.bold=True
    for i,r in enumerate(results):
        c=stbl.rows[i+1].cells
        c[0].text=r['f'];c[1].text=r['desc'][:25]
        c[2].text='%+.1f%%'%r['avg_ret'];c[3].text='%+.1f%%'%r['min_ret']
        c[4].text='%.2f'%r['avg_pf'];c[5].text='%.2f'%r['min_pf']
        c[6].text='%.1f%%'%r['avg_mdd'];c[7].text='%.1f%%'%r['max_mdd']
        c[8].text='%.0f'%r['avg_tr']
        if r['min_ret']>0 and r['max_mdd']<=30:c[9].text='추천'
        elif r['min_ret']>0:c[9].text='주의(MDD)'
        elif r['min_ret']<=0:c[9].text='폐기'
        for j in range(10):
            for p in c[j].paragraphs:
                for run in p.runs:run.font.size=Pt(6)

    fp=os.path.join(OUT,'BTC_기획서_전체재검토_6엔진교차검증.docx')
    doc.save(fp)
    print('Word saved: %s'%fp)
    print('\n'+'='*70)
    print('  COMPLETE: %d backtests in %.1f min'%(cnt,(time.time()-t0)/60))
    print('='*70)

if __name__=='__main__':
    main()
